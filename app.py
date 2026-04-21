import logging
import json
import requests
import io
import os
import pandas as pd
import re
from flask import Flask, request, jsonify, render_template, send_file, session, redirect, url_for, abort, g
from config import Config
from models import db, Message, MessageStatus, Contact, Tag, contact_tags, Campaign, CampaignLog, ConversationTopic, ConversationSession, RagDocument, ChatbotConfig, ConversationNote, AutoTagRule, AutoTagLog, FollowUpSequence, FollowUpStep, FollowUpEnrollment, CrmUserTagVisibility, CatalogProduct, Order, OrderItem
import threading
import time as time_module
from event_handlers import process_event
from sqlalchemy import func, or_, and_, text
from sqlalchemy.orm import joinedload
from datetime import datetime, timedelta, timezone
import logging
import pytz
import mimetypes
import uuid

# Fix MIME types for Windows/Local
mimetypes.add_type('audio/ogg', '.oga')
mimetypes.add_type('audio/ogg', '.ogg')
mimetypes.add_type('audio/ogg', '.opus')

app = Flask(__name__)
logger = logging.getLogger(__name__)

# Zona horaria de Argentina
ARGENTINA_TZ = pytz.timezone('America/Argentina/Buenos_Aires')

# Filtro Jinja2 para convertir UTC a hora Argentina
WHATSAPP_ERROR_MESSAGES = {
    131026: "El mensaje no pudo ser entregado al destinatario.",
    131047: "El usuario no inició una conversación recientemente (mensaje fuera de ventana).",
    130429: "Se alcanzó el límite de velocidad de envío. Intente más tarde.",
    131021: "El número de remitente no coincide con el destinatario.",
    131031: "El remitente no está en la cuenta de negocios.",
    130472: "El número del usuario es parte de un experimento de Meta.",
    131000: "Fallo genérico al enviar el mensaje.",
    131051: "Tipo de mensaje no soportado.",
    132000: "El número de parámetros de la plantilla no coincide.",
    132001: "La plantilla no existe o fue eliminada.",
    132007: "La plantilla contiene contenido no permitido por Meta.",
    133000: "El número no está registrado en WhatsApp.",
    133004: "El servidor de WhatsApp no está disponible temporalmente.",
    133010: "El número no está registrado o fue desactivado.",
    135000: "Error genérico del servidor de Meta.",
}

@app.template_filter('format_whatsapp_error')
def format_whatsapp_error_filter(error_details):
    """Convierte el JSON de error de WhatsApp a texto legible."""
    if not error_details:
        return ''
    try:
        import json as _json
        errors = _json.loads(error_details) if isinstance(error_details, str) else error_details
        if isinstance(errors, list) and errors:
            err = errors[0]
            code = err.get('code')
            friendly = WHATSAPP_ERROR_MESSAGES.get(code)
            if friendly:
                return friendly
            # Fallback: usar el mensaje del API
            msg = err.get('message') or err.get('title') or ''
            details = (err.get('error_data') or {}).get('details', '')
            return details or msg or str(error_details)
    except Exception:
        pass
    return str(error_details)

@app.template_filter('to_argentina')
def to_argentina_filter(dt):
    """Convierte datetime UTC a hora de Argentina."""
    if dt is None:
        return ''
    # Si el datetime es naive, asumir que es UTC
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(ARGENTINA_TZ)

# Configuración de la base de datos
app.config['SQLALCHEMY_DATABASE_URI'] = Config.DATABASE_URL
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = Config.SECRET_KEY
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,   # Detecta conexiones muertas antes de usarlas
    'pool_recycle': 300,     # Recicla conexiones cada 5 min (evita timeout de Supabase)
}

# Inicializar SQLAlchemy
db.init_app(app)

# Crear tablas al iniciar
with app.app_context():
    db.create_all()

# Importar servicio de WhatsApp (después de crear app)
from whatsapp_service import whatsapp_api, init_all_buckets

# Inicializar buckets de MinIO al arrancar
with app.app_context():
    init_all_buckets()

# Crear tag del sistema "Asistencia Humana" al arrancar
# Envuelto en try/except: si is_system no existe aún (pre-migración), no falla
with app.app_context():
    try:
        system_tag = Tag.query.filter_by(name='Asistencia Humana').first()
        if not system_tag:
            system_tag = Tag(name='Asistencia Humana', color='red', is_system=True, is_active=True)
            db.session.add(system_tag)
            db.session.commit()
            logger.info("System tag 'Asistencia Humana' created")
        elif not system_tag.is_system:
            system_tag.is_system = True
            db.session.commit()
            logger.info("System tag 'Asistencia Humana' marked as system")
    except Exception as e:
        db.session.rollback()
        logger.warning(f"Could not ensure system tag (run migrate_human_assistance.py first): {e}")

# Rutas públicas que no requieren autenticación
PUBLIC_PATHS = {'/', '/login', '/logout', '/webhook', '/chatwoot-webhook', '/api/minio/diagnose', '/sw.js', '/static/manifest.json', '/api/whatsapp/send-text', '/api/whatsapp/send-media', '/api/bot/catalog', '/api/bot/audios', '/api/bot/send-audio'}

def ensure_admin_exists():
    """Crea el usuario admin inicial desde .env si no hay usuarios en la BD."""
    from models import CrmUser
    try:
        db.session.rollback()  # Limpiar cualquier transacción pendiente
        if CrmUser.query.count() == 0:
            admin = CrmUser(
                username='admin',
                display_name='Administrador',
                is_admin=True,
                is_active=True
            )
            admin.set_password(Config.LOGIN_PASSWORD or 'admin')
            db.session.add(admin)
            db.session.commit()
            logger.info("✅ Usuario admin creado automáticamente desde .env")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creando admin inicial: {e}")

# Mapa de permisos requeridos por ruta
ROUTE_PERMISSIONS = {
    '/dashboard': 'dashboard',
    '/contacts': 'contacts',
    '/tags': 'tags',
    '/campaigns': 'campaigns',
    '/failed-messages': 'failed_messages',
    '/analytics': 'analytics',
    '/sessions': 'sessions',
    '/topics': 'topics',
    '/chatbot': 'chatbot',
    '/whatsapp-settings': 'settings',
}

@app.before_request
def check_auth():
    if request.path in PUBLIC_PATHS or request.path.startswith('/static/') or request.path.startswith('/media/'):
        return None
    if request.path.startswith('/api/rag/documents/') and request.path.endswith('/status'):
        return None
    if request.path == '/api/extract-docx':
        return None
    if request.path.startswith('/api/contact/') and request.path.endswith('/bot-status'):
        return None
    if request.path == '/api/escalate-to-human':
        return None
    if request.path == '/api/contacts/assign-tag':
        return None

    if not session.get('user_id'):
        if request.path.startswith('/api/'):
            return jsonify({'error': 'Unauthorized'}), 401
        return redirect(url_for('login'))

    # Verificar permiso para la ruta actual
    from models import CrmUser
    user = CrmUser.query.get(session['user_id'])
    if not user or not user.is_active:
        session.clear()
        return redirect(url_for('login'))

    # Admin panel solo para admins
    if request.path.startswith('/admin') and not user.is_admin:
        return jsonify({'error': 'Forbidden'}), 403 if request.path.startswith('/api/') else abort(403)

    # Chequear permiso de sección
    for route_prefix, permission in ROUTE_PERMISSIONS.items():
        if request.path == route_prefix or request.path.startswith(route_prefix + '/'):
            if not user.has_permission(permission):
                if request.path.startswith('/api/'):
                    return jsonify({'error': 'Forbidden'}), 403
                return render_template('403.html'), 403
            break

    # Inyectar usuario en g para uso en templates y vistas
    g.current_user = user

@app.route("/sw.js")
def service_worker():
    """Serve service worker from root for proper scope."""
    import os
    return send_file(
        os.path.join(os.path.dirname(__file__), 'static', 'service-worker.js'),
        mimetype='application/javascript'
    )

@app.context_processor
def inject_current_user():
    """Inyecta el usuario actual en todos los templates."""
    from models import CrmUser
    user_id = session.get('user_id')
    if user_id:
        user = CrmUser.query.get(user_id)
        if user:
            return {'current_user': user}
    return {'current_user': None}

@app.route("/", methods=["GET"])
def index():
    return redirect(url_for('dashboard'))

@app.route("/login", methods=["GET"])
def login():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    return render_template('login.html', error=None)

@app.route("/login", methods=["POST"])
def login_post():
    from models import CrmUser
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')

    ensure_admin_exists()

    user = CrmUser.query.filter_by(username=username, is_active=True).first()
    if user and user.check_password(password):
        session['user_id'] = user.id
        session['username'] = user.username
        session['display_name'] = user.display_name
        session['is_admin'] = user.is_admin
        return redirect(url_for('dashboard'))
    return render_template('login.html', error='Usuario o contraseña incorrectos.')

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route("/webhook", methods=["GET"])
def verify_webhook():
    """
    Endpoint de verificación para Meta (Facebook).
    Meta enviará un GET request con hub.mode, hub.verify_token y hub.challenge.
    """
    mode = request.args.get("hub.mode")
    token = request.args.get("hub.verify_token")
    challenge = request.args.get("hub.challenge")

    if mode and token:
        if mode == "subscribe" and token == Config.VERIFY_TOKEN:
            logger.info("WEBHOOK_VERIFIED")
            return challenge, 200
        else:
            logger.error("Verificación fallida. Token recibido no coincide.")
            return "Verification token mismatch", 403
    
    return "Hello world", 200

@app.route("/media/<path:filename>")
def media_proxy(filename):
    """
    Proxy para servir archivos de MinIO o almacenamiento local.
    Primero intenta MinIO, si falla busca en static/media/.
    """
    from whatsapp_service import get_s3_client, ensure_bucket_exists

    # Intentar MinIO primero
    try:
        s3 = get_s3_client()
        bucket = Config.MINIO_BUCKET

        if bucket:
            response = s3.get_object(Bucket=bucket, Key=filename)
            content_type = response.get('ContentType', 'application/octet-stream')
            return send_file(
                io.BytesIO(response['Body'].read()),
                mimetype=content_type,
                download_name=filename
            )
    except Exception as e:
        logger.warning(f"MinIO failed for {filename}, trying local: {str(e)}")

    # Fallback: buscar en static/media/
    local_path = os.path.join(os.path.dirname(__file__), 'static', 'media', filename)
    if os.path.exists(local_path):
        return send_file(local_path, download_name=filename)

    logger.error(f"File not found: {filename}")
    return "File not found", 404


@app.route("/api/media/retry/<int:message_id>", methods=["POST"])
def api_retry_media_download(message_id):
    """
    Re-intenta descargar un archivo multimedia que falló.
    Útil para corregir mensajes con media_url inválido.
    """
    try:
        msg = Message.query.get(message_id)
        if not msg:
            return jsonify({'error': 'Mensaje no encontrado'}), 404

        if not msg.media_id:
            return jsonify({'error': 'El mensaje no tiene media_id'}), 400

        # Re-descargar el archivo
        new_url = whatsapp_api.download_media(msg.media_id)

        if not new_url:
            return jsonify({'error': 'No se pudo descargar el archivo'}), 500

        # Actualizar la URL en la BD
        msg.media_url = new_url
        db.session.commit()

        return jsonify({
            'success': True,
            'message_id': message_id,
            'new_url': new_url
        })

    except Exception as e:
        logger.error(f"Error re-descargando media {message_id}: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/media/fix-extensions", methods=["POST"])
def api_fix_media_extensions():
    """
    Corrige extensiones de audio de .oga a .ogg en la BD.
    NO re-descarga archivos, solo actualiza las URLs.
    """
    try:
        # Buscar mensajes con extensión .oga
        oga_msgs = Message.query.filter(
            Message.media_url.like('%.oga')
        ).all()

        fixed = 0
        results = []

        for msg in oga_msgs:
            old_url = msg.media_url
            # Cambiar .oga a .ogg
            new_url = old_url.replace('.oga', '.ogg')
            msg.media_url = new_url
            fixed += 1
            results.append({
                'id': msg.id,
                'old_url': old_url,
                'new_url': new_url
            })

        db.session.commit()

        return jsonify({
            'success': True,
            'total_fixed': fixed,
            'details': results
        })

    except Exception as e:
        logger.error(f"Error fixing extensions: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/media/fix-paths", methods=["POST"])
def api_fix_media_paths():
    """
    Corrige URLs de static/media/ a /media/ y re-descarga archivos a MinIO.
    """
    try:
        # Buscar mensajes con URLs que empiezan con static/media/
        broken_msgs = Message.query.filter(
            Message.media_id.isnot(None),
            Message.media_url.like('static/media/%')
        ).all()

        fixed = 0
        failed = 0
        results = []

        for msg in broken_msgs:
            try:
                # Re-descargar el archivo a MinIO
                new_url = whatsapp_api.download_media(msg.media_id)
                if new_url:
                    msg.media_url = new_url
                    fixed += 1
                    results.append({'id': msg.id, 'status': 'fixed', 'url': new_url})
                else:
                    failed += 1
                    results.append({'id': msg.id, 'status': 'failed', 'error': 'Download returned None'})
            except Exception as e:
                failed += 1
                results.append({'id': msg.id, 'status': 'failed', 'error': str(e)})

        db.session.commit()

        return jsonify({
            'success': True,
            'total_found': len(broken_msgs),
            'fixed': fixed,
            'failed': failed,
            'details': results
        })

    except Exception as e:
        logger.error(f"Error fixing paths: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/media/fix-broken", methods=["POST"])
def api_fix_broken_media():
    """
    Identifica y corrige mensajes con media_url inválido.
    Re-descarga archivos desde WhatsApp.
    """
    try:
        from sqlalchemy import or_

        # Buscar mensajes con URLs problemáticas (NULL, sin ruta, o con extensión .oga)
        broken_msgs = Message.query.filter(
            Message.media_id.isnot(None),
            or_(
                Message.media_url.is_(None),
                ~Message.media_url.like('/%'),
                ~Message.media_url.like('http%'),
                Message.media_url.like('%.oga')  # También incluir .oga para re-descargar
            )
        ).all()

        fixed = 0
        failed = 0
        results = []

        for msg in broken_msgs:
            try:
                new_url = whatsapp_api.download_media(msg.media_id)
                if new_url:
                    msg.media_url = new_url
                    fixed += 1
                    results.append({'id': msg.id, 'status': 'fixed', 'url': new_url})
                else:
                    failed += 1
                    results.append({'id': msg.id, 'status': 'failed', 'error': 'Download returned None'})
            except Exception as e:
                failed += 1
                results.append({'id': msg.id, 'status': 'failed', 'error': str(e)})

        db.session.commit()

        return jsonify({
            'success': True,
            'total_broken': len(broken_msgs),
            'fixed': fixed,
            'failed': failed,
            'details': results
        })

    except Exception as e:
        logger.error(f"Error fixing broken media: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/geocode", methods=["GET"])
def api_geocode():
    """Resuelve coordenadas directas, URL de Google Maps, plus code o dirección."""
    import re, urllib.request, urllib.parse, json as _json
    q = request.args.get("q", "").strip()
    logger.info(f"[geocode] query recibida: '{q}'")
    if not q:
        return jsonify({"error": "q requerido"}), 400
    try:
        # 1. Coordenadas directas: "-24.7821, -65.4232"
        coord_match = re.match(r'^(-?\d{1,3}\.?\d*)\s*,\s*(-?\d{1,3}\.?\d*)$', q)
        if coord_match:
            lat, lng = float(coord_match.group(1)), float(coord_match.group(2))
            logger.info(f"[geocode] coordenadas directas: {lat}, {lng}")
            return jsonify({"lat": round(lat, 7), "lng": round(lng, 7), "label": f"{lat}, {lng}"})

        # 2. URL de Google Maps → seguir redirección y extraer coordenadas de la URL final
        if q.startswith("http"):
            logger.info(f"[geocode] detectada URL de Google Maps, siguiendo redirección...")
            req = urllib.request.Request(q, headers={"User-Agent": "Mozilla/5.0"})
            try:
                with urllib.request.urlopen(req, timeout=8) as r:
                    final_url = r.geturl()
            except Exception as url_err:
                # En algunos casos la redirección falla pero la URL final queda en el error
                final_url = str(url_err)
            logger.info(f"[geocode] URL final: {final_url}")
            # 1er intento: !3d{lat}!4d{lng} = pin real del lugar (más preciso)
            # Tomar el ÚLTIMO par porque el primero puede ser la ubicación del usuario
            all_pins = re.findall(r'!3d(-?\d+\.?\d+)!4d(-?\d+\.?\d+)', final_url)
            coord_in_url = None
            if all_pins:
                last_pin = all_pins[-1]
                logger.info(f"[geocode] pines encontrados: {all_pins}, usando último: {last_pin}")
                # Crear objeto compatible con .group()
                class _Match:
                    def __init__(self, g1, g2): self._g = [None, g1, g2]
                    def group(self, n): return self._g[n]
                coord_in_url = _Match(last_pin[0], last_pin[1])
            if not coord_in_url:
                # 2do intento: @lat,lng = centro del viewport
                coord_in_url = re.search(r'@(-?\d+\.?\d+),(-?\d+\.?\d+)', final_url)
            if not coord_in_url:
                # Intentar formato ?q=lat,lng
                coord_in_url = re.search(r'[?&]q=(-?\d+\.?\d+),(-?\d+\.?\d+)', final_url)
            if not coord_in_url:
                # Intentar formato /place/lat,lng
                coord_in_url = re.search(r'place/(-?\d+\.?\d+),(-?\d+\.?\d+)', final_url)
            if coord_in_url:
                lat, lng = float(coord_in_url.group(1)), float(coord_in_url.group(2))
                logger.info(f"[geocode] coordenadas extraídas de URL: {lat}, {lng}")
                return jsonify({"lat": round(lat, 7), "lng": round(lng, 7), "label": "Google Maps"})
            logger.warning(f"[geocode] no se pudieron extraer coordenadas de la URL")
            return jsonify({"error": "No se pudieron extraer coordenadas de esa URL de Google Maps"}), 404

        # 3. Plus code: "6HFG+R6Q Salta"
        plus_code_pattern = re.compile(
            r'^([23456789CFGHJMPQRVWX]{2,8}\+[23456789CFGHJMPQRVWX]*)\s*(.*)?$', re.IGNORECASE
        )
        pc_match = plus_code_pattern.match(q)
        logger.info(f"[geocode] ¿es plus code? {'sí' if pc_match else 'no'}")

        if pc_match:
            from openlocationcode import openlocationcode as olc
            code = pc_match.group(1).upper()
            ref = (pc_match.group(2) or "").strip()
            logger.info(f"[geocode] plus code: '{code}' | referencia: '{ref}'")
            full_code = code
            is_full = olc.isFull(code)
            logger.info(f"[geocode] isFull: {is_full}")
            if not is_full:
                if not ref:
                    return jsonify({"error": "Código corto requiere ciudad de referencia (ej: 6HFG+R6Q Salta)"}), 400
                nom_url = f"https://nominatim.openstreetmap.org/search?q={urllib.parse.quote(ref)}&format=json&limit=5&countrycodes=ar"
                logger.info(f"[geocode] buscando referencia: {nom_url}")
                req = urllib.request.Request(nom_url, headers={"User-Agent": "WhatsAppCRM/1.0"})
                with urllib.request.urlopen(req, timeout=5) as r:
                    city_data = _json.loads(r.read())
                logger.info(f"[geocode] resultados: {[{'name': r.get('name'), 'place_rank': r.get('place_rank')} for r in city_data]}")
                if not city_data:
                    return jsonify({"error": f"No se encontró '{ref}' como ciudad de referencia"}), 404
                best = city_data[0]
                ref_lat, ref_lng = float(best["lat"]), float(best["lon"])
                logger.info(f"[geocode] referencia: {ref_lat}, {ref_lng}")
                full_code = olc.recoverNearest(code, ref_lat, ref_lng)
                logger.info(f"[geocode] full code: '{full_code}'")
            decoded = olc.decode(full_code)
            result = {"lat": round(decoded.latitudeCenter, 7), "lng": round(decoded.longitudeCenter, 7),
                      "label": f"Plus code: {code}" + (f" ({ref})" if ref else "")}
            logger.info(f"[geocode] resultado plus code: {result}")
            return jsonify(result)

        # 4. Dirección normal — Nominatim
        nom_url = f"https://nominatim.openstreetmap.org/search?q={urllib.parse.quote(q)}&format=json&limit=1&countrycodes=ar"
        logger.info(f"[geocode] dirección normal, Nominatim: {nom_url}")
        req = urllib.request.Request(nom_url, headers={"User-Agent": "WhatsAppCRM/1.0"})
        with urllib.request.urlopen(req, timeout=5) as r:
            data = _json.loads(r.read())
        logger.info(f"[geocode] respuesta Nominatim: {data}")
        if not data:
            return jsonify({"error": "No se encontraron coordenadas para esa dirección"}), 404
        label = ", ".join(data[0]["display_name"].split(",")[:3])
        result = {"lat": round(float(data[0]["lat"]), 7), "lng": round(float(data[0]["lon"]), 7), "label": label}
        logger.info(f"[geocode] resultado dirección: {result}")
        return jsonify(result)

    except Exception as e:
        logger.error(f"[geocode] excepción para '{q}': {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/orders/fix-message-content", methods=["POST"])
def api_fix_order_message_content():
    """
    Migración: actualiza mensajes de tipo 'order' que tengan retailer_id crudo en su contenido.
    Reemplaza IDs no reconocibles por 'Artículo'.
    """
    import re
    try:
        order_msgs = Message.query.filter_by(message_type='order').all()
        fixed = 0

        for msg in order_msgs:
            if not msg.content:
                continue
            # Extraer el interior: "[Pedido: X ×N, Y ×M]"
            m = re.match(r'^\[Pedido: (.+)\]$', msg.content)
            if not m:
                continue
            inner = m.group(1)
            parts = inner.split(', ')
            new_parts = []
            changed = False
            for part in parts:
                # Detectar formato "texto ×N"
                pm = re.match(r'^(.+?) ×(\d+)$', part)
                if pm:
                    item_name = pm.group(1)
                    qty = pm.group(2)
                    # Buscar en DB por nombre o por retailer_id
                    prod = CatalogProduct.query.get(item_name)
                    if prod and prod.name:
                        new_parts.append(f"{prod.name} ×{qty}")
                        if prod.name != item_name:
                            changed = True
                    elif not any(c in item_name for c in [' ', 'á', 'é', 'í', 'ó', 'ú', 'ñ']) and len(item_name) > 8:
                        # Parece un retailer_id crudo → reemplazar
                        new_parts.append(f"Artículo ×{qty}")
                        changed = True
                    else:
                        new_parts.append(part)
                else:
                    new_parts.append(part)
            if changed:
                msg.content = "[Pedido: " + ", ".join(new_parts) + "]"
                fixed += 1

        db.session.commit()
        return jsonify({'success': True, 'total_order_msgs': len(order_msgs), 'fixed': fixed})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error fixing order message content: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/minio/diagnose", methods=["GET"])
def api_minio_diagnose():
    """Diagnostica la conexión a MinIO."""
    from whatsapp_service import get_s3_client
    import requests

    results = {
        'config': {
            'endpoint': Config.MINIO_ENDPOINT,
            'bucket': Config.MINIO_BUCKET,
            'use_ssl': Config.MINIO_USE_SSL,
            'access_key_set': bool(Config.MINIO_ACCESS_KEY),
            'secret_key_set': bool(Config.MINIO_SECRET_KEY),
        },
        'tests': {}
    }

    protocol = "https" if Config.MINIO_USE_SSL else "http"
    endpoint_url = f"{protocol}://{Config.MINIO_ENDPOINT}"

    # Test 1: HTTP connectivity
    try:
        resp = requests.get(f"{endpoint_url}/minio/health/live", timeout=5, verify=False)
        results['tests']['health_check'] = {
            'status': 'ok' if resp.status_code == 200 else 'error',
            'http_code': resp.status_code,
            'response': resp.text[:200] if resp.text else None
        }
    except Exception as e:
        results['tests']['health_check'] = {'status': 'error', 'error': str(e)}

    # Test 2: List buckets
    try:
        s3 = get_s3_client()
        buckets = s3.list_buckets()
        results['tests']['list_buckets'] = {
            'status': 'ok',
            'buckets': [b['Name'] for b in buckets.get('Buckets', [])]
        }
    except Exception as e:
        results['tests']['list_buckets'] = {'status': 'error', 'error': str(e)}

    # Test 3: Head bucket
    try:
        s3 = get_s3_client()
        s3.head_bucket(Bucket=Config.MINIO_BUCKET)
        results['tests']['head_bucket'] = {'status': 'ok', 'bucket_exists': True}
    except Exception as e:
        results['tests']['head_bucket'] = {'status': 'error', 'error': str(e)}

    # Test 4: Try to create bucket
    try:
        s3 = get_s3_client()
        s3.create_bucket(Bucket=Config.MINIO_BUCKET)
        results['tests']['create_bucket'] = {'status': 'ok', 'created': True}
    except Exception as e:
        error_str = str(e)
        if 'BucketAlreadyOwnedByYou' in error_str or 'BucketAlreadyExists' in error_str:
            results['tests']['create_bucket'] = {'status': 'ok', 'already_exists': True}
        else:
            results['tests']['create_bucket'] = {'status': 'error', 'error': error_str}

    return jsonify(results)


@app.route("/webhook", methods=["POST"])
def webhook_handler():
    """
    Endpoint principal para recibir eventos de WhatsApp.
    """
    try:
        data = request.json
        if not data:
            return "No data received", 400
            
        # Procesar el evento (loguear, guardar, reenviar)
        process_event(data)
        
        return "EVENT_RECEIVED", 200
        
    except Exception as e:
        logger.error(f"Error procesando el webhook: {e}")
        return "Internal Server Error", 500

@app.route("/dashboard")
def dashboard():
    """Dashboard para visualizar conversaciones tipo WhatsApp - OPTIMIZADO."""
    selected_phone = request.args.get('phone')

    # OPTIMIZACIÓN: Removidas queries de stats generales (no se usan en UI del chat)
    # Si se necesitan, se pueden cargar via AJAX o en /analytics
    stats = {'total': 0, 'sent': 0, 'read': 0, 'failed': 0}

    CONTACTS_LIMIT = 25
    from sqlalchemy import text

    # Query optimizada: un solo JOIN entre mensajes y contactos (con sus tags).
    # Usa lateral join para obtener el último mensaje por teléfono sin doble sort.
    # El índice ix_messages_phone_ts (phone_number, timestamp) hace esto muy rápido.
    # Las conversaciones con tag "Asistencia Humana" se muestran primero.
    vis_sql, vis_params = build_visibility_sql(g.current_user)
    combined_query = text(f"""
        SELECT
            m.phone_number,
            m.content        AS last_message,
            m.timestamp      AS last_timestamp,
            c.id             AS contact_id,
            c.name           AS contact_name,
            CASE WHEN ha.contact_id IS NOT NULL THEN 1 ELSE 0 END AS has_human
        FROM (
            SELECT DISTINCT ON (phone_number) phone_number, content, timestamp
            FROM whatsapp_messages
            WHERE phone_number NOT IN ('unknown', 'outbound', '')
            ORDER BY phone_number, timestamp DESC
        ) m
        LEFT JOIN whatsapp_contacts c ON c.phone_number = m.phone_number
        LEFT JOIN LATERAL (
            SELECT ct.contact_id
            FROM whatsapp_contact_tags ct
            JOIN whatsapp_tags t ON t.id = ct.tag_id
            WHERE ct.contact_id = c.id AND t.name = 'Asistencia Humana'
            LIMIT 1
        ) ha ON true
        WHERE 1=1 {vis_sql}
        ORDER BY has_human DESC, m.timestamp DESC
        LIMIT :lim
    """)

    rows = db.session.execute(combined_query, {'lim': CONTACTS_LIMIT + 1, **vis_params}).fetchall()

    has_more_contacts = len(rows) > CONTACTS_LIMIT
    rows = rows[:CONTACTS_LIMIT]

    # Cargar tags solo para los contactos encontrados (una sola query con IN)
    contact_ids = [r.contact_id for r in rows if r.contact_id]
    tags_map = {}  # contact_id -> [Tag, ...]
    if contact_ids:
        from models import contact_tags as ct_table
        tag_rows = db.session.execute(
            text("""
                SELECT ct.contact_id, t.name
                FROM whatsapp_contact_tags ct
                JOIN whatsapp_tags t ON t.id = ct.tag_id
                WHERE ct.contact_id = ANY(:ids)
            """),
            {'ids': contact_ids}
        ).fetchall()
        for row in tag_rows:
            tags_map.setdefault(row.contact_id, []).append(row.name)

    contacts = []
    for r in rows:
        tag_names = tags_map.get(r.contact_id, [])
        last_msg = r.last_message
        contacts.append({
            'phone_number': r.phone_number,
            'last_timestamp': r.last_timestamp,
            'last_message': (last_msg[:50] + '...') if last_msg and len(last_msg) > 50 else last_msg,
            'name': r.contact_name,
            'tags': tag_names,
            'has_human_assistance': 'Asistencia Humana' in tag_names
        })

    # Si hay contacto seleccionado, obtener sus mensajes
    messages = []
    contact_stats = {}
    selected_contact = None
    contact_details = None
    
    # Límite de mensajes para mostrar en el chat
    MESSAGE_LIMIT = 60
    from sqlalchemy.orm import joinedload as _jl

    if selected_phone and not user_can_access_phone(g.current_user, selected_phone):
        return redirect(url_for('dashboard'))

    if selected_phone:
        selected_contact = selected_phone

        # Query raw: trae mensajes + latest_status en una sola pasada sin JOIN pesado.
        # Usa DISTINCT ON en statuses para traer solo el último estado por mensaje.
        msg_query = text("""
            SELECT
                m.id, m.wa_message_id, m.phone_number, m.direction,
                m.message_type, m.content, m.media_url, m.caption,
                m.timestamp, m.media_id,
                s.status AS latest_status
            FROM whatsapp_messages m
            LEFT JOIN LATERAL (
                SELECT status
                FROM whatsapp_message_statuses
                WHERE wa_message_id = m.wa_message_id
                ORDER BY timestamp DESC
                LIMIT 1
            ) s ON true
            WHERE m.phone_number = :phone
            ORDER BY m.timestamp DESC
            LIMIT :lim
        """)
        raw_msgs = db.session.execute(msg_query, {'phone': selected_phone, 'lim': MESSAGE_LIMIT}).fetchall()
        messages = list(reversed(raw_msgs))  # orden cronológico

        # joinedload evita query separado para las tags del contacto
        contact_details = Contact.query.options(_jl(Contact.tags))\
            .filter_by(phone_number=selected_phone).first()

        # Stats simplificadas
        outbound_msgs = [m for m in messages if m.direction == 'outbound']
        contact_stats = {
            'message_count': len(messages),
            'sent': sum(1 for m in outbound_msgs if m.latest_status in ['sent', 'delivered', 'read']),
            'delivered': sum(1 for m in outbound_msgs if m.latest_status in ['delivered', 'read']),
            'read': sum(1 for m in outbound_msgs if m.latest_status == 'read')
        }
    elif contacts:
        # Sin ?phone= en la URL: NO cargar mensajes en SSR.
        # El JS se encarga de cargar el chat del primer contacto via AJAX al iniciar.
        selected_contact = None
        messages = []
        contact_stats = {}
    
    # OPTIMIZACIÓN: Removidas queries de gráficos (se cargan en /analytics)
    chart_data = {
        'messages_by_day': [],
        'messages_by_hour': [],
        'direction_stats': {'inbound': 0, 'outbound': 0}
    }
    
    # Verificar ventana de 24 horas para envío de mensajes
    can_send_free_text = False
    last_inbound_msg = None
    templates = []
    whatsapp_configured = whatsapp_api.is_configured()
    
    if selected_contact and whatsapp_configured:
        # Buscar último mensaje entrante del contacto (optimizado: buscar en mensajes ya cargados)
        twenty_four_hours_ago = datetime.utcnow() - timedelta(hours=24)
        inbound_recent = [m for m in messages if m.direction == 'inbound' and m.timestamp >= twenty_four_hours_ago]
        
        if inbound_recent:
            can_send_free_text = True
            last_inbound_msg = max(m.timestamp for m in inbound_recent)
        
        # Templates se cargan async via AJAX para no bloquear el render
        # (ver /api/whatsapp/templates)
    
    # Verificar si el bot está pausado para el contacto seleccionado
    bot_paused = False
    if contact_details:
        bot_paused = any(t.name == 'Asistencia Humana' for t in contact_details.tags)

    # Etiquetas activas para el filtro del sidebar
    available_tags = Tag.query.filter_by(is_active=True).order_by(Tag.name).all()

    # Detectar si el usuario no tiene etiquetas asignadas (y no es admin)
    user = g.current_user
    _vis_ids = get_visible_tag_ids(user)
    has_no_visibility = _vis_ids is not None and len(_vis_ids) == 0 and not user.can_see_untagged

    # Mapa wa_message_id → order_id para el SSR de burbujas de pedido
    order_wamids = [m.wa_message_id for m in messages if m.message_type == 'order' and m.wa_message_id]
    order_id_map = {}
    if order_wamids:
        from models import Order as _Order
        _rows = _Order.query.filter(_Order.wa_message_id.in_(order_wamids)).with_entities(_Order.wa_message_id, _Order.id).all()
        order_id_map = {row[0]: row[1] for row in _rows}

    return render_template('dashboard.html',
                         stats=stats,
                         contacts=contacts,
                         messages=messages,
                         selected_contact=selected_contact,
                         contact_details=contact_details,
                         contact_stats=contact_stats,
                         chart_data=chart_data,
                         can_send_free_text=can_send_free_text,
                         last_inbound_msg=last_inbound_msg,
                         templates=templates,
                         whatsapp_configured=whatsapp_configured,
                         has_more_contacts=has_more_contacts,
                         bot_paused=bot_paused,
                         available_tags=available_tags,
                         has_no_visibility=has_no_visibility,
                         order_id_map=order_id_map)

@app.route("/analytics")
def analytics():
    """Página de analytics con estadísticas detalladas - OPTIMIZADO."""
    # Zona horaria de Argentina
    ARGENTINA_TZ = 'America/Argentina/Buenos_Aires'

    # Período de análisis configurable (default: 30 días)
    period = request.args.get('period', 30, type=int)

    # Calcular fecha de inicio según el período (0 = todo el historial)
    if period > 0:
        since_date = datetime.utcnow() - timedelta(days=period)
    else:
        since_date = None  # Sin filtro = todo el historial

    # OPTIMIZACIÓN: Una sola query para stats de mensajes con GROUP BY
    msg_query = db.session.query(
        Message.direction,
        func.count(Message.id).label('count')
    )
    if since_date:
        msg_query = msg_query.filter(Message.timestamp >= since_date)
    message_stats = msg_query.group_by(Message.direction).all()

    # Convertir a dict
    direction_counts = {row.direction: row.count for row in message_stats}
    outbound = direction_counts.get('outbound', 0)
    inbound = direction_counts.get('inbound', 0)
    total_messages = outbound + inbound

    # OPTIMIZACIÓN: Una sola query para todos los estados con GROUP BY
    status_query = db.session.query(
        MessageStatus.status,
        func.count(MessageStatus.id).label('count')
    )
    if since_date:
        status_query = status_query.filter(MessageStatus.timestamp >= since_date)
    status_stats = status_query.group_by(MessageStatus.status).all()

    # Convertir a dict
    status_counts = {row.status: row.count for row in status_stats}
    read = status_counts.get('read', 0)
    delivered = status_counts.get('delivered', 0)
    sent = status_counts.get('sent', 0)
    failed = status_counts.get('failed', 0)

    stats = {
        'total_messages': total_messages,
        'outbound': outbound,
        'inbound': inbound,
        'read': read,
        'delivered': delivered,
        'sent': sent,
        'failed': failed
    }

    # Fecha para gráficos (usa el mismo período seleccionado)
    chart_since = since_date if since_date else datetime.utcnow() - timedelta(days=365 * 10)  # 10 años si es "todo"
    
    # Mensajes por día (hora Argentina) - usa período seleccionado
    messages_by_day = db.session.execute(db.text(f"""
        SELECT
            DATE(timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}') as date,
            direction,
            COUNT(*) as count
        FROM whatsapp_messages
        WHERE timestamp >= :since
        GROUP BY DATE(timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}'), direction
    """), {'since': chart_since}).fetchall()

    # Formatear datos por día
    day_data = {}
    for row in messages_by_day:
        date_str = str(row.date) if row.date else ''
        if date_str not in day_data:
            day_data[date_str] = {'date': date_str, 'inbound': 0, 'outbound': 0}
        if row.direction == 'inbound':
            day_data[date_str]['inbound'] = row.count
        else:
            day_data[date_str]['outbound'] = row.count

    # Mensajes enviados por hora (hora Argentina) - usa período seleccionado
    sent_by_hour = db.session.execute(db.text(f"""
        SELECT
            EXTRACT(HOUR FROM timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}')::int as hour,
            COUNT(*) as count
        FROM whatsapp_messages
        WHERE direction = 'outbound' AND timestamp >= :since
        GROUP BY EXTRACT(HOUR FROM timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}')
        ORDER BY hour
    """), {'since': chart_since}).fetchall()

    # Mensajes leídos por hora - cuenta mensajes ENVIADOS en esa hora que fueron leídos (en cualquier momento)
    # Corrige el problema de tasas > 100% comparando mensajes con sus estados correctamente
    read_by_hour = db.session.execute(db.text(f"""
        SELECT
            EXTRACT(HOUR FROM m.timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}')::int as hour,
            COUNT(DISTINCT m.id) as count
        FROM whatsapp_messages m
        INNER JOIN whatsapp_message_statuses s ON m.wa_message_id = s.wa_message_id
        WHERE m.direction = 'outbound'
          AND m.timestamp >= :since
          AND s.status = 'read'
        GROUP BY EXTRACT(HOUR FROM m.timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}')
        ORDER BY hour
    """), {'since': chart_since}).fetchall()

    # Mensajes por día de la semana (hora Argentina) - usa período seleccionado
    by_day_of_week = db.session.execute(db.text(f"""
        SELECT
            EXTRACT(DOW FROM timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}')::int as dow,
            COUNT(*) as count
        FROM whatsapp_messages
        WHERE timestamp >= :since
        GROUP BY EXTRACT(DOW FROM timestamp AT TIME ZONE 'UTC' AT TIME ZONE '{ARGENTINA_TZ}')
    """), {'since': chart_since}).fetchall()
    
    dow_counts = [0] * 7
    for row in by_day_of_week:
        if row.dow is not None:
            idx = int(row.dow)
            # Ajustar para que lunes sea 0
            idx = (idx - 1) % 7
            dow_counts[idx] = row.count
    
    # Top contactos
    top_contacts = db.session.query(
        Message.phone_number,
        func.count(Message.id).label('count')
    ).filter(
        Message.phone_number.notin_(['unknown', 'outbound', ''])
    ).group_by(Message.phone_number).order_by(func.count(Message.id).desc()).limit(5).all()
    
    chart_data = {
        'messages_by_day': sorted(day_data.values(), key=lambda x: x['date']),
        'status_dist': {'read': read, 'delivered': delivered, 'sent': sent, 'failed': failed},
        'sent_by_hour': [{'hour': int(h.hour) if h.hour else 0, 'count': h.count} for h in sent_by_hour],
        'read_by_hour': [{'hour': int(h.hour) if h.hour else 0, 'count': h.count} for h in read_by_hour],
        'by_day_of_week': dow_counts,
        'direction': {'inbound': inbound, 'outbound': outbound},
        'top_contacts': [{'phone': c.phone_number, 'count': c.count} for c in top_contacts]
    }
    
    # ========== ESTADÍSTICAS DE TEMPLATES ==========
    # Obtener mensajes salientes del período seleccionado
    outbound_messages = Message.query.filter(
        Message.direction == 'outbound',
        Message.timestamp >= chart_since
    ).all()
    
    # Obtener templates de la API para mapeo
    templates_info = whatsapp_api.get_templates().get("templates", [])
    
    # Agrupar por "nombre" de template
    stats_by_template = {} # key: template_name, value: {sent: 0, read: 0}
    
    import re
    # Pre-calcular patrones de templates para mayor velocidad
    template_patterns = []
    for t in templates_info:
        for comp in t.get("components", []):
            if comp.get("type") == "BODY":
                body = comp.get("text", "").strip()
                if not body: continue
                # Limpiar el escape de re.escape para que sea más flexible
                # En lugar de re.escape completo, escapamos solo caracteres especiales pero no espacios
                pattern = re.escape(body)
                # Reemplazar variables {{n}} por un comodín
                pattern = re.sub(r'\\\{\\\{\d+\\\}\\\}', '.*?', pattern)
                # Permitir cualquier cantidad de espacios/newslines donde haya uno
                pattern = re.sub(r'\\ ', r'\\s+', pattern)
                pattern = re.sub(r'\\n', r'\\s*', pattern)
                
                template_patterns.append({
                    'name': t.get("name"), 
                    'regex': f".*{pattern}.*" # Permitir que esté contenido (por si hay Header/Footer)
                })
    
    for msg in outbound_messages:
        t_name = None
        content = (msg.content or "").strip()
        if not content: continue
        
        # 1. Prioridad: Mensajes ya marcados con [Template: nombre]
        match_name = re.match(r'^\[Template: ([^\]]+)\]', content)
        if match_name:
            t_name = match_name.group(1)
        
        # 2. Si no, intentar por coincidencia de patrones (independiente de message_type)
        if not t_name:
            for tp in template_patterns:
                try:
                    if re.match(tp['regex'], content, re.DOTALL | re.IGNORECASE):
                        t_name = tp['name']
                        break
                except: continue
        
        # 3. Si aún no hay nombre pero es tipo template, usar contenido truncado
        if not t_name and msg.message_type == 'template':
            t_name = content[:50] + "..." if len(content) > 50 else content
        
        # Si detectamos que es un template, sumar a stats
        if t_name:
            if t_name not in stats_by_template:
                stats_by_template[t_name] = {'sent': 0, 'read': 0}
            
            stats_by_template[t_name]['sent'] += 1
            is_read = any(s.status == 'read' for s in msg.statuses)
            if is_read:
                stats_by_template[t_name]['read'] += 1

    # Convertir a lista para el template
    template_performance = []
    for name, s in stats_by_template.items():
        read_rate = round((s['read'] / s['sent'] * 100) if s['sent'] > 0 else 0, 1)
        template_performance.append({
            'name': name,
            'sent': s['sent'],
            'read': s['read'],
            'read_rate': read_rate
        })
    
    # Ordenar por más enviados
    template_performance = sorted(template_performance, key=lambda x: x['sent'], reverse=True)[:10]
    
    # ========== MEJORES HORARIOS PARA LECTURA ==========
    # Convertir datos de lectura por hora a un formato más útil
    read_by_hour_dict = {int(h.hour) if h.hour else 0: h.count for h in read_by_hour}
    sent_by_hour_dict = {int(h.hour) if h.hour else 0: h.count for h in sent_by_hour}
    
    # Calcular tasa de lectura por hora
    hourly_read_rate = []
    for hour in range(24):
        sent_at_hour = sent_by_hour_dict.get(hour, 0)
        read_at_hour = read_by_hour_dict.get(hour, 0)
        rate = round((read_at_hour / sent_at_hour * 100) if sent_at_hour > 0 else 0, 1)
        hourly_read_rate.append({
            'hour': hour,
            'sent': sent_at_hour,
            'read': read_at_hour,
            'rate': rate
        })
    
    # Encontrar las mejores horas para enviar (mayor tasa de lectura)
    # Solo considerar horas con al menos 5 mensajes enviados
    best_hours = sorted(
        [h for h in hourly_read_rate if h['sent'] >= 5],
        key=lambda x: x['rate'],
        reverse=True
    )[:3]
    
    # Hora con más lecturas (no tasa, cantidad absoluta)
    peak_read_hour = max(hourly_read_rate, key=lambda x: x['read']) if hourly_read_rate else None
    
    # Insights
    peak_hour = max(sent_by_hour, key=lambda x: x.count) if sent_by_hour else None
    busiest_dow = dow_counts.index(max(dow_counts)) if dow_counts else 0
    days_names = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
    
    insights = {
        'peak_hour': int(peak_hour.hour) if peak_hour and peak_hour.hour else 12,
        'peak_hour_count': peak_hour.count if peak_hour else 0,
        'read_rate': round((read / outbound * 100) if outbound > 0 else 0, 1),
        'busiest_day': days_names[busiest_dow],
        'avg_daily': round(total_messages / 30, 1) if total_messages > 0 else 0,
        'best_hours': best_hours,
        'peak_read_hour': peak_read_hour['hour'] if peak_read_hour else 12,
        'template_count': len(template_performance)
    }
    
    # ==========================================
    # CONVERSATION CATEGORIZATION ANALYTICS
    # ==========================================
    topic_distribution = []
    rating_distribution = []
    
    # Topic distribution
    topic_stats = db.session.query(
        ConversationTopic.name,
        ConversationTopic.color,
        db.func.count(ConversationSession.id).label('count')
    ).outerjoin(ConversationSession, ConversationTopic.id == ConversationSession.topic_id)\
     .group_by(ConversationTopic.id)\
     .order_by(db.func.count(ConversationSession.id).desc())\
     .all()
    
    for stat in topic_stats:
        topic_distribution.append({
            'name': stat.name,
            'color': stat.color,
            'count': stat.count or 0
        })
    
    # Add uncategorized sessions count
    uncategorized_count = ConversationSession.query.filter(ConversationSession.topic_id.is_(None)).count()
    if uncategorized_count > 0:
        topic_distribution.append({
            'name': 'Sin categorizar',
            'color': 'gray',
            'count': uncategorized_count
        })
    
    # Rating distribution
    rating_stats = db.session.query(
        ConversationSession.rating,
        db.func.count(ConversationSession.id).label('count')
    ).filter(ConversationSession.rating.isnot(None))\
     .group_by(ConversationSession.rating)\
     .all()
    
    rating_map = {
        'buena': ('Buena', '#22c55e'),
        'neutral': ('Neutral', '#f59e0b'),
        'mala': ('Mala', '#ef4444')
    }
    
    for stat in rating_stats:
        if stat.rating in rating_map:
            name, color = rating_map[stat.rating]
            rating_distribution.append({
                'name': name,
                'color': color,
                'count': stat.count
            })
    
    session_analytics = {
        'topic_distribution': topic_distribution,
        'rating_distribution': rating_distribution,
        'total_sessions': ConversationSession.query.count()
    }
    
    return render_template('analytics.html',
                         stats=stats,
                         chart_data=chart_data,
                         insights=insights,
                         template_performance=template_performance,
                         hourly_read_rate=hourly_read_rate,
                         session_analytics=session_analytics,
                         period=period)

@app.route("/api/stats")
def api_stats():
    """API endpoint para obtener estadísticas en JSON."""
    twenty_four_hours_ago = datetime.utcnow() - timedelta(hours=24)
    
    total = Message.query.filter(Message.timestamp >= twenty_four_hours_ago).count()
    
    sent = db.session.query(func.count(MessageStatus.id)).filter(
        MessageStatus.status == 'sent',
        MessageStatus.timestamp >= twenty_four_hours_ago
    ).scalar() or 0
    
    delivered = db.session.query(func.count(MessageStatus.id)).filter(
        MessageStatus.status == 'delivered',
        MessageStatus.timestamp >= twenty_four_hours_ago
    ).scalar() or 0
    
    read = db.session.query(func.count(MessageStatus.id)).filter(
        MessageStatus.status == 'read',
        MessageStatus.timestamp >= twenty_four_hours_ago
    ).scalar() or 0
    
    failed = db.session.query(func.count(MessageStatus.id)).filter(
        MessageStatus.status == 'failed',
        MessageStatus.timestamp >= twenty_four_hours_ago
    ).scalar() or 0
    
    total_attempts = sent + delivered + read + failed
    success_rate = round(((delivered + read) / total_attempts * 100) if total_attempts > 0 else 100, 1)
    
    return jsonify({
        'total': total,
        'sent': sent,
        'delivered': delivered,
        'read': read,
        'failed': failed,
        'success_rate': success_rate
    })

def get_visible_tag_ids(user):
    """
    Retorna la lista de tag_ids visibles para el usuario, o None si no hay restricción (admin).
    Lista vacía significa que no ve ninguna conversación.
    """
    if user.is_admin:
        return None
    return [v.tag_id for v in user.tag_visibility]


def user_can_access_phone(user, phone):
    """
    Verifica si el usuario tiene acceso a un contacto por su número de teléfono.
    Retorna True si tiene acceso, False si no.
    """
    vis_tag_ids = get_visible_tag_ids(user)
    if vis_tag_ids is None:
        return True  # admin → acceso total
    contact = Contact.query.filter_by(phone_number=phone).first()
    if contact:
        contact_tag_ids = {t.id for t in contact.tags}
        if contact_tag_ids & set(vis_tag_ids):
            return True
        if user.can_see_untagged and not contact_tag_ids:
            return True
        return False
    else:
        return bool(user.can_see_untagged)


def build_visibility_sql(user, phone_alias='m.phone_number'):
    """
    Retorna (sql_fragment, params) para filtrar conversaciones por visibilidad de etiquetas.
    sql_fragment empieza con AND y se puede concatenar directamente al WHERE.
    """
    if user.is_admin:
        return '', {}

    tag_ids = get_visible_tag_ids(user)

    if not tag_ids:
        return 'AND 1=0', {}  # Sin etiquetas → no ve nada

    conditions = [f"""
        EXISTS (
            SELECT 1 FROM whatsapp_contacts _vc
            JOIN whatsapp_contact_tags _vct ON _vct.contact_id = _vc.id
            WHERE _vc.phone_number = {phone_alias} AND _vct.tag_id = ANY(:_vis_ids)
        )
    """]

    if user.can_see_untagged:
        conditions.append(f"""(
            NOT EXISTS (SELECT 1 FROM whatsapp_contacts _vc2 WHERE _vc2.phone_number = {phone_alias})
            OR EXISTS (
                SELECT 1 FROM whatsapp_contacts _vc3
                WHERE _vc3.phone_number = {phone_alias}
                AND NOT EXISTS (SELECT 1 FROM whatsapp_contact_tags _vct3 WHERE _vct3.contact_id = _vc3.id)
            )
        )""")

    return f'AND ({" OR ".join(conditions)})', {'_vis_ids': tag_ids}


def format_utc_iso(dt):
    """Convierte datetime a string ISO 8601 con sufijo Z para UTC."""
    if not dt:
        return None
    if dt.tzinfo is None:
        # Asumir UTC si es naive
        return dt.isoformat() + 'Z'
    # Si tiene zona horaria, convertir a UTC explícitamente
    return dt.astimezone(timezone.utc).isoformat().replace('+00:00', 'Z')

def normalize_phone(phone):
    """
    Normaliza un número de teléfono eliminando el '+' inicial y espacios.
    Ejemplo: '+5493874882011' -> '5493874882011'
    Busca primero con el número normalizado; si no encuentra, intenta con el original.
    """
    if not phone:
        return phone
    return phone.strip().lstrip('+')


def find_contact_by_phone(phone):
    """
    Busca un contacto tolerando variantes con/sin '+'.
    Prueba: número normalizado (sin +), luego con '+'.
    """
    normalized = normalize_phone(phone)
    contact = Contact.query.filter_by(phone_number=normalized).first()
    if not contact and not phone.startswith('+'):
        # Intentar también con el '+' por si se guardó con él
        contact = Contact.query.filter_by(phone_number='+' + normalized).first()
    return contact


def register_contact_if_new(phone_number, name=None):
    """Registra un contacto si no existe."""
    try:
        if not phone_number or phone_number in ['unknown', 'outbound', '']:
            return
            
        contact = Contact.query.filter_by(phone_number=phone_number).first()
        if not contact:
            new_contact = Contact(phone_number=phone_number, name=name)
            db.session.add(new_contact)
            db.session.commit()
            logger.info(f"🆕 Nuevo contacto registrado: {phone_number}")
        elif name and not contact.name:
            contact.name = name
            db.session.commit()
    except Exception as e:
        logger.error(f"Error registrando contacto auto: {e}")
        db.session.rollback()

# ==========================================
# API DASHBOARD CONTACTS (con scroll infinito)
# ==========================================

@app.route("/api/dashboard/contacts", methods=["GET"])
def api_dashboard_contacts():
    """API para obtener contactos del dashboard con paginación, búsqueda y filtro de etiqueta."""
    search = request.args.get('search', '').strip()
    tag_filter = request.args.get('tag', '').strip()
    offset = request.args.get('offset', 0, type=int)
    limit = request.args.get('limit', 30, type=int)

    # Limitar el máximo de resultados por request
    limit = min(limit, 50)

    from sqlalchemy import text

    # Filtro de etiqueta explícito (barra lateral)
    tag_join = ""
    tag_params = {}
    if tag_filter:
        tag_join = """
            JOIN whatsapp_contacts c_tag ON c_tag.phone_number = m.phone_number
            JOIN whatsapp_contact_tags ct_tag ON ct_tag.contact_id = c_tag.id
            JOIN whatsapp_tags t_tag ON t_tag.id = ct_tag.tag_id AND t_tag.name = :tag_filter
        """
        tag_params['tag_filter'] = tag_filter

    # Filtro de visibilidad por usuario
    vis_sql, vis_params = build_visibility_sql(g.current_user)

    if search:
        search_query = text(f"""
            SELECT sub.phone_number, sub.last_message, sub.last_timestamp
            FROM (
                SELECT DISTINCT ON (m.phone_number) m.phone_number, m.content AS last_message, m.timestamp AS last_timestamp
                FROM whatsapp_messages m
                {tag_join}
                WHERE m.phone_number NOT IN ('unknown', 'outbound', '')
                  AND (
                    m.phone_number ILIKE :pattern
                    OR m.phone_number IN (
                        SELECT c.phone_number FROM whatsapp_contacts c
                        WHERE c.name ILIKE :pattern OR c.phone_number ILIKE :pattern
                    )
                    OR m.phone_number IN (
                        SELECT DISTINCT m2.phone_number FROM whatsapp_messages m2
                        WHERE m2.content ILIKE :pattern
                          AND m2.phone_number NOT IN ('unknown', 'outbound', '')
                    )
                  )
                  {vis_sql}
                ORDER BY m.phone_number, m.timestamp DESC
            ) sub
            LEFT JOIN whatsapp_contacts c2 ON c2.phone_number = sub.phone_number
            LEFT JOIN LATERAL (
                SELECT ct.contact_id FROM whatsapp_contact_tags ct
                JOIN whatsapp_tags t ON t.id = ct.tag_id
                WHERE ct.contact_id = c2.id AND t.name = 'Asistencia Humana' LIMIT 1
            ) ha ON true
            ORDER BY CASE WHEN ha.contact_id IS NOT NULL THEN 1 ELSE 0 END DESC, sub.last_timestamp DESC
            OFFSET :off LIMIT :lim
        """)
        results = db.session.execute(search_query, {
            'pattern': f'%{search}%', 'off': offset, 'lim': limit + 1, **tag_params, **vis_params
        }).fetchall()
    else:
        distinct_query = text(f"""
            SELECT sub.phone_number, sub.last_message, sub.last_timestamp
            FROM (
                SELECT DISTINCT ON (m.phone_number) m.phone_number, m.content AS last_message, m.timestamp AS last_timestamp
                FROM whatsapp_messages m
                {tag_join}
                WHERE m.phone_number NOT IN ('unknown', 'outbound', '')
                {vis_sql}
                ORDER BY m.phone_number, m.timestamp DESC
            ) sub
            LEFT JOIN whatsapp_contacts c ON c.phone_number = sub.phone_number
            LEFT JOIN LATERAL (
                SELECT ct.contact_id FROM whatsapp_contact_tags ct
                JOIN whatsapp_tags t ON t.id = ct.tag_id
                WHERE ct.contact_id = c.id AND t.name = 'Asistencia Humana' LIMIT 1
            ) ha ON true
            ORDER BY CASE WHEN ha.contact_id IS NOT NULL THEN 1 ELSE 0 END DESC, sub.last_timestamp DESC
            OFFSET :off LIMIT :lim
        """)
        results = db.session.execute(distinct_query, {
            'off': offset, 'lim': limit + 1, **tag_params, **vis_params
        }).fetchall()

    has_more = len(results) > limit
    results = results[:limit]

    # Obtener nombres de contactos
    phones = [r.phone_number for r in results]
    contacts_map = {}
    if phones:
        found_contacts = Contact.query.filter(Contact.phone_number.in_(phones)).all()
        contacts_map = {c.phone_number: c for c in found_contacts}

    # Formatear respuesta
    contacts = []
    for r in results:
        contact = contacts_map.get(r.phone_number)
        last_msg = r.last_message
        # Manejar timestamp de forma robusta (raw SQL puede devolver string o datetime)
        ts = r.last_timestamp
        if ts:
            ts_str = ts.isoformat() if hasattr(ts, 'isoformat') else str(ts)
        else:
            ts_str = None
        has_human = any(t.name == 'Asistencia Humana' for t in contact.tags) if contact else False
        contacts.append({
            'phone_number': r.phone_number,
            'last_timestamp': ts_str,
            'last_message': (last_msg[:50] + '...') if last_msg and len(last_msg) > 50 else last_msg,
            'name': contact.name if contact else None,
            'tags': [t.name for t in contact.tags] if contact else [],
            'has_human_assistance': has_human
        })

    return jsonify({
        'contacts': contacts,
        'total': offset + len(contacts) + (1 if has_more else 0),
        'offset': offset,
        'limit': limit,
        'has_more': has_more
    })


# ==========================================
# API CRM CONTACTOS
# ==========================================

@app.route("/api/contacts", methods=["POST"])
def api_create_contact():
    """API para crear un nuevo contacto."""
    data = request.get_json()
    
    phone = data.get('phone_number', '').strip()
    if not phone:
        return jsonify({'success': False, 'error': 'El teléfono es requerido'}), 400
    
    # Crear nuevo contacto
    contact = Contact(
        phone_number=phone,
        contact_id=data.get('contact_id') or None,
        name=data.get('name', '').strip() or None,
        first_name=data.get('first_name', '').strip() or None,
        last_name=data.get('last_name', '').strip() or None,
        notes=data.get('notes', '').strip() or None,
        custom_field_1=data.get('custom_field_1', '').strip() or None,
        custom_field_2=data.get('custom_field_2', '').strip() or None,
        custom_field_3=data.get('custom_field_3', '').strip() or None,
        custom_field_4=data.get('custom_field_4', '').strip() or None,
        custom_field_5=data.get('custom_field_5', '').strip() or None,
        custom_field_6=data.get('custom_field_6', '').strip() or None,
        custom_field_7=data.get('custom_field_7', '').strip() or None
    )
    
    db.session.add(contact)
    db.session.flush()  # Para obtener el ID
    
    # Procesar tags
    new_tags = data.get('tags', [])
    if new_tags:
        for tag_name in new_tags:
            tag = Tag.query.filter_by(name=tag_name).first()
            if not tag:
                tag = Tag(name=tag_name)
                db.session.add(tag)
                db.session.flush()
            contact.tags.append(tag)
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'id': contact.id,
        'message': 'Contacto creado correctamente'
    })

@app.route("/api/contacts/<identifier>", methods=["GET", "POST", "DELETE"])
def api_contact_detail(identifier):
    """API para obtener, actualizar o eliminar un contacto.

    El identificador puede ser:
    - Un ID numérico interno (ej: 123) - solo dígitos cortos
    - Un contact_id externo (ej: CLI-001) - si contiene letras/guiones
    - Un número de teléfono (ej: 5491123456789) - solo dígitos largos

    POST permite cambiar el phone_number y contact_id.
    DELETE elimina el contacto permanentemente.
    """
    # Determinar tipo de identificador
    contact = None
    is_internal_id = identifier.isdigit() and len(identifier) <= 10  # IDs internos son cortos
    is_phone = identifier.isdigit() and len(identifier) > 10  # Teléfonos son largos

    if is_internal_id:
        contact = Contact.query.get(int(identifier))
    elif is_phone:
        contact = Contact.query.filter_by(phone_number=identifier).first()
    else:
        # Buscar por contact_id externo
        contact = Contact.query.filter_by(contact_id=identifier).first()

    # DELETE method - Eliminar contacto
    if request.method == "DELETE":
        if not contact:
            return jsonify({'error': f'Contacto no encontrado: {identifier}'}), 404

        try:
            contact_info = f"ID={contact.id}, Tel={contact.phone_number}, Nombre={contact.name or 'Sin nombre'}"

            # Eliminar registros relacionados en campaign_logs
            CampaignLog.query.filter_by(contact_id=contact.id).delete()

            # Eliminar el contacto (las tags se desvinculan automáticamente)
            db.session.delete(contact)
            db.session.commit()

            logger.info(f"🗑️ Contacto eliminado: {contact_info}")
            return jsonify({'success': True, 'message': f'Contacto eliminado correctamente'})
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error eliminando contacto {identifier}: {str(e)}")
            return jsonify({'error': str(e)}), 500

    if request.method == "POST":
        data = request.json
        is_new = False

        if not contact:
            # Crear nuevo contacto (solo si se pasa un teléfono, no un ID)
            if is_internal_id:
                return jsonify({'error': f'Contacto con ID {identifier} no encontrado'}), 404
            if not is_phone:
                return jsonify({'error': f'Contacto con contact_id "{identifier}" no encontrado'}), 404
            contact = Contact(phone_number=identifier)
            db.session.add(contact)
            is_new = True

        # Permitir cambio de contact_id (ID externo editable)
        if 'contact_id' in data:
            new_contact_id = data['contact_id'].strip() if data['contact_id'] else None
            if new_contact_id and new_contact_id != contact.contact_id:
                # Verificar que no exista otro contacto con ese contact_id
                existing = Contact.query.filter_by(contact_id=new_contact_id).first()
                if existing and existing.id != contact.id:
                    return jsonify({
                        'error': f'El ID externo "{new_contact_id}" ya pertenece a otro contacto (ID interno: {existing.id}, Nombre: {existing.name or "Sin nombre"})'
                    }), 400
                contact.contact_id = new_contact_id
                logger.info(f"🆔 Contact ID actualizado para contacto ID {contact.id}: → {new_contact_id}")
            elif not new_contact_id:
                contact.contact_id = None

        # Permitir cambio de teléfono si viene en el payload
        if 'phone_number' in data:
            new_phone = data['phone_number'].strip() if data['phone_number'] else None
            if new_phone and new_phone != contact.phone_number:
                contact.phone_number = new_phone
                logger.info(f"📱 Teléfono actualizado para contacto ID {contact.id}: {identifier} → {new_phone}")

        # Mapeo de campos
        fields = ['name', 'first_name', 'last_name', 'notes',
                  'custom_field_1', 'custom_field_2', 'custom_field_3',
                  'custom_field_4', 'custom_field_5', 'custom_field_6', 'custom_field_7']

        for field in fields:
            if field in data:
                setattr(contact, field, data[field])

        newly_added_tag_ids = []
        if 'tags' in data:
            new_tag_names = set(data['tags'])
            current_tag_names = {t.name for t in contact.tags}
            editor_name = g.current_user.username if g.current_user else 'manual'
            # Tags a agregar
            for name in new_tag_names - current_tag_names:
                tag = Tag.query.filter_by(name=name).first()
                if not tag:
                    tag = Tag(name=name)
                    db.session.add(tag)
                    db.session.flush()
                contact.tags.append(tag)
                newly_added_tag_ids.append(tag.id)
                _record_tag_history(contact.id, tag, 'added', 'manual', editor_name)
            # Tags a eliminar
            to_remove = current_tag_names - new_tag_names
            removed_tags = [t for t in contact.tags if t.name in to_remove]
            for tag in removed_tags:
                _record_tag_history(contact.id, tag, 'removed', 'manual', editor_name)
            contact.tags = [t for t in contact.tags if t.name not in to_remove]

        try:
            db.session.commit()
            action = "creado" if is_new else "actualizado"
            logger.info(f"✅ Contacto {action}: ID={contact.id}, Tel={contact.phone_number}")
            # Enrollar en secuencias para cada tag nuevo asignado manualmente
            if newly_added_tag_ids:
                from auto_tagger import enroll_in_sequences
                for tid in newly_added_tag_ids:
                    enroll_in_sequences(db, contact, tid, FollowUpSequence, FollowUpEnrollment)
            return jsonify({'success': True, 'contact': contact.to_dict()})
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500

    # GET method
    if not contact:
        return jsonify({'found': False, 'details': {'phone_number': identifier if is_phone else None, 'id': int(identifier) if is_internal_id else None}})

    return jsonify({'found': True, 'details': contact.to_dict()})

@app.route("/api/contacts/import", methods=["POST"])

def api_import_contacts():
    """Importar contactos desde Excel/CSV con mapeo estricto y optimización por lotes.

    Prioridad de búsqueda:
    1. Si existe columna ID y tiene valor → buscar por ID (permite cambiar teléfono)
    2. Si no hay ID → buscar por Teléfono (comportamiento tradicional)
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        # Leer archivo
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)

        # Normalizar columnas (strip solamente, mantener case)
        df.columns = [str(c).strip() for c in df.columns]

        # Identificar columna de ID interno (opcional, para actualizaciones)
        id_cols = ['ID', 'Id', 'id']
        id_col = next((c for c in df.columns if c in id_cols), None)

        # Identificar columna de Contact ID externo (opcional)
        contact_id_cols = ['Contact ID', 'ContactID', 'contact_id', 'ID Externo', 'External ID']
        contact_id_col = next((c for c in df.columns if c in contact_id_cols), None)

        # Identificar columna de teléfono (requerida)
        phone_cols = ['Telefono', 'Teléfono', 'Phone', 'Celular']
        phone_col = next((c for c in df.columns if c in phone_cols), None)

        if not phone_col:
            return jsonify({'error': f'Columna de teléfono no encontrada. Se busca una de: {", ".join(phone_cols)}'}), 400

        col_map = {
            'Nombre completo': 'name',
            'Nombre': 'first_name',
            'Apellido': 'last_name',
            'Notas': 'notes',
            'Campo 1': 'custom_field_1',
            'Campo 2': 'custom_field_2',
            'Campo 3': 'custom_field_3',
            'Campo 4': 'custom_field_4',
            'Campo 5': 'custom_field_5',
            'Campo 6': 'custom_field_6',
            'Campo 7': 'custom_field_7'
        }

        # Pre-calcular tag de importación
        import_tag = None
        import_tag_name = request.form.get('assign_tag', '').strip()
        if import_tag_name:
            import_tag = Tag.query.filter_by(name=import_tag_name).first()
            if not import_tag:
                import_tag = Tag(name=import_tag_name)
                db.session.add(import_tag)
                db.session.flush()

        # =====================================================
        # OPTIMIZACIÓN POR LOTES (BATCH PROCESSING)
        # =====================================================
        
        # 1. Recolectar todos los identificadores del archivo para hacer un solo query
        all_contact_ids = set()
        all_phones = set()
        all_internal_ids = set()
        
        # Normalizar datos en el dataframe para facilitar procesamiento
        # Convertir a string y limpiar
        df['clean_phone'] = df[phone_col].apply(lambda x: str(x).replace('.0', '').strip() if pd.notna(x) else '')
        
        if contact_id_col:
            df['clean_contact_id'] = df[contact_id_col].apply(lambda x: str(x).strip() if pd.notna(x) else None)
            all_contact_ids = set(df['clean_contact_id'].dropna().unique())
            # Eliminar strings vacíos
            all_contact_ids = {cid for cid in all_contact_ids if cid}
            
        if id_col:
            # IDs internos suelen ser enteros
            def clean_id(x):
                try: 
                    return int(float(x)) 
                except: 
                    return None
            df['clean_internal_id'] = df[id_col].apply(clean_id)
            all_internal_ids = set(df['clean_internal_id'].dropna().unique())

        all_phones = set(df['clean_phone'].dropna().unique())
        all_phones = {p for p in all_phones if p} # Eliminar vacíos
        
        # 2. Pre-cargar contactos existentes de la base de datos
        existing_contacts_by_cid = {}
        existing_contacts_by_phone = {}
        existing_contacts_by_iid = {}
        
        # Buscar por Contact ID
        if all_contact_ids:
            results_cid = Contact.query.filter(Contact.contact_id.in_(all_contact_ids)).all()
            for c in results_cid:
                existing_contacts_by_cid[c.contact_id] = c
                existing_contacts_by_phone[c.phone_number] = c # También indexar por tel para evitar dups
                existing_contacts_by_iid[c.id] = c

        # Buscar por Phone (que no hayamos traído ya)
        phones_to_fetch = all_phones - set(existing_contacts_by_phone.keys())
        if phones_to_fetch:
            # Optimización: Consultar en chunks si son muchísimos (>1000)
            # SQLAlchemy maneja bien IN clauses grandes pero postgres tiene límites de parámetros (~65k)
            # Para 3000 filas es seguro hacerlo de una vez
            results_phone = Contact.query.filter(Contact.phone_number.in_(phones_to_fetch)).all()
            for c in results_phone:
                existing_contacts_by_phone[c.phone_number] = c
                if c.contact_id: existing_contacts_by_cid[c.contact_id] = c
                existing_contacts_by_iid[c.id] = c
                
        # Buscar por Internal ID (fallback)
        ids_to_fetch = all_internal_ids - set(existing_contacts_by_iid.keys())
        if ids_to_fetch:
            results_iid = Contact.query.filter(Contact.id.in_(ids_to_fetch)).all()
            for c in results_iid:
                existing_contacts_by_iid[c.id] = c
                existing_contacts_by_phone[c.phone_number] = c
                if c.contact_id: existing_contacts_by_cid[c.contact_id] = c

        count = 0
        updated = 0
        phone_updated = 0
        errors = []
        
        # 3. Procesar filas en memoria
        for idx, row in df.iterrows():
            phone = row['clean_phone']
            if not phone:
                continue

            contact = None
            found_by = None
            
            # A. Buscar en memoria
            # 1. Contact ID
            ext_id = row.get('clean_contact_id') if contact_id_col else None
            if ext_id and ext_id in existing_contacts_by_cid:
                contact = existing_contacts_by_cid[ext_id]
                found_by = 'contact_id'
            
            # 2. Internal ID
            int_id = row.get('clean_internal_id') if id_col else None
            if not contact and int_id and int_id in existing_contacts_by_iid:
                contact = existing_contacts_by_iid[int_id]
                found_by = 'id'
                
            # 3. Phone
            if not contact and phone in existing_contacts_by_phone:
                contact = existing_contacts_by_phone[phone]
                found_by = 'phone'
            
            is_new = False
            
            if not contact:
                # CREACIÓN
                
                # Validar Client ID obligatorio
                if not ext_id: # ext_id ya está limpio y verificado
                    errors.append(f"Fila {idx+2}: Ignorado - Se requiere Client ID (contact_id) para crear nuevos contactos")
                    continue
                
                # Crear nuevo
                contact = Contact(phone_number=phone)
                contact.contact_id = ext_id
                
                db.session.add(contact)
                
                # Actualizar índices en memoria para futuras filas en este mismo loop (por si hay reps)
                existing_contacts_by_cid[ext_id] = contact
                existing_contacts_by_phone[phone] = contact
                
                is_new = True
                count += 1
            else:
                # ACTUALIZACIÓN
                if found_by in ('contact_id', 'id') and contact.phone_number != phone:
                    old_phone = contact.phone_number
                    # Actualizar índice en memoria: quitar el viejo teléfono
                    if old_phone in existing_contacts_by_phone:
                        # Solo si apunta a este contacto (cuidado con colisiones)
                        if existing_contacts_by_phone[old_phone] == contact:
                            del existing_contacts_by_phone[old_phone]
                            
                    contact.phone_number = phone
                    # Actualizar índice con nuevo teléfono
                    existing_contacts_by_phone[phone] = contact
                    
                    phone_updated += 1
                    # logger.info(...) # Evitar exceso de logs en loop
                updated += 1

            # Actualizar campos mapeados
            for excel_col, model_attr in col_map.items():
                if excel_col in df.columns:
                    val = row[excel_col]
                    if pd.notna(val):
                        setattr(contact, model_attr, str(val))

            # Actualizar Contact ID si fue encontrado por teléfono y el archivo tiene uno nuevo
            if found_by in ('phone', 'id') and ext_id:
                if ext_id != contact.contact_id:
                    # Verificar unicidad (en los ya cargados o en DB)
                    # Si ya existe otro contacto con ese ID en memoria...
                    if ext_id in existing_contacts_by_cid and existing_contacts_by_cid[ext_id] != contact:
                         errors.append(f"Fila {idx+2}: El Contact ID '{ext_id}' ya existe en otro contacto")
                    else:
                        contact.contact_id = ext_id
                        existing_contacts_by_cid[ext_id] = contact # Actualizar índice

            # Asignar tag
            if import_tag:
                # Verificar si ya tiene el tag. 
                # Nota: acceder a contact.tags dispara query si no está cargado.
                # Para optimización extrema se podría hacer eager loading al principio join tags.
                # Al ser lazy='select', esto hará N queries si son updates. 
                # Pero como SQLAlchemy tiene identity map, si ya cargamos tags quizas reusa.
                # Una optimización simple: si es nuevo, append directo.
                if is_new:
                    contact.tags.append(import_tag)
                else:
                    if import_tag not in contact.tags:
                        contact.tags.append(import_tag)

        # 4. Commit masivo
        db.session.commit()

        message = f'Procesados {count + updated} contactos ({count} nuevos, {updated} actualizados)'
        if phone_updated > 0:
            message += f', {phone_updated} teléfonos actualizados'

        result = {'success': True, 'message': message}
        if errors:
            result['warnings'] = errors[:100] # Limitar warnings para no saturar respuesta

        return jsonify(result)

    except Exception as e:
        logger.error(f"Error importando contactos: {e}")
        return jsonify({'error': f"Error procesando archivo: {str(e)}"}), 500

@app.route("/api/contacts/export", methods=["GET"])
def api_export_contacts():
    """Exportar contactos a Excel (optimizado para grandes volúmenes).

    Incluye columna ID y Contact ID para permitir reimportar y actualizar.
    Usa eager loading y procesamiento por lotes para mejor rendimiento.
    """
    from sqlalchemy.orm import joinedload
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter
    
    try:
        # Crear workbook en modo write_only para memoria optimizada
        wb = Workbook(write_only=True)
        ws = wb.create_sheet('Contactos')
        
        # Headers
        headers = ['ID', 'Contact ID', 'Telefono', 'Nombre completo', 'Nombre', 
                   'Apellido', 'Campo 1', 'Campo 2', 'Campo 3', 'Campo 4', 
                   'Campo 5', 'Campo 6', 'Campo 7', 'Notas', 'Etiquetas', 'Fecha Creacion']
        ws.append(headers)
        
        # Procesar por lotes con eager loading de tags
        BATCH_SIZE = 5000
        offset = 0
        
        while True:
            # Cargar lote con tags pre-cargados (evita N+1 queries)
            contacts = Contact.query.options(
                joinedload(Contact.tags)
            ).order_by(Contact.id).offset(offset).limit(BATCH_SIZE).all()
            
            if not contacts:
                break
            
            for c in contacts:
                ws.append([
                    c.id,
                    c.contact_id,
                    c.phone_number,
                    c.name,
                    c.first_name,
                    c.last_name,
                    c.custom_field_1,
                    c.custom_field_2,
                    c.custom_field_3,
                    c.custom_field_4,
                    c.custom_field_5,
                    c.custom_field_6,
                    c.custom_field_7,
                    c.notes,
                    ', '.join(t.name for t in c.tags) if c.tags else '',
                    c.created_at.strftime('%Y-%m-%d %H:%M:%S') if c.created_at else ''
                ])
            
            offset += BATCH_SIZE
            
            # Liberar memoria
            db.session.expire_all()
        
        # Guardar a BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'contactos_crm_{datetime.now().strftime("%Y%m%d")}.xlsx'
        )
        
    except Exception as e:
        logger.error(f"Error exportando contactos: {e}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/contacts/template", methods=["GET"])
def api_contacts_template():
    """Descargar plantilla Excel para importar contactos."""
    try:
        # Crear datos de ejemplo - Contact ID es el identificador principal
        example_data = [
            {
                'Contact ID': 'CLI-001',
                'Telefono': '5491123456789',
                'Nombre': 'Juan',
                'Apellido': 'Perez',
                'Nombre completo': 'Juan Perez',
                'Etiquetas': 'cliente, vip',
                'Notas': 'Cliente desde 2024',
                'Campo 1': 'Valor personalizado 1',
                'Campo 2': 'Valor personalizado 2',
                'Campo 3': '',
                'Campo 4': '',
                'Campo 5': '',
                'Campo 6': '',
                'Campo 7': ''
            },
            {
                'Contact ID': 'CLI-002',
                'Telefono': '5491198765432',
                'Nombre': 'Maria',
                'Apellido': 'Garcia',
                'Nombre completo': 'Maria Garcia',
                'Etiquetas': 'prospecto',
                'Notas': 'Interesada en servicios',
                'Campo 1': '',
                'Campo 2': '',
                'Campo 3': '',
                'Campo 4': '',
                'Campo 5': '',
                'Campo 6': '',
                'Campo 7': ''
            }
        ]

        df = pd.DataFrame(example_data)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Contactos')

            # Agregar hoja de instrucciones
            instructions = pd.DataFrame({
                'Instrucciones': [
                    'PLANTILLA PARA IMPORTAR CONTACTOS',
                    '',
                    '============================================',
                    'IDENTIFICADOR PRINCIPAL: Contact ID',
                    '============================================',
                    '',
                    '- Contact ID es TU identificador para cada contacto',
                    '- Usa codigos de tu sistema: legajo, DNI, expediente, etc.',
                    '- Al importar, el sistema busca por Contact ID',
                    '- Si existe, actualiza el contacto (incluyendo telefono)',
                    '- Si no existe, crea uno nuevo',
                    '',
                    'COLUMNAS:',
                    '',
                    '- Contact ID: Tu identificador unico (legajo, DNI, codigo cliente)',
                    '- Telefono (REQUERIDO): Con codigo de pais, sin + ni espacios',
                    '  Ejemplo: 5491123456789',
                    '- Nombre, Apellido: Datos del contacto',
                    '- Etiquetas: Separadas por coma. Ej: cliente, vip',
                    '- Notas: Notas adicionales',
                    '- Campo 1-7: Campos personalizados',
                    '',
                    'COMO ACTUALIZAR CONTACTOS:',
                    '1. Exporta tus contactos actuales',
                    '2. Modifica lo que necesites (telefono, nombre, etc)',
                    '3. Mantene el Contact ID igual (es la clave de busqueda)',
                    '4. Reimporta el archivo',
                    '5. El sistema actualizara los telefonos manteniendo etiquetas y datos',
                    '',
                    'IMPORTANTE:',
                    '- Solo la columna Telefono es obligatoria',
                    '- El numero debe incluir codigo de pais (54 para Argentina)',
                    '- NO uses el simbolo + al inicio',
                    '- NO uses espacios, guiones ni parentesis',
                    '- Las etiquetas se crean automaticamente si no existen'
                ]
            })
            instructions.to_excel(writer, index=False, sheet_name='Instrucciones')

        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='plantilla_contactos.xlsx'
        )

    except Exception as e:
        logger.error(f"Error generando plantilla: {e}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/messages/<phone>")
def api_get_messages(phone):
    """API para obtener mensajes de un contacto (optimizado para AJAX)."""
    try:
        # Validar visibilidad
        if not user_can_access_phone(g.current_user, phone):
            return jsonify({'error': 'Sin acceso'}), 403

        # Límite de mensajes
        MESSAGE_LIMIT = 100

        # Obtener mensajes recientes
        recent_messages = Message.query.filter_by(phone_number=phone)\
            .order_by(Message.timestamp.desc())\
            .limit(MESSAGE_LIMIT).all()
        
        # Invertir para orden cronológico (O(n) vs O(n log n) de sorted)
        messages = recent_messages[::-1]
        
        # Obtener info de contacto
        contact = Contact.query.filter_by(phone_number=phone).first()
        contact_dict = contact.to_dict() if contact else None
        
        # Calcular stats básicos
        outbound_msgs = [m for m in messages if m.direction == 'outbound']
        stats = {
            'sent': sum(1 for m in outbound_msgs if m.latest_status in ['sent', 'delivered', 'read']),
            'delivered': sum(1 for m in outbound_msgs if m.latest_status in ['delivered', 'read']),
            'read': sum(1 for m in outbound_msgs if m.latest_status == 'read')
        }
        
        # Verificar ventana de 24hs
        can_send_free_text = False
        last_inbound_msg = None
        
        whatsapp_configured = whatsapp_api.is_configured()
        if whatsapp_configured:
            twenty_four_hours_ago = datetime.utcnow() - timedelta(hours=24)
            # Solo buscamos en lo que ya trajimos para ser rápidos, 
            # o hacemos query específica si no hay inbound recientes en los últimos 100
            
            # Buscar en los mensajes cargados primero
            inbound_loaded = [m for m in messages if m.direction == 'inbound']
            if inbound_loaded:
                last_msg = inbound_loaded[-1] # El más reciente de los cargados
                if last_msg.timestamp >= twenty_four_hours_ago:
                    can_send_free_text = True
                    last_inbound_msg = last_msg.timestamp
            
            # Si no encontramos en los últimos 100, quizás hay uno anterior pero dentro de 24h
            if not can_send_free_text:
                # Query específica rápida
                last_inbound = Message.query.filter_by(
                    phone_number=phone,
                    direction='inbound'
                ).filter(Message.timestamp >= twenty_four_hours_ago).order_by(Message.timestamp.desc()).first()
                
                if last_inbound:
                    can_send_free_text = True
                    last_inbound_msg = last_inbound.timestamp

        # Pre-cargar order_ids para mensajes de tipo order
        order_msg_ids = [m.wa_message_id for m in messages if m.message_type == 'order' and m.wa_message_id]
        order_id_by_wamid = {}
        if order_msg_ids:
            from models import Order as _Order
            order_rows = _Order.query.filter(_Order.wa_message_id.in_(order_msg_ids)).with_entities(_Order.wa_message_id, _Order.id).all()
            order_id_by_wamid = {row[0]: row[1] for row in order_rows}

        # Serializar mensajes
        messages_data = []
        for m in messages:
            # Convertir a hora argentina
            dt_arg = to_argentina_filter(m.timestamp)
            time_str = dt_arg.strftime('%H:%M') if dt_arg else ''
            date_str = dt_arg.strftime('%d/%m/%Y') if dt_arg else ''

            messages_data.append({
                'id': m.id,
                'content': m.content,
                'direction': m.direction,
                'time': time_str,
                'date': date_str,
                'status': m.latest_status,
                'message_type': m.message_type,
                'media_url': m.media_url,
                'caption': m.caption,
                'sent_by': m.sent_by,
                'order_id': order_id_by_wamid.get(m.wa_message_id) if m.message_type == 'order' else None,
            })

        # Verificar si el bot está pausado para este contacto
        bot_paused = False
        if contact:
            bot_paused = any(t.name == 'Asistencia Humana' for t in contact.tags)

        return jsonify({
            'success': True,
            'contact': contact_dict,
            'messages': messages_data,
            'stats': stats,
            'can_send_free_text': can_send_free_text,
            'whatsapp_configured': whatsapp_configured,
            'bot_paused': bot_paused
        })
        
    except Exception as e:
        logger.error(f"Error fetching messages API: {e}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/messages/<phone>/mark-read", methods=["POST"])
def api_mark_messages_read(phone):
    """Marca todos los mensajes inbound no leídos de un contacto como leídos."""
    now = datetime.utcnow()
    Message.query.filter_by(phone_number=phone, direction='inbound')\
        .filter(Message.read_at == None)\
        .update({'read_at': now}, synchronize_session=False)
    db.session.commit()
    return jsonify({'success': True})


@app.route("/api/unread-counts")
def api_unread_counts():
    """Devuelve el conteo de mensajes inbound no leídos por contacto."""
    from sqlalchemy import func
    rows = db.session.query(
        Message.phone_number,
        func.count(Message.id).label('unread')
    ).filter(
        Message.direction == 'inbound',
        Message.read_at == None
    ).group_by(Message.phone_number).all()
    return jsonify({row.phone_number: row.unread for row in rows})


@app.route("/contacts")
def contacts_page():
    """Página para ver listado de contactos con paginación y búsqueda."""
    tag_filter = request.args.get('tag')
    exclude_tag = request.args.get('exclude_tag')
    search_query = request.args.get('search', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)

    from sqlalchemy.orm import joinedload

    # Base query con joinedload para evitar N+1 en tags
    query = Contact.query.options(joinedload(Contact.tags))

    # Filtro de visibilidad por usuario
    user = g.current_user
    vis_tag_ids = get_visible_tag_ids(user)
    if vis_tag_ids is not None:
        vis_conditions = []
        if vis_tag_ids:
            vis_conditions.append(Contact.tags.any(Tag.id.in_(vis_tag_ids)))
        if user.can_see_untagged:
            vis_conditions.append(~Contact.tags.any())
        if vis_conditions:
            query = query.filter(or_(*vis_conditions))
        else:
            query = query.filter(False)

    # Apply tag filter (Include)
    if tag_filter:
        query = query.filter(Contact.tags.any(Tag.name == tag_filter))

    # Apply exclude tag filter
    if exclude_tag:
        query = query.filter(~Contact.tags.any(Tag.name == exclude_tag))

    # Apply search filter
    if search_query:
        search_pattern = f"%{search_query}%"
        query = query.filter(
            or_(
                Contact.name.ilike(search_pattern),
                Contact.phone_number.ilike(search_pattern),
                Contact.contact_id.ilike(search_pattern)
            )
        )

    # Order and paginate
    query = query.order_by(Contact.created_at.desc())
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    return render_template(
        'contacts.html',
        contacts=pagination.items,
        pagination=pagination,
        tag_filter=tag_filter,
        exclude_tag=exclude_tag,
        search_query=search_query
    )

@app.route("/tags")
def tags_page():
    """Página para ver etiquetas y estadísticas."""
    tags_with_count = db.session.query(
        Tag,
        func.count(contact_tags.c.contact_id).label('cnt')
    ).outerjoin(
        contact_tags, Tag.id == contact_tags.c.tag_id
    ).group_by(Tag.id).order_by(func.count(contact_tags.c.contact_id).desc()).all()

    active_tags = [(tag.name, cnt, tag.is_system) for tag, cnt in tags_with_count if tag.is_active]
    disabled_tags = [(tag.name, cnt, tag.is_system) for tag, cnt in tags_with_count if not tag.is_active]
    total_contacts = Contact.query.count()
    return render_template('tags.html', tags=active_tags, disabled_tags=disabled_tags, total_contacts=total_contacts)

@app.route("/api/tags", methods=["GET"])
def api_list_tags():
    """Lista etiquetas activas con conteo de contactos. Usar ?include_inactive=true para incluir deshabilitadas."""
    include_inactive = request.args.get('include_inactive', 'false').lower() == 'true'
    
    query = db.session.query(
        Tag,
        func.count(contact_tags.c.contact_id).label('cnt')
    ).outerjoin(
        contact_tags, Tag.id == contact_tags.c.tag_id
    )
    
    if not include_inactive:
        query = query.filter(Tag.is_active == True)
    
    tags_with_count = query.group_by(Tag.id).order_by(func.count(contact_tags.c.contact_id).desc()).all()

    return jsonify([{'name': tag.name, 'count': cnt, 'is_active': tag.is_active, 'is_system': tag.is_system} for tag, cnt in tags_with_count])

@app.route("/api/tags", methods=["POST"])
def api_create_tag():
    """Crea una nueva etiqueta."""
    data = request.json
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'error': 'Nombre requerido'}), 400

    if Tag.query.filter_by(name=name).first():
        return jsonify({'error': 'Etiqueta ya existe'}), 409

    tag = Tag(name=name)
    db.session.add(tag)
    db.session.commit()
    return jsonify({'success': True, 'name': name}), 201

@app.route("/api/tags/<tag_name>", methods=["DELETE"])
def api_delete_tag(tag_name):
    """Elimina o deshabilita una etiqueta. Si tiene campañas, la deshabilita en lugar de eliminarla."""
    try:
        tag = Tag.query.filter_by(name=tag_name).first()
        if not tag:
            return jsonify({'error': 'Tag no encontrado'}), 404

        # Proteger tags del sistema
        if tag.is_system:
            return jsonify({'error': 'No se puede eliminar una etiqueta del sistema'}), 403

        # Verificar si hay campañas asociadas a este tag
        campaigns_count = Campaign.query.filter_by(tag_id=tag.id).count()
        
        # Quitar tag de todos los contactos
        removed_count = db.session.execute(
            contact_tags.delete().where(contact_tags.c.tag_id == tag.id)
        ).rowcount
        
        if campaigns_count > 0:
            # Deshabilitar en lugar de eliminar
            tag.is_active = False
            db.session.commit()
            logger.info(f"🏷️ Tag '{tag_name}' deshabilitado (tiene {campaigns_count} campañas). {removed_count} contactos desvinculados.")
            return jsonify({
                'success': True,
                'action': 'disabled',
                'message': f'Etiqueta deshabilitada ({campaigns_count} campañas asociadas). Se quitó de {removed_count} contacto(s).'
            })
        else:
            # Eliminar permanentemente
            db.session.delete(tag)
            db.session.commit()
            logger.info(f"🗑️ Tag '{tag_name}' eliminado permanentemente. {removed_count} contactos desvinculados.")
            return jsonify({'success': True, 'action': 'deleted'})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error eliminando/deshabilitando tag '{tag_name}': {e}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/tags/<tag_name>/toggle", methods=["POST"])
def api_toggle_tag(tag_name):
    """Rehabilita o deshabilita una etiqueta."""
    try:
        tag = Tag.query.filter_by(name=tag_name).first()
        if not tag:
            return jsonify({'error': 'Tag no encontrado'}), 404

        # Proteger tags del sistema
        if tag.is_system:
            return jsonify({'error': 'No se puede deshabilitar una etiqueta del sistema'}), 403

        if tag.is_active:
            # Deshabilitar: quitar de todos los contactos
            removed = db.session.execute(
                contact_tags.delete().where(contact_tags.c.tag_id == tag.id)
            ).rowcount
            tag.is_active = False
            db.session.commit()
            logger.info(f"🏷️ Tag '{tag_name}' deshabilitado. {removed} contactos desvinculados.")
            return jsonify({
                'success': True,
                'is_active': False,
                'message': f'Etiqueta deshabilitada. Se quitó de {removed} contacto(s).'
            })
        else:
            # Rehabilitar
            tag.is_active = True
            db.session.commit()
            logger.info(f"✅ Tag '{tag_name}' rehabilitado.")
            return jsonify({
                'success': True,
                'is_active': True,
                'message': 'Etiqueta rehabilitada exitosamente.'
            })
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error toggling tag '{tag_name}': {e}")
        return jsonify({'error': str(e)}), 500


# ==========================================
# ASISTENCIA HUMANA — Escalación y Bot Status
# ==========================================

@app.route("/api/escalate-to-human", methods=["POST"])
def api_escalate_to_human():
    """Asigna etiqueta 'Asistencia Humana' a un contacto. Llamado por n8n."""
    try:
        data = request.get_json() or {}
        phone = (data.get('phone_number') or '').strip()
        if not phone:
            return jsonify({'error': 'phone_number is required'}), 400

        # Normalizar número (eliminar '+' inicial si viene de n8n)
        phone_normalized = normalize_phone(phone)

        # Buscar contacto tolerando variantes con/sin '+'
        contact = find_contact_by_phone(phone_normalized)
        if not contact:
            # Crear con el número normalizado (sin '+')
            contact = Contact(phone_number=phone_normalized)
            db.session.add(contact)
            db.session.flush()
            logger.info(f"Created new contact for escalation: {phone_normalized} (original: {phone})")

        # Buscar tag del sistema
        tag = Tag.query.filter_by(name='Asistencia Humana').first()
        if not tag:
            return jsonify({'error': 'System tag not found'}), 500

        # Agregar tag si no la tiene
        if tag not in contact.tags:
            contact.tags.append(tag)
            db.session.commit()
            logger.info(f"Escalated to human: {phone}")

        return jsonify({'success': True, 'phone': phone})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error escalating to human: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/contacts/assign-tag", methods=["POST"])
def api_assign_tag():
    """Asigna un tag a un contacto por número de teléfono. Llamado por el bot de n8n."""
    try:
        data = request.get_json() or {}
        phone = (data.get('phone_number') or '').strip()
        tag_name = (data.get('tag_name') or '').strip()

        if not phone or not tag_name:
            return jsonify({'error': 'phone_number y tag_name son requeridos'}), 400

        phone_normalized = normalize_phone(phone)
        contact = find_contact_by_phone(phone_normalized)
        if not contact:
            contact = Contact(phone_number=phone_normalized)
            db.session.add(contact)
            db.session.flush()

        tag = Tag.query.filter_by(name=tag_name).first()
        if not tag:
            tag = Tag(name=tag_name, color='blue', is_active=True)
            db.session.add(tag)
            db.session.flush()

        if tag not in contact.tags:
            contact.tags.append(tag)

        db.session.commit()
        logger.info(f"Tag '{tag_name}' asignado a {phone_normalized}")
        return jsonify({'success': True, 'phone': phone_normalized, 'tag': tag_name})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error asignando tag: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/contacts/<phone>/toggle-tag", methods=["POST"])
def api_chat_toggle_tag(phone):
    """Agrega o quita un tag de un contacto desde el chat. Retorna el estado nuevo."""
    try:
        data = request.get_json() or {}
        tag_id = data.get('tag_id')
        if not tag_id:
            return jsonify({'error': 'tag_id requerido'}), 400

        phone_normalized = normalize_phone(phone)
        contact = find_contact_by_phone(phone_normalized)
        if not contact:
            contact = Contact(phone_number=phone_normalized)
            db.session.add(contact)
            db.session.flush()

        tag = Tag.query.get(tag_id)
        if not tag:
            return jsonify({'error': 'Tag no encontrado'}), 404

        current_user_name = g.current_user.username if g.current_user else 'manual'
        if tag in contact.tags:
            contact.tags.remove(tag)
            assigned = False
            _record_tag_history(contact.id, tag, 'removed', 'manual', current_user_name)
        else:
            contact.tags.append(tag)
            assigned = True
            _record_tag_history(contact.id, tag, 'added', 'manual', current_user_name)

        db.session.commit()

        # Si se asignó la etiqueta manualmente, enrollar en secuencias activas
        if assigned:
            from auto_tagger import enroll_in_sequences
            enroll_in_sequences(db, contact, tag_id, FollowUpSequence, FollowUpEnrollment)

        return jsonify({'success': True, 'assigned': assigned, 'tag_id': tag_id, 'tag_name': tag.name})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error toggling tag en chat: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/tags/active", methods=["GET"])
def api_tags_active():
    """Lista todos los tags activos. Usado por el panel de tags del chat."""
    try:
        include_system = request.args.get('include_system', 'false').lower() == 'true'
        q = Tag.query.filter_by(is_active=True)
        if not include_system:
            q = q.filter_by(is_system=False)
        tags = q.order_by(Tag.name).all()
        return jsonify([{'id': t.id, 'name': t.name, 'color': t.color} for t in tags])
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route("/api/contact/<phone>/bot-status", methods=["GET"])
def api_bot_status(phone):
    """Retorna si el bot está activo para este contacto. Llamado por n8n."""
    try:
        phone_normalized = normalize_phone(phone)
        contact = find_contact_by_phone(phone_normalized)
        if not contact:
            logger.info(f"[BOT-STATUS] {phone} → normalizado: {phone_normalized} → contacto NO encontrado → bot_active=True")
            return jsonify({'bot_active': True, 'phone': phone_normalized})

        # Verificar si tiene la etiqueta "Asistencia Humana"
        contact_tags = [t.name for t in contact.tags]
        has_human_tag = any(t == 'Asistencia Humana' for t in contact_tags)

        logger.info(f"[BOT-STATUS] {phone} → normalizado: {phone_normalized} → contacto id={contact.id} nombre='{contact.name}' → tags={contact_tags} → has_human_tag={has_human_tag} → bot_active={not has_human_tag}")

        return jsonify({
            'bot_active': not has_human_tag,
            'phone': phone_normalized,
            'has_human_assistance_tag': has_human_tag
        })
    except Exception as e:
        logger.error(f"[BOT-STATUS] Error checking bot status para {phone}: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/contact/<phone>/pause-bot", methods=["POST"])
def api_pause_bot(phone):
    """Agrega etiqueta 'Asistencia Humana' al contacto. Llamado desde dashboard."""
    try:
        phone_normalized = normalize_phone(phone)
        if not user_can_access_phone(g.current_user, phone_normalized):
            return jsonify({'error': 'Sin acceso'}), 403
        contact = find_contact_by_phone(phone_normalized)
        if not contact:
            return jsonify({'error': 'Contacto no encontrado'}), 404

        tag = Tag.query.filter_by(name='Asistencia Humana').first()
        if not tag:
            return jsonify({'error': 'System tag not found'}), 500

        if tag not in contact.tags:
            contact.tags.append(tag)
            db.session.commit()
            logger.info(f"Bot paused for: {phone}")

        return jsonify({'success': True, 'phone': phone})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error pausing bot: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/contact/<phone>/resume-bot", methods=["POST"])
def api_resume_bot(phone):
    """Quita etiqueta 'Asistencia Humana' del contacto. Llamado desde dashboard."""
    try:
        phone_normalized = normalize_phone(phone)
        if not user_can_access_phone(g.current_user, phone_normalized):
            return jsonify({'error': 'Sin acceso'}), 403
        contact = find_contact_by_phone(phone_normalized)
        if not contact:
            return jsonify({'error': 'Contacto no encontrado'}), 404

        tag = Tag.query.filter_by(name='Asistencia Humana').first()
        if not tag:
            return jsonify({'error': 'System tag not found'}), 500

        if tag in contact.tags:
            contact.tags.remove(tag)
            db.session.commit()
            logger.info(f"Bot resumed for: {phone}")

        return jsonify({'success': True, 'phone': phone})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error resuming bot: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/human-assistance/pending", methods=["GET"])
def api_human_assistance_pending():
    """Retorna la cantidad y lista de contactos que necesitan asistencia humana."""
    try:
        tag = Tag.query.filter_by(name='Asistencia Humana').first()
        if not tag:
            return jsonify({'count': 0, 'contacts': []})

        contacts_with_tag = tag.contacts  # relación many-to-many
        result = []
        for c in contacts_with_tag:
            result.append({
                'phone_number': c.phone_number,
                'name': c.name or c.phone_number
            })

        return jsonify({'count': len(result), 'contacts': result})
    except Exception as e:
        logger.error(f"Error getting human assistance pending: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/contacts/bulk-delete", methods=["POST"])
def api_bulk_delete_contacts():
    """Eliminar múltiples contactos de una vez."""
    data = request.json
    contact_ids = data.get('contact_ids', [])

    if not contact_ids:
        return jsonify({'error': 'contact_ids requeridos'}), 400

    try:
        # Eliminar registros relacionados en campaign_logs
        CampaignLog.query.filter(CampaignLog.contact_id.in_(contact_ids)).delete(synchronize_session=False)

        # Eliminar contactos
        deleted = Contact.query.filter(Contact.id.in_(contact_ids)).delete(synchronize_session=False)
        db.session.commit()

        logger.info(f"🗑️ {deleted} contactos eliminados en lote")
        return jsonify({'success': True, 'deleted': deleted})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error eliminando contactos en lote: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/contacts/bulk-tags", methods=["POST"])
def api_bulk_tags():
    """Asignar o remover tags de múltiples contactos.

    Acepta contact_ids (preferido) o phones (legacy) para identificar contactos.
    """
    data = request.json
    contact_ids = data.get('contact_ids', [])
    phones = data.get('phones', [])  # Legacy support
    tag_name = (data.get('tag') or '').strip()
    action = data.get('action')  # 'add' or 'remove'

    if (not contact_ids and not phones) or not tag_name or action not in ('add', 'remove'):
        return jsonify({'error': 'contact_ids (o phones), tag y action (add/remove) requeridos'}), 400

    try:
        # Preferir IDs sobre phones
        if contact_ids:
            contacts = Contact.query.filter(Contact.id.in_(contact_ids)).all()
        else:
            contacts = Contact.query.filter(Contact.phone_number.in_(phones)).all()

        # Prepare IDs
        target_contact_ids = [c.id for c in contacts]
        if not target_contact_ids:
             return jsonify({'success': True, 'affected': 0})

        if action == 'add':
            tag = Tag.query.filter_by(name=tag_name).first()
            if not tag:
                tag = Tag(name=tag_name)
                db.session.add(tag)
                db.session.flush()
            
            # Find IDs that already have this tag
            existing_relations = db.session.query(contact_tags.c.contact_id).filter(
                contact_tags.c.tag_id == tag.id,
                contact_tags.c.contact_id.in_(target_contact_ids)
            ).all()
            existing_ids = {r[0] for r in existing_relations}

            # Insert only new relations
            new_relations = [
                {'contact_id': cid, 'tag_id': tag.id} 
                for cid in target_contact_ids if cid not in existing_ids
            ]
            
            if new_relations:
                db.session.execute(contact_tags.insert(), new_relations)

        elif action == 'remove':
            tag = Tag.query.filter_by(name=tag_name).first()
            if tag:
                db.session.execute(
                    contact_tags.delete().where(
                        and_(
                            contact_tags.c.tag_id == tag.id,
                            contact_tags.c.contact_id.in_(target_contact_ids)
                        )
                    )
                )

        db.session.commit()
        return jsonify({'success': True, 'affected': len(contacts)})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route("/api/tags/bulk-action", methods=["POST"])
def api_tags_bulk_action():
    """Agregar o quitar etiqueta de múltiples contactos via archivo Excel/CSV con optimización por lotes."""
    tag_name = request.form.get('tag_name', '').strip()
    action = request.form.get('action', '').strip()

    if not tag_name or action not in ('add', 'remove'):
        return jsonify({'error': 'tag_name y action (add/remove) requeridos'}), 400

    if 'file' not in request.files or request.files['file'].filename == '':
        return jsonify({'error': 'Archivo requerido'}), 400

    file = request.files['file']
    try:
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)

        # Normalizar columnas
        df.columns = [str(c).lower().strip() for c in df.columns]

        # Buscar columna de Contact ID (PRIORIDAD)
        contact_id_candidates = ['contact id', 'contact_id', 'contactid', 'id externo', 'external_id']
        contact_id_col = next((c for c in df.columns if c in contact_id_candidates), None)

        # Buscar columna de teléfono (fallback)
        phone_candidates = ['telefono', 'phone', 'phone_number', 'numero', 'celular']
        phone_col = next((c for c in df.columns if c in phone_candidates), None)

        if not contact_id_col and not phone_col:
            return jsonify({'error': 'Se requiere columna "Contact ID" o "Telefono"'}), 400

        # Obtener o crear tag
        tag = Tag.query.filter_by(name=tag_name).first()
        if not tag:
            if action == 'remove':
                return jsonify({'success': True, 'added': 0, 'removed': 0, 'skipped': 0})
            tag = Tag(name=tag_name)
            db.session.add(tag)
            db.session.flush()

        # =====================================================
        # OPTIMIZACIÓN POR LOTES
        # =====================================================

        # 1. Recolectar identificadores
        all_contact_ids = set()
        all_phones = set()

        # Limpiar datos
        if contact_id_col:
            df['clean_contact_id'] = df[contact_id_col].apply(lambda x: str(x).strip() if pd.notna(x) else None)
            all_contact_ids = {cid for cid in df['clean_contact_id'].dropna().unique() if cid}

        if phone_col:
            df['clean_phone'] = df[phone_col].apply(lambda x: str(x).replace('.0', '').strip() if pd.notna(x) else '')
            all_phones = {p for p in df['clean_phone'].dropna().unique() if p}

        # 2. Batch Fetch de contactos existentes
        existing_contacts_by_cid = {}
        existing_contacts_by_phone = {}

        # Fetch por Contact ID
        if all_contact_ids:
            results_cid = Contact.query.filter(Contact.contact_id.in_(all_contact_ids)).all()
            for c in results_cid:
                existing_contacts_by_cid[c.contact_id] = c
                existing_contacts_by_phone[c.phone_number] = c  # Indexar también por teléfono

        # Fetch por Phone (evitando duplicados)
        phones_to_fetch = all_phones - set(existing_contacts_by_phone.keys())
        if phones_to_fetch:
             results_phone = Contact.query.filter(Contact.phone_number.in_(phones_to_fetch)).all()
             for c in results_phone:
                 existing_contacts_by_phone[c.phone_number] = c
                 if c.contact_id:
                     existing_contacts_by_cid[c.contact_id] = c

        added = 0
        removed = 0
        skipped = 0
        not_found = 0

        # 3. Recolectar IDs de contactos encontrados
        found_contact_ids = set()
        
        for _, row in df.iterrows():
            contact = None

            # Buscar
            if contact_id_col:
                ext_id = row.get('clean_contact_id')
                if ext_id and ext_id in existing_contacts_by_cid:
                    contact = existing_contacts_by_cid[ext_id]
            
            if not contact and phone_col:
                phone = row.get('clean_phone')
                if phone and phone in existing_contacts_by_phone:
                    contact = existing_contacts_by_phone[phone]

            if not contact:
                not_found += 1
                continue

            found_contact_ids.add(contact.id)
        
        # 4. Ejecutar operación en lote según acción
        if action == 'add':
            # Obtener contactos que YA tienen el tag
            existing_relations = db.session.query(contact_tags.c.contact_id).filter(
                contact_tags.c.tag_id == tag.id,
                contact_tags.c.contact_id.in_(found_contact_ids)
            ).all()
            existing_ids = {r[0] for r in existing_relations}
            
            # Insertar solo relaciones nuevas
            new_relations = [
                {'contact_id': cid, 'tag_id': tag.id} 
                for cid in found_contact_ids if cid not in existing_ids
            ]
            
            if new_relations:
                db.session.execute(contact_tags.insert(), new_relations)
            
            added = len(new_relations)
            skipped = len(existing_ids)
            
        elif action == 'remove':
            # Obtener contactos que tienen el tag
            existing_relations = db.session.query(contact_tags.c.contact_id).filter(
                contact_tags.c.tag_id == tag.id,
                contact_tags.c.contact_id.in_(found_contact_ids)
            ).all()
            contacts_with_tag = {r[0] for r in existing_relations}
            
            # Eliminar en lote con un solo DELETE
            if contacts_with_tag:
                db.session.execute(
                    contact_tags.delete().where(
                        and_(
                            contact_tags.c.tag_id == tag.id,
                            contact_tags.c.contact_id.in_(contacts_with_tag)
                        )
                    )
                )
            
            removed = len(contacts_with_tag)
            skipped = len(found_contact_ids) - len(contacts_with_tag)

        db.session.commit()
        
        return jsonify({
            'success': True,
            'added': added,
            'removed': removed,
            'skipped': skipped,
            'not_found': not_found
        })

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error en bulk tag action: {e}")
        return jsonify({'error': str(e)}), 500

@app.route("/failed-messages")
def failed_messages_page():
    """Página para ver mensajes fallidos."""
    # Buscar mensajes con estado 'failed' en su último status
    # Hacemos un join con MessageStatus
    
    # Subquery para obtener el último status de cada mensaje
    last_status_subquery = db.session.query(
        MessageStatus.wa_message_id,
        func.max(MessageStatus.timestamp).label('max_timestamp')
    ).group_by(MessageStatus.wa_message_id).subquery()
    
    failed_msgs = db.session.query(Message, MessageStatus).join(
        MessageStatus, Message.wa_message_id == MessageStatus.wa_message_id
    ).join(
        last_status_subquery, 
        (MessageStatus.wa_message_id == last_status_subquery.c.wa_message_id) & 
        (MessageStatus.timestamp == last_status_subquery.c.max_timestamp)
    ).filter(
        MessageStatus.status == 'failed'
    ).order_by(Message.timestamp.desc()).all()
    
    # Batch load contactos para evitar N+1 queries
    phones = list({msg.phone_number for msg, _ in failed_msgs})
    contacts_map = {}
    if phones:
        contacts_map = {c.phone_number: c for c in Contact.query.filter(Contact.phone_number.in_(phones)).all()}

    enriched_failures = []
    for msg, status in failed_msgs:
        contact = contacts_map.get(msg.phone_number)
        enriched_failures.append({
            'message': msg,
            'status': status,
            'contact_name': contact.name if contact else None
        })
        
    return render_template('failed_messages.html', failures=enriched_failures)

@app.route("/chatwoot-webhook", methods=["POST"])
def chatwoot_webhook():
    """
    Endpoint para recibir webhooks de Chatwoot.
    Captura mensajes salientes enviados por agentes.
    """
    try:
        data = request.json
        if not data:
            return "No data received", 400
        
        event = data.get("event")
        logger.info(f"📬 CHATWOOT WEBHOOK: {event}")
        
        # Manejar mensaje creado o actualizado
        if event in ["message_created", "message_updated"]:
            # En Chatwoot, el contenido viene en el nivel raíz
            content = data.get("content", "")
            conversation = data.get("conversation", {})
            
            # Obtener el message_type desde messages[0] dentro de conversation
            messages = conversation.get("messages", [])
            message_type_cw = None
            source_id = None
            cw_msg_id = None
            
            if messages:
                first_msg = messages[0]
                message_type_cw = first_msg.get("message_type")
                source_id = first_msg.get("source_id")  # wa_message_id de WhatsApp
                cw_msg_id = first_msg.get("id")  # ID interno de Chatwoot
            
            # Obtener el número de teléfono desde contact_inbox
            contact_inbox = conversation.get("contact_inbox", {})
            phone_number = contact_inbox.get("source_id", "").replace("+", "")
            
            logger.info(f"📝 message_type: {message_type_cw}, cw_id: {cw_msg_id}, source_id: {source_id}")
            
            # Solo mensajes salientes (message_type=1 en Chatwoot)
            if message_type_cw == 1 and content and cw_msg_id:
                logger.info(f"📤 MENSAJE SALIENTE: '{content[:50]}...' para {phone_number}")
                
                # Buscar primero por source_id (wa_message_id) si existe
                existing = None
                if source_id:
                    existing = Message.query.filter_by(wa_message_id=source_id).first()
                
                # Si no existe por source_id, buscar por cw_id
                cw_id_str = f"cw_{cw_msg_id}"
                if not existing:
                    existing = Message.query.filter_by(wa_message_id=cw_id_str).first()
                
                if existing:
                    # Actualizar mensaje existente
                    updated = False
                    if not existing.content and content:
                        existing.content = content
                        updated = True
                    if phone_number and existing.phone_number in ["outbound", "unknown"]:
                        existing.phone_number = phone_number
                        updated = True
                    # Si tenemos source_id y el mensaje tenía cw_id, actualizar al wa_message_id real
                    if source_id and existing.wa_message_id.startswith("cw_"):
                        existing.wa_message_id = source_id
                        updated = True
                    if updated:
                        db.session.commit()
                        logger.info(f"✅ Mensaje actualizado: {existing.wa_message_id}")
                else:
                    # Crear nuevo mensaje
                    # Usar source_id si está disponible, sino usar cw_id
                    msg_id = source_id if source_id else cw_id_str
                    
                    # Intentar detectar si el contenido es un template
                    detected_type = "text"
                    final_content = content
                    
                    try:
                        templates_result = whatsapp_api.get_templates()
                        for t in templates_result.get("templates", []):
                            if t.get("status") == "APPROVED":
                                for comp in t.get("components", []):
                                    if comp.get("type") == "BODY":
                                        template_body = comp.get("text", "")
                                        # Limpieza básica para comparación (quitar variables {{1}}, etc)
                                        import re
                                        pattern = re.escape(template_body)
                                        pattern = re.sub(r'\\\{\\\{\d+\\\}\\\}', '.*', pattern)
                                        
                                        if re.match(f"^{pattern}$", content, re.DOTALL):
                                            detected_type = "template"
                                            # Opcional: podríamos normalizar el contenido al template original
                                            # pero mejor dejar el texto real enviado.
                                            break
                                if detected_type == "template": break
                    except Exception as te:
                        logger.error(f"Error detecting template in webhook: {te}")

                    new_msg = Message(
                        wa_message_id=msg_id,
                        phone_number=phone_number or "unknown",
                        direction="outbound",
                        message_type=detected_type,
                        content=final_content,
                        timestamp=datetime.utcnow()
                    )
                    db.session.add(new_msg)
                    db.session.commit()
                    logger.info(f"✅ Mensaje saliente creado ({detected_type}): {msg_id}")
        
        return "OK", 200
        
    except Exception as e:
        logger.error(f"Error en chatwoot webhook: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return "Internal Server Error", 500

# ==================== Admin: Usuarios ====================

AVAILABLE_PERMISSIONS = [
    ('dashboard',      'Dashboard',          'chat'),
    ('contacts',       'Contactos',          'group'),
    ('tags',           'Etiquetas',          'label'),
    ('campaigns',      'Campañas',           'campaign'),
    ('failed_messages','Mensajes Fallidos',  'error'),
    ('analytics',      'Analytics',          'bar_chart'),
    ('sessions',       'Sesiones',           'forum'),
    ('topics',         'Temas',              'category'),
    ('chatbot',        'Chatbot',            'smart_toy'),
    ('reengagement',   'Re-engagement',      'autorenew'),
    ('orders',         'Órdenes',            'shopping_bag'),
    ('catalog',        'Catálogo',           'inventory_2'),
    ('settings',       'Configuración',      'settings'),
]

@app.route("/admin/users")
def admin_users():
    from models import CrmUser
    users = CrmUser.query.order_by(CrmUser.created_at).all()
    return render_template('admin_users.html', users=users, available_permissions=AVAILABLE_PERMISSIONS)

@app.route("/api/admin/users", methods=["POST"])
def api_admin_create_user():
    from models import CrmUser, CrmUserPermission
    data = request.json
    if not data:
        return jsonify({'error': 'No data'}), 400
    username = data.get('username', '').strip()
    display_name = data.get('display_name', '').strip()
    password = data.get('password', '')
    is_admin = data.get('is_admin', False)
    permissions = data.get('permissions', [])

    if not username or not display_name or not password:
        return jsonify({'error': 'username, display_name y password son requeridos'}), 400
    if CrmUser.query.filter_by(username=username).first():
        return jsonify({'error': 'El usuario ya existe'}), 409

    can_see_untagged = data.get('can_see_untagged', False)
    tag_visibility = data.get('tag_visibility', [])

    user = CrmUser(username=username, display_name=display_name, is_admin=is_admin, is_active=True, can_see_untagged=can_see_untagged)
    user.set_password(password)
    db.session.add(user)
    db.session.flush()
    for perm in permissions:
        db.session.add(CrmUserPermission(user_id=user.id, permission=perm))
    for tid in tag_visibility:
        if tid is not None:
            db.session.add(CrmUserTagVisibility(user_id=user.id, tag_id=int(tid)))
    db.session.commit()
    return jsonify({'success': True, 'user': user.to_dict()})

@app.route("/api/admin/users/<int:user_id>", methods=["PUT"])
def api_admin_update_user(user_id):
    from models import CrmUser, CrmUserPermission
    user = CrmUser.query.get_or_404(user_id)
    data = request.json or {}

    if 'display_name' in data:
        user.display_name = data['display_name'].strip()
    if 'is_admin' in data:
        user.is_admin = data['is_admin']
    if 'is_active' in data:
        user.is_active = data['is_active']
    if 'password' in data and data['password']:
        user.set_password(data['password'])
    if 'permissions' in data:
        CrmUserPermission.query.filter_by(user_id=user.id).delete()
        for perm in data['permissions']:
            db.session.add(CrmUserPermission(user_id=user.id, permission=perm))
    if 'can_see_untagged' in data:
        user.can_see_untagged = bool(data['can_see_untagged'])
    if 'tag_visibility' in data:
        CrmUserTagVisibility.query.filter_by(user_id=user.id).delete()
        for tid in (data['tag_visibility'] or []):
            if tid is not None:
                db.session.add(CrmUserTagVisibility(user_id=user.id, tag_id=int(tid)))

    db.session.commit()
    return jsonify({'success': True, 'user': user.to_dict()})

@app.route("/api/admin/users/<int:user_id>", methods=["DELETE"])
def api_admin_delete_user(user_id):
    from models import CrmUser
    user = CrmUser.query.get_or_404(user_id)
    # No borrar el último admin
    if user.is_admin and CrmUser.query.filter_by(is_admin=True, is_active=True).count() <= 1:
        return jsonify({'error': 'No podés eliminar el único administrador activo'}), 400
    db.session.delete(user)
    db.session.commit()
    return jsonify({'success': True})

# ==================== WhatsApp Settings ====================

@app.route("/whatsapp-settings")
def whatsapp_settings():
    """Página de configuración y templates de WhatsApp."""
    is_configured = whatsapp_api.is_configured()
    
    templates = []
    phone_numbers = []
    profile = {}
    error = None
    
    if is_configured:
        # Obtener templates
        templates_result = whatsapp_api.get_templates()
        if "error" in templates_result:
            error = templates_result["error"]
        templates = templates_result.get("templates", [])
        
        # Obtener números
        numbers_result = whatsapp_api.get_phone_numbers()
        phone_numbers = numbers_result.get("phone_numbers", [])
        
        # Obtener perfil
        profile_result = whatsapp_api.get_business_profile()
        profile = profile_result.get("profile", {})
    
    return render_template('whatsapp_settings.html',
                         is_configured=is_configured,
                         templates=templates,
                         phone_numbers=phone_numbers,
                         profile=profile,
                         error=error)

@app.route("/templates/new")
def create_template_page():
    """Página dedicada para crear nuevas plantillas de WhatsApp con vista previa en vivo."""
    return render_template('create_template.html')

@app.route("/api/whatsapp/templates")
def api_whatsapp_templates():
    """API para obtener templates."""
    return jsonify(whatsapp_api.get_templates())

@app.route("/api/whatsapp/phone-numbers")
def api_whatsapp_phone_numbers():
    """API para obtener números de teléfono."""
    return jsonify(whatsapp_api.get_phone_numbers())

@app.route("/api/whatsapp/profile")
def api_whatsapp_profile():
    """API para obtener perfil del negocio."""
    return jsonify(whatsapp_api.get_business_profile())

@app.route("/api/whatsapp/profile", methods=["POST"])
def api_update_whatsapp_profile():
    """API para actualizar el perfil de WhatsApp Business."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    # Validar campos opcionales
    profile_data = {}

    # About (descripción corta) - máx 256 caracteres
    if "about" in data:
        about = data["about"].strip()
        if len(about) > 256:
            return jsonify({"error": "about no puede exceder 256 caracteres"}), 400
        profile_data["about"] = about

    # Description (descripción larga) - máx 512 caracteres
    if "description" in data:
        description = data["description"].strip()
        if len(description) > 512:
            return jsonify({"error": "description no puede exceder 512 caracteres"}), 400
        profile_data["description"] = description

    # Address (dirección)
    if "address" in data:
        profile_data["address"] = data["address"].strip()

    # Email
    if "email" in data:
        email = data["email"].strip()
        if email and "@" not in email:
            return jsonify({"error": "email inválido"}), 400
        profile_data["email"] = email

    # Vertical (categoría del negocio)
    if "vertical" in data:
        profile_data["vertical"] = data["vertical"]

    # Websites - máx 2 URLs
    if "websites" in data:
        websites = data["websites"]
        if isinstance(websites, str):
            # Si es un string, dividir por comas o crear lista de 1 elemento
            websites = [w.strip() for w in websites.split(",") if w.strip()]
        if not isinstance(websites, list):
            return jsonify({"error": "websites debe ser una lista o string separado por comas"}), 400
        if len(websites) > 2:
            return jsonify({"error": "Solo se permiten hasta 2 sitios web"}), 400
        profile_data["websites"] = websites

    if not profile_data:
        return jsonify({"error": "No se proporcionaron campos para actualizar"}), 400

    result = whatsapp_api.update_business_profile(profile_data)

    if result.get("error"):
        return jsonify(result), 400
    return jsonify(result)

@app.route("/api/whatsapp/create-template", methods=["POST"])
def api_create_template():
    """API para crear una nueva plantilla de mensaje."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    name = data.get("name", "").strip().lower().replace(" ", "_")
    category = data.get("category")
    language = data.get("language", "es_AR")

    if not name or not category:
        return jsonify({"error": "name y category son requeridos"}), 400

    if category not in ["MARKETING", "UTILITY", "AUTHENTICATION"]:
        return jsonify({"error": "category debe ser MARKETING, UTILITY o AUTHENTICATION"}), 400

    # Construir componentes
    components = []

    # Header (opcional)
    header = data.get("header")
    if header and header.get("format"):
        header_component = {"type": "HEADER", "format": header["format"]}
        if header["format"] == "TEXT" and header.get("text"):
            header_component["text"] = header["text"]
        components.append(header_component)

    # Body (requerido)
    body_text = data.get("body", "").strip()
    if not body_text:
        return jsonify({"error": "body es requerido"}), 400
    body_component = {"type": "BODY", "text": body_text}

    # Agregar ejemplos de variables si existen (Meta los requiere para aprobar templates con variables)
    body_examples = data.get("body_examples", [])
    if body_examples:
        body_component["example"] = {"body_text": [body_examples]}

    components.append(body_component)

    # Footer (opcional)
    footer_text = data.get("footer", "").strip()
    if footer_text:
        components.append({"type": "FOOTER", "text": footer_text})

    # Buttons (opcional)
    buttons = data.get("buttons", [])
    if buttons:
        button_components = []
        for btn in buttons:
            if btn.get("type") == "QUICK_REPLY" and btn.get("text"):
                button_components.append({"type": "QUICK_REPLY", "text": btn["text"]})
            elif btn.get("type") == "URL" and btn.get("text") and btn.get("url"):
                btn_url = btn["url"].lower()
                if "wa.me" in btn_url or "whatsapp.com" in btn_url:
                    return jsonify({"error": "No se permiten enlaces directos a WhatsApp en los botones de templates."}), 400
                button_components.append({"type": "URL", "text": btn["text"], "url": btn["url"]})
            elif btn.get("type") == "PHONE_NUMBER" and btn.get("text") and btn.get("phone_number"):
                button_components.append({"type": "PHONE_NUMBER", "text": btn["text"], "phone_number": btn["phone_number"]})
        if button_components:
            components.append({"type": "BUTTONS", "buttons": button_components})

    result = whatsapp_api.create_template(name, category, language, components)

    if result.get("error"):
        return jsonify(result), 400
    return jsonify(result)

@app.route("/api/whatsapp/send-template", methods=["POST"])
def api_send_template():
    """API para enviar mensaje con template y variables dinámicas."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    to_phone = data.get("to")
    template_name = data.get("template_name")
    language = data.get("language", "es_AR")
    # variable_mapping es una lista de nombres de campos del contacto, ej: ['first_name', 'custom_field_1']
    variable_mapping = data.get("variable_mapping", []) 
    
    if not to_phone or not template_name:
        return jsonify({"error": "to y template_name son requeridos"}), 400

    _cu = getattr(g, 'current_user', None)
    if _cu and not user_can_access_phone(_cu, to_phone):
        return jsonify({"error": "Sin acceso"}), 403

    # Construir componentes si hay mapeo de variables
    components = data.get("components") # Permitir componentes manuales si se envían
    
    if variable_mapping and not components:
        contact = Contact.query.filter_by(phone_number=to_phone).first()
        parameters = []
        
        for field in variable_mapping:
            value = "-" # Fallback para evitar errores de API
            
            # Valores especiales
            if field == 'phone_number':
                value = to_phone
            elif contact:
                # Obtener valor del atributo del contacto
                val = getattr(contact, field, None)
                if val:
                    value = str(val)
            
            parameters.append({
                "type": "text",
                "text": value
            })
            
        if parameters:
            components = [{
                "type": "body",
                "parameters": parameters
            }]
    
    # Obtener el contenido del template para guardarlo en el historial
    # Esto es una aproximación, ya que no tenemos el texto final renderizado por WhatsApp
    template_content = f"[Template: {template_name}]"
    templates_result = whatsapp_api.get_templates()
    for t in templates_result.get("templates", []):
        if t.get("name") == template_name and t.get("language") == language:
            for comp in t.get("components", []):
                if comp.get("type") == "BODY":
                    text = comp.get("text", "")
                    # Intentar rellenar variables para el historial local
                    if variable_mapping and contact:
                        # Caso 1: variable_mapping con campos del contacto
                        for i, field in enumerate(variable_mapping):
                            val = getattr(contact, field, "") or "-"
                            text = text.replace(f"{{{{{i+1}}}}}", str(val))
                    elif components:
                        # Caso 2: components enviados directamente desde el dashboard
                        # Extraer los valores de los body parameters
                        for c in components:
                            if c.get("type") == "body":
                                for i, param in enumerate(c.get("parameters", [])):
                                    val = param.get("text", "-")
                                    text = text.replace(f"{{{{{i+1}}}}}", str(val))
                    template_content = text
                    break
            break
            
    # Fallback si por alguna razón el contenido está vacío
    if not template_content:
        template_content = f"[Template: {template_name}]"
    
    result = whatsapp_api.send_template_message(to_phone, template_name, language, components)
    
    if result.get("success"):
        wa_id = result.get("message_id")
        if wa_id:
            # Verificar si ya existe (creado por save_status en event_handlers.py por ejemplo)
            existing = Message.query.filter_by(wa_message_id=wa_id).first()
            if existing:
                existing.content = template_content
                existing.message_type = "template"
                existing.phone_number = to_phone
                logger.info(f"✅ Mensaje existente actualizado con contenido del template: {wa_id}")
            else:
                new_msg = Message(
                    wa_message_id=wa_id,
                    phone_number=to_phone,
                    direction="outbound",
                    message_type="template",
                    content=template_content,
                    timestamp=datetime.utcnow()
                )
                db.session.add(new_msg)
            
            try:
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                logger.error(f"Error al guardar mensaje en BD: {e}")
    
    return jsonify(result)

@app.route("/api/whatsapp/templates", methods=["GET"])
def api_list_templates():
    """API para obtener templates aprobados (usado para carga async en dashboard)."""
    if not whatsapp_api.is_configured():
        return jsonify({"templates": [], "error": "WhatsApp API no configurada"})
    
    templates_result = whatsapp_api.get_templates()
    approved = [t for t in templates_result.get("templates", []) if t.get("status") == "APPROVED"]
    return jsonify({"templates": approved})

@app.route("/api/whatsapp/send-text", methods=["POST"])
def api_send_text():
    """API para enviar mensaje de texto."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400

    to_phone = re.sub(r'[^\d]', '', str(data.get("to") or ""))
    text = data.get("text")

    if not to_phone or not text:
        return jsonify({"error": "to y text son requeridos"}), 400

    current_user = getattr(g, 'current_user', None)
    if current_user and not user_can_access_phone(current_user, to_phone):
        return jsonify({"error": "Sin acceso"}), 403

    # Determinar autor: si hay sesión es un agente, si no es el bot
    sent_by = session.get('username', 'bot')

    result = whatsapp_api.send_text_message(to_phone, text)

    if result.get("success"):
        wa_id = result.get("message_id")
        if wa_id:
            existing = Message.query.filter_by(wa_message_id=wa_id).first()
            if existing:
                existing.content = text
                existing.phone_number = to_phone
                existing.sent_by = sent_by
                logger.info(f"✅ Mensaje existente actualizado con texto: {wa_id}")
            else:
                new_msg = Message(
                    wa_message_id=wa_id,
                    phone_number=to_phone,
                    direction="outbound",
                    message_type="text",
                    content=text,
                    sent_by=sent_by,
                    timestamp=datetime.utcnow()
                )
                db.session.add(new_msg)
            
            try:
                db.session.commit()
            except Exception as e:
                db.session.rollback()
                logger.error(f"Error al guardar mensaje en BD: {e}")
    
    return jsonify(result)

@app.route("/api/whatsapp/send-media", methods=["POST"])
def api_send_media():
    """API para enviar imagen/documento/video/audio desde el dashboard."""
    to_phone = request.form.get("to")
    caption = request.form.get("caption", "").strip() or None
    file = request.files.get("file")

    if not to_phone:
        return jsonify({"error": "'to' es requerido"}), 400
    if not file:
        return jsonify({"error": "No se envió ningún archivo"}), 400
    _cu = getattr(g, 'current_user', None)
    if _cu and not user_can_access_phone(_cu, re.sub(r'[^\d]', '', to_phone)):
        return jsonify({"error": "Sin acceso"}), 403
    
    # Leer archivo
    file_bytes = file.read()
    if len(file_bytes) == 0:
        return jsonify({"error": "El archivo está vacío"}), 400
    
    # Límite de 16MB (WhatsApp limit para imágenes es 5MB, documentos 100MB, pero mantenemos razonable)
    if len(file_bytes) > 16 * 1024 * 1024:
        return jsonify({"error": "El archivo excede el límite de 16MB"}), 400
    
    original_filename = file.filename or "archivo"
    mime_type = file.content_type or mimetypes.guess_type(original_filename)[0] or "application/octet-stream"
    
    # Convertir audio no soportado por WhatsApp a MP3 usando PyAV (sin dependencias externas)
    WHATSAPP_AUDIO_OK = {'audio/aac', 'audio/mp4', 'audio/mpeg', 'audio/amr', 'audio/ogg', 'audio/opus'}
    base_mime = mime_type.split(';')[0].strip().lower()
    if base_mime.startswith("audio/") and base_mime not in WHATSAPP_AUDIO_OK:
        try:
            import av, io
            input_buf = io.BytesIO(file_bytes)
            output_buf = io.BytesIO()
            with av.open(input_buf) as in_container:
                in_stream = in_container.streams.audio[0]
                with av.open(output_buf, 'w', format='mp3') as out_container:
                    out_stream = out_container.add_stream('libmp3lame', rate=44100)
                    out_stream.layout = 'mono'
                    for frame in in_container.decode(in_stream):
                        frame.pts = None
                        for packet in out_stream.encode(frame):
                            out_container.mux(packet)
                    for packet in out_stream.encode(None):
                        out_container.mux(packet)
            file_bytes = output_buf.getvalue()
            mime_type = 'audio/mpeg'
            original_filename = original_filename.rsplit('.', 1)[0] + '.mp3'
            logger.info(f"✅ Audio convertido a mp3 con PyAV ({len(file_bytes)} bytes)")
        except Exception as e:
            logger.warning(f"Error convirtiendo audio con PyAV: {e}")

    # Determinar media_type para WhatsApp
    if mime_type.startswith("image/"):
        media_type = "image"
    elif mime_type.startswith("video/"):
        media_type = "video"
    elif mime_type.startswith("audio/"):
        media_type = "audio"
    else:
        media_type = "document"

    # 1. Subir a WhatsApp API
    upload_result = whatsapp_api.upload_media(file_bytes, mime_type, original_filename)
    if not upload_result.get("success"):
        return jsonify({"error": "Error subiendo archivo a WhatsApp: " + upload_result.get("error", "Desconocido")}), 500
    
    wa_media_id = upload_result["media_id"]
    
    # 2. Enviar mensaje multimedia
    send_result = whatsapp_api.send_media_message(
        to_phone, media_type, wa_media_id, 
        caption=caption,
        filename=original_filename if media_type == "document" else None
    )
    
    if not send_result.get("success"):
        return jsonify({"error": "Error enviando mensaje: " + send_result.get("error", "Desconocido")}), 500
    
    # 3. Subir a MinIO para almacenamiento persistente
    ext = mimetypes.guess_extension(mime_type) or ""
    if ext in ['.oga', '.opus']:
        ext = '.ogg'
    minio_filename = f"sent_{uuid.uuid4().hex[:12]}{ext}"
    media_url = whatsapp_api.upload_to_minio(file_bytes, mime_type, minio_filename)
    
    # 4. Guardar en BD
    wa_id = send_result.get("message_id")
    if wa_id:
        existing = Message.query.filter_by(wa_message_id=wa_id).first()
        if existing:
            existing.content = caption or f"[{media_type}]"
            existing.phone_number = to_phone
            existing.message_type = media_type
            existing.media_url = media_url
            existing.caption = caption
        else:
            new_msg = Message(
                wa_message_id=wa_id,
                phone_number=to_phone,
                direction="outbound",
                message_type=media_type,
                content=caption or f"[{media_type}]",
                media_url=media_url,
                caption=caption,
                timestamp=datetime.utcnow()
            )
            db.session.add(new_msg)
        
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error al guardar mensaje media en BD: {e}")
    
    return jsonify({
        "success": True,
        "message_id": wa_id,
        "media_type": media_type,
        "media_url": media_url
    })

# ==================== CAMPAÑAS ====================

@app.route("/campaigns")
def campaigns_page():
    """Página de campañas — OPTIMIZADO: stats en SQL + templates cargados async."""
    from sqlalchemy import case

    # Filtro de visibilidad por usuario
    vis_tag_ids = get_visible_tag_ids(g.current_user)

    # Stats en SQL — GROUP BY solo sobre Campaign.id (joinedload rompe GROUP BY en Postgres)
    campaigns_q = db.session.query(
        Campaign,
        func.count(CampaignLog.id).label('total'),
        func.count(case(
            (CampaignLog.status.in_(['sent', 'delivered', 'read']), 1)
        )).label('sent'),
        func.count(case(
            (CampaignLog.status == 'failed', 1)
        )).label('failed')
    ).outerjoin(
        CampaignLog, Campaign.id == CampaignLog.campaign_id
    ).group_by(Campaign.id)

    if vis_tag_ids is not None:
        if vis_tag_ids:
            campaigns_q = campaigns_q.filter(Campaign.tag_id.in_(vis_tag_ids))
        else:
            campaigns_q = campaigns_q.filter(False)

    campaigns_with_stats = campaigns_q.order_by(Campaign.created_at.desc()).all()

    # Cargar tags en batch (una sola query IN) para evitar N+1 lazy load en el template
    tag_ids_used = list({c.tag_id for c, *_ in campaigns_with_stats if c.tag_id})
    tags_by_id = {}
    if tag_ids_used:
        for t in Tag.query.filter(Tag.id.in_(tag_ids_used)).all():
            tags_by_id[t.id] = t
    for c, *_ in campaigns_with_stats:
        if c.tag_id and c.tag_id in tags_by_id:
            c.__dict__['tag'] = tags_by_id[c.tag_id]

    campaigns_data = []
    for c, total, sent, failed in campaigns_with_stats:
        campaigns_data.append({
            'campaign': c,
            'stats': {'total': total, 'sent': sent, 'failed': failed}
        })

    # Etiquetas visibles para el dropdown de "crear campaña"
    tags_q = Tag.query.filter_by(is_active=True)
    if vis_tag_ids is not None:
        tags_q = tags_q.filter(Tag.id.in_(vis_tag_ids)) if vis_tag_ids else tags_q.filter(False)
    tags = tags_q.all()

    # Los templates se cargan async via AJAX para no bloquear el SSR
    # (la llamada HTTP a Meta puede tardar 500ms-2s cuando el cache está frío)
    return render_template('campaigns.html',
                         campaigns=campaigns_data,
                         tags=tags)

@app.route("/campaigns/<int:campaign_id>")
def campaign_details_page(campaign_id):
    """Página de detalles de campaña."""
    campaign = Campaign.query.get_or_404(campaign_id)
    return render_template('campaign_details.html', campaign_id=campaign_id)

@app.route("/campaigns/compare")
def campaigns_compare_page():
    """Página de comparación de campañas."""
    return render_template('campaign_compare.html')

# ==========================================
# AUTO TAG RULES API
# ==========================================

@app.route("/api/auto-tag-rules", methods=["GET"])
def api_auto_tag_rules_list():
    tag_ids = get_visible_tag_ids(g.current_user)
    q = AutoTagRule.query
    if tag_ids is not None:
        q = q.filter(AutoTagRule.tag_id.in_(tag_ids)) if tag_ids else q.filter(False)
    rules = q.order_by(AutoTagRule.id).all()
    return jsonify([r.to_dict() for r in rules])

@app.route("/api/auto-tag-rules", methods=["POST"])
def api_auto_tag_rules_create():
    data = request.get_json() or {}
    tag_id = data.get('tag_id')
    prompt_condition = (data.get('prompt_condition') or '').strip()
    inactivity_minutes = data.get('inactivity_minutes', 30)

    if not tag_id or not prompt_condition:
        return jsonify({'error': 'tag_id y prompt_condition son requeridos'}), 400

    tag = Tag.query.get(tag_id)
    if not tag:
        return jsonify({'error': 'Tag no encontrado'}), 404

    rule = AutoTagRule(
        tag_id=tag_id,
        prompt_condition=prompt_condition,
        inactivity_minutes=inactivity_minutes,
        is_active=True
    )
    db.session.add(rule)
    db.session.commit()
    return jsonify(rule.to_dict()), 201

@app.route("/api/auto-tag-rules/<int:rule_id>", methods=["PUT"])
def api_auto_tag_rules_update(rule_id):
    rule = AutoTagRule.query.get_or_404(rule_id)
    data = request.get_json() or {}
    if 'tag_id' in data:
        rule.tag_id = data['tag_id']
    if 'prompt_condition' in data:
        rule.prompt_condition = data['prompt_condition'].strip()
    if 'inactivity_minutes' in data:
        rule.inactivity_minutes = data['inactivity_minutes']
    if 'is_active' in data:
        was_active = rule.is_active
        rule.is_active = bool(data['is_active'])
        # Si se está reactivando, actualizar activated_at para ignorar mensajes del período apagado
        if rule.is_active and not was_active:
            from datetime import datetime as _dt
            rule.activated_at = _dt.utcnow()
    db.session.commit()
    return jsonify(rule.to_dict())

@app.route("/api/auto-tag-rules/<int:rule_id>", methods=["DELETE"])
def api_auto_tag_rules_delete(rule_id):
    rule = AutoTagRule.query.get_or_404(rule_id)
    db.session.delete(rule)
    db.session.commit()
    return jsonify({'success': True})


# ==========================================
# FOLLOW-UP SEQUENCES API
# ==========================================

@app.route("/api/followup-sequences", methods=["GET"])
def api_followup_sequences_list():
    tag_ids = get_visible_tag_ids(g.current_user)
    q = FollowUpSequence.query
    if tag_ids is not None:
        q = q.filter(FollowUpSequence.tag_id.in_(tag_ids)) if tag_ids else q.filter(False)
    sequences = q.order_by(FollowUpSequence.id).all()
    return jsonify([s.to_dict() for s in sequences])

@app.route("/api/followup-sequences", methods=["POST"])
def api_followup_sequences_create():
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    tag_ids = data.get('tag_ids') or ([data.get('tag_id')] if data.get('tag_id') else [])
    steps_data = data.get('steps', [])

    if not name or not tag_ids:
        return jsonify({'error': 'name y al menos una etiqueta son requeridos'}), 400

    tags = Tag.query.filter(Tag.id.in_(tag_ids)).all()
    if not tags:
        return jsonify({'error': 'Etiquetas no encontradas'}), 404

    seq = FollowUpSequence(
        name=name,
        tag_id=tags[0].id,  # legacy
        is_active=True,
        add_tag_on_complete=bool(data.get('add_tag_on_complete', False)),
        send_window_start=data.get('send_window_start') or None,
        send_window_end=data.get('send_window_end') or None,
        send_weekdays=data.get('send_weekdays') or None
    )
    seq.trigger_tags = tags
    db.session.add(seq)
    db.session.flush()

    for i, step in enumerate(steps_data, start=1):
        s = FollowUpStep(
            sequence_id=seq.id,
            order=i,
            delay_hours=float(step.get('delay_hours', 24)),
            template_name=step.get('template_name', ''),
            template_language=step.get('template_language', 'es_AR'),
            template_params=step.get('template_params'),
            remove_tag_on_execute=bool(step.get('remove_tag_on_execute', False)),
            schedule_type=step.get('schedule_type', 'delay'),
            scheduled_weekday=step.get('scheduled_weekday'),
            scheduled_time=step.get('scheduled_time')
        )
        db.session.add(s)

    db.session.commit()
    return jsonify(seq.to_dict()), 201

@app.route("/api/followup-sequences/<int:seq_id>", methods=["PUT"])
def api_followup_sequences_update(seq_id):
    seq = FollowUpSequence.query.get_or_404(seq_id)
    data = request.get_json() or {}

    if 'name' in data:
        seq.name = data['name'].strip()
    if 'tag_ids' in data:
        tags = Tag.query.filter(Tag.id.in_(data['tag_ids'])).all()
        seq.trigger_tags = tags
        if tags:
            seq.tag_id = tags[0].id  # legacy
    elif 'tag_id' in data:
        seq.tag_id = data['tag_id']
    if 'is_active' in data:
        seq.is_active = bool(data['is_active'])
    if 'add_tag_on_complete' in data:
        seq.add_tag_on_complete = bool(data['add_tag_on_complete'])
    if 'send_window_start' in data:
        seq.send_window_start = data['send_window_start'] or None
    if 'send_window_end' in data:
        seq.send_window_end = data['send_window_end'] or None
    if 'send_weekdays' in data:
        seq.send_weekdays = data['send_weekdays'] or None

    if 'steps' in data:
        # Reemplazar pasos completamente
        FollowUpStep.query.filter_by(sequence_id=seq.id).delete()
        for i, step in enumerate(data['steps'], start=1):
            s = FollowUpStep(
                sequence_id=seq.id,
                order=i,
                delay_hours=float(step.get('delay_hours', 24)),
                template_name=step.get('template_name', ''),
                template_language=step.get('template_language', 'es_AR'),
                template_params=step.get('template_params'),
                remove_tag_on_execute=bool(step.get('remove_tag_on_execute', False)),
                schedule_type=step.get('schedule_type', 'delay'),
                scheduled_weekday=step.get('scheduled_weekday'),
                scheduled_time=step.get('scheduled_time')
            )
            db.session.add(s)

    db.session.commit()
    return jsonify(seq.to_dict())

@app.route("/api/followup-sequences/<int:seq_id>", methods=["DELETE"])
def api_followup_sequences_delete(seq_id):
    seq = FollowUpSequence.query.get_or_404(seq_id)
    db.session.delete(seq)
    db.session.commit()
    return jsonify({'success': True})

@app.route("/api/followup-sequences/<int:seq_id>/enrollments", methods=["GET"])
def api_followup_enrollments_list(seq_id):
    FollowUpSequence.query.get_or_404(seq_id)
    enrollments = FollowUpEnrollment.query.filter_by(sequence_id=seq_id)\
        .order_by(FollowUpEnrollment.enrolled_at.desc()).all()
    return jsonify([e.to_dict() for e in enrollments])

@app.route("/api/followup-enrollments", methods=["GET"])
def api_all_enrollments():
    """Todos los enrollments activos para el dashboard."""
    status_filter = request.args.get('status')
    q = FollowUpEnrollment.query
    if status_filter:
        q = q.filter_by(status=status_filter)
    enrollments = q.order_by(FollowUpEnrollment.enrolled_at.desc()).limit(200).all()
    return jsonify([e.to_dict() for e in enrollments])


@app.route("/api/followup-enrollments/<int:enrollment_id>/cancel", methods=["PATCH"])
def api_cancel_enrollment(enrollment_id):
    """Cancelar un enrollment activo (pendiente)."""
    enrollment = FollowUpEnrollment.query.get_or_404(enrollment_id)
    if enrollment.status != 'pending':
        return jsonify({'error': 'Solo se pueden cancelar enrollments pendientes'}), 400
    enrollment.status = 'cancelled'
    enrollment.cancelled_at = datetime.utcnow()
    db.session.commit()
    logger.info(f"🚫 Enrollment #{enrollment_id} cancelado (contacto {enrollment.contact.phone_number if enrollment.contact else '?'}, secuencia {enrollment.sequence.name if enrollment.sequence else '?'})")
    return jsonify({'success': True, 'enrollment': enrollment.to_dict()})


@app.route("/api/followup-sequences/<int:seq_id>/enroll", methods=["POST"])
def api_followup_enroll_manual(seq_id):
    """Enrollar manualmente un contacto en una secuencia (re-enrollment o testing)."""
    seq = FollowUpSequence.query.get_or_404(seq_id)
    data = request.get_json() or {}

    phone = (data.get('phone_number') or '').strip()
    if not phone:
        return jsonify({'error': 'phone_number requerido'}), 400

    phone_normalized = normalize_phone(phone)
    contact = find_contact_by_phone(phone_normalized)
    if not contact:
        return jsonify({'error': 'Contacto no encontrado'}), 404

    if not seq.steps:
        return jsonify({'error': 'La secuencia no tiene pasos'}), 400

    # Si ya existe un enrollment, cancelarlo y crear uno nuevo (re-enrollment)
    existing = FollowUpEnrollment.query.filter_by(
        contact_id=contact.id,
        sequence_id=seq.id
    ).first()
    if existing:
        existing.status = 'cancelled'
        existing.cancelled_at = datetime.utcnow()
        db.session.flush()
        # Eliminar para poder crear uno nuevo (unique constraint)
        db.session.delete(existing)
        db.session.flush()

    first_step = seq.steps[0]
    _now = datetime.utcnow()
    if (first_step.schedule_type or 'delay') == 'fixed_time' and first_step.scheduled_weekday is not None and first_step.scheduled_time:
        from followup_sender import _next_fixed_time
        _next_send = _next_fixed_time(_now, first_step.scheduled_weekday, first_step.scheduled_time)
    else:
        _next_send = _now + timedelta(hours=first_step.delay_hours)
    enrollment = FollowUpEnrollment(
        contact_id=contact.id,
        sequence_id=seq.id,
        current_step=1,
        status='pending',
        next_send_at=_next_send
    )
    db.session.add(enrollment)
    db.session.commit()
    logger.info(f"📋 [MANUAL] {contact.phone_number} enrollado en '{seq.name}' (paso 1)")
    return jsonify({'success': True, 'enrollment': enrollment.to_dict()})


@app.route("/api/followup-sequences/<int:seq_id>/enroll-tagged", methods=["POST"])
def api_followup_enroll_tagged(seq_id):
    """Enrolla masivamente todos los contactos que ya tienen las etiquetas disparadoras de la secuencia."""
    seq = FollowUpSequence.query.get_or_404(seq_id)
    if not seq.steps:
        return jsonify({'error': 'La secuencia no tiene pasos'}), 400

    trigger_tags = seq.get_trigger_tags()
    if not trigger_tags:
        return jsonify({'error': 'La secuencia no tiene etiquetas disparadoras configuradas'}), 400

    tag_ids = [t.id for t in trigger_tags]
    contacts = Contact.query.filter(Contact.tags.any(Tag.id.in_(tag_ids))).all()

    enrolled_count = 0
    skipped_count = 0
    _now = datetime.utcnow()
    first_step = seq.steps[0]

    for contact in contacts:
        existing = FollowUpEnrollment.query.filter_by(
            contact_id=contact.id,
            sequence_id=seq.id
        ).first()
        if existing:
            if existing.status == 'pending':
                skipped_count += 1
                continue
            db.session.delete(existing)
            db.session.flush()

        if (first_step.schedule_type or 'delay') == 'fixed_time' and first_step.scheduled_weekday is not None and first_step.scheduled_time:
            from followup_sender import _next_fixed_time
            next_send = _next_fixed_time(_now, first_step.scheduled_weekday, first_step.scheduled_time)
        else:
            next_send = _now + timedelta(hours=first_step.delay_hours)

        enrollment = FollowUpEnrollment(
            contact_id=contact.id,
            sequence_id=seq.id,
            current_step=1,
            status='pending',
            next_send_at=next_send
        )
        db.session.add(enrollment)
        enrolled_count += 1
        logger.info(f"📋 [MANUAL-BULK] {contact.phone_number} enrollado en '{seq.name}'")

    db.session.commit()
    return jsonify({
        'success': True,
        'enrolled': enrolled_count,
        'skipped': skipped_count,
        'tag_names': [t.name for t in trigger_tags]
    })


@app.route("/api/contacts/search", methods=["GET"])
def api_contacts_search():
    """Búsqueda rápida de contactos por nombre o teléfono (para modales)."""
    q = (request.args.get('q') or '').strip()
    if len(q) < 2:
        return jsonify([])
    results = Contact.query.filter(
        db.or_(
            Contact.phone_number.ilike(f'%{q}%'),
            Contact.name.ilike(f'%{q}%'),
            Contact.first_name.ilike(f'%{q}%'),
            Contact.last_name.ilike(f'%{q}%'),
        )
    ).limit(20).all()
    return jsonify([{'id': c.id, 'name': c.name or c.phone_number, 'phone_number': c.phone_number} for c in results])


@app.route("/api/campaigns", methods=["GET"])
def api_list_campaigns():
    """Lista campañas."""
    tag_ids = get_visible_tag_ids(g.current_user)
    q = Campaign.query
    if tag_ids is not None:
        q = q.filter(Campaign.tag_id.in_(tag_ids)) if tag_ids else q.filter(False)
    campaigns = q.order_by(Campaign.created_at.desc()).all()

    # Obtener templates para preview de mensajes
    templates_map = {}
    try:
        templates_data = whatsapp_api.get_templates()
        for t in templates_data.get("templates", []):
            # Extraer el primer componente BODY del template
            body_text = ""
            for comp in t.get("components", []):
                if comp.get("type") == "BODY":
                    body_text = comp.get("text", "")
                    break
            templates_map[t.get("name")] = body_text
    except Exception as e:
        logger.warning(f"Error obteniendo templates para preview: {e}")
        # Continuar sin previews de mensajes

    result = []
    for c in campaigns:
        total = len(c.logs)
        sent = sum(1 for l in c.logs if l.status in ('sent', 'delivered', 'read'))
        failed = sum(1 for l in c.logs if l.status == 'failed')

        # Obtener preview del mensaje
        message_preview = templates_map.get(c.template_name, "")
        # Eliminar las variables {{1}}, {{2}}, etc. para el preview
        message_preview = re.sub(r'\{\{\d+\}\}', '...', message_preview)

        result.append({
            'id': c.id,
            'name': c.name,
            'status': c.status,
            'tag': c.tag.name if c.tag else None,
            'template_name': c.template_name,
            'message_preview': message_preview[:100] + ('...' if len(message_preview) > 100 else ''),
            'total': total,
            'sent': sent,
            'failed': failed,
            'created_at': format_utc_iso(c.created_at),
            'started_at': format_utc_iso(c.started_at),
            'completed_at': format_utc_iso(c.completed_at),
            'created_by': c.created_by
        })
    return jsonify(result)

@app.route("/api/campaigns", methods=["POST"])
def api_create_campaign():
    """Crea una nueva campaña."""
    data = request.json
    name = data.get('name')
    tag_id = data.get('tag_id')
    template_name = data.get('template_name')
    template_language = data.get('template_language', 'es_AR')
    scheduled_at_str = data.get('scheduled_at')
    variables = data.get('variables') # Dict {"1": "first_name", ...}

    if not name or not template_name:
        return jsonify({'error': 'name y template_name requeridos'}), 400

    if not tag_id:
        return jsonify({'error': 'tag_id requerido'}), 400

    tag = Tag.query.get(tag_id)
    if not tag:
        return jsonify({'error': 'Tag no encontrado'}), 404
        
    scheduled_at = None
    status = 'draft'
    
    if scheduled_at_str:
        # Asumimos que viene en ISO format o timestamp
        try:
            # Si viene con timezone, convertir a UTC. Si no, asumir que es UTC o manejarlo.
            # Simplificación: el frontend debe enviar ISO string.
            dt = datetime.fromisoformat(scheduled_at_str.replace('Z', '+00:00'))
            
            # Si es naive (no tiene timezone), asumir que es hora de Argentina
            if dt.tzinfo is None:
                ar_tz = pytz.timezone('America/Argentina/Buenos_Aires')
                dt = ar_tz.localize(dt)
            
            # Convertir a UTC para almacenamiento
            dt = dt.astimezone(timezone.utc).replace(tzinfo=None)
            
            scheduled_at = dt
            status = 'scheduled'
        except Exception as e:
            return jsonify({'error': f'Error en fecha programada: {str(e)}'}), 400

    campaign = Campaign(
        name=name,
        template_name=template_name,
        template_language=template_language,
        tag_id=tag_id,
        status=status,
        scheduled_at=scheduled_at,
        variables=variables,
        created_by=session.get('username')
    )
    try:
        db.session.add(campaign)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating campaign: {e}")
        return jsonify({'error': str(e)}), 500

    return jsonify({
        'success': True,
        'id': campaign.id,
        'name': campaign.name,
        'status': campaign.status
    }), 201

@app.route("/api/campaigns/<int:campaign_id>", methods=["DELETE"])
def api_delete_campaign(campaign_id):
    """Elimina una campaña (solo si está en draft o scheduled)."""
    campaign = Campaign.query.get(campaign_id)
    if not campaign:
        return jsonify({'error': 'Campaña no encontrada'}), 404
    if campaign.status not in ('draft', 'scheduled'):
        return jsonify({'error': 'Solo se puede eliminar una campaña en estado draft o programada'}), 400

    CampaignLog.query.filter_by(campaign_id=campaign_id).delete()
    db.session.delete(campaign)
    db.session.commit()
    return jsonify({'success': True})

@app.route("/api/campaigns/<int:campaign_id>/send", methods=["POST"])
def api_send_campaign(campaign_id):
    """Inicia el envío de una campaña en background."""
    # Bloqueo de fila para prevenir race condition
    campaign = Campaign.query.with_for_update().get(campaign_id)
    if not campaign:
        return jsonify({'error': 'Campaña no encontrada'}), 404
    # Permitir enviar si es draft O si es scheduled (para "Iniciar Ahora")
    if campaign.status not in ('draft', 'scheduled'):
        return jsonify({'error': 'La campaña ya está en curso o completada'}), 400

    if not campaign.tag_id:
        return jsonify({'error': 'La campaña debe tener un tag asignado'}), 400

    # Contar contactos con el tag (rápido, sin cargar en memoria)
    contact_count = Contact.query.filter(
        Contact.tags.any(Tag.id == campaign.tag_id)
    ).count()

    if contact_count == 0:
        return jsonify({'error': 'No hay contactos con ese tag'}), 400

    # Actualizar estado
    campaign.status = 'sending'
    campaign.started_at = datetime.utcnow()
    db.session.commit()

    # Crear logs pendientes - SUPER OPTIMIZADO CON SQL DIRECTO + ON CONFLICT
    # Inserta todos los logs en una sola operación SQL sin cargar contactos en Python
    now = datetime.utcnow()
    try:
        # Usar ON CONFLICT DO NOTHING para evitar subconsulta NOT EXISTS (mucho más rápido)
        result = db.session.execute(text("""
            INSERT INTO whatsapp_campaign_logs (campaign_id, contact_id, contact_phone, status, created_at)
            SELECT :cid, c.id, c.phone_number, 'pending', :now
            FROM whatsapp_contacts c
            JOIN whatsapp_contact_tags ct ON c.id = ct.contact_id
            WHERE ct.tag_id = :tid
            ON CONFLICT (campaign_id, contact_id) DO NOTHING
        """), {'cid': campaign.id, 'tid': campaign.tag_id, 'now': now})
        db.session.commit()
        logger.info(f"📊 Logs creados para campaña {campaign.id}, tag {campaign.tag_id}, contactos: {contact_count}, insertados: {result.rowcount}")
    except Exception as e:
        db.session.rollback()
        logger.error(f"❌ Error creando logs con ON CONFLICT para campaña {campaign.id}: {e}")
        # Fallback: INSERT sin ON CONFLICT (para SQLite u otras BD)
        try:
            result = db.session.execute(text("""
                INSERT INTO whatsapp_campaign_logs (campaign_id, contact_id, contact_phone, status, created_at)
                SELECT :cid, c.id, c.phone_number, 'pending', :now
                FROM whatsapp_contacts c
                JOIN whatsapp_contact_tags ct ON c.id = ct.contact_id
                WHERE ct.tag_id = :tid
                AND NOT EXISTS (
                    SELECT 1 FROM whatsapp_campaign_logs cl 
                    WHERE cl.campaign_id = :cid AND cl.contact_id = c.id
                )
            """), {'cid': campaign.id, 'tid': campaign.tag_id, 'now': now})
            db.session.commit()
            logger.info(f"📊 Logs creados con fallback SQL para campaña {campaign.id}, insertados: {result.rowcount}")
        except Exception as e2:
            db.session.rollback()
            logger.error(f"❌ Error creando logs con fallback para campaña {campaign.id}: {e2}")

    ctx = app.app_context()
    t = threading.Thread(target=send_campaign_bg, args=(ctx, campaign.id))
    t.daemon = True
    t.start()

    return jsonify({
        'success': True,
        'status': 'sending',
        'total_contacts': contact_count
    })

def send_campaign_bg(app_context, cid):
    """Función de envío en background con procesamiento por lotes."""
    with app_context:
        camp = Campaign.query.get(cid)
        if not camp: return
        
        BATCH_SIZE = 100  # Procesar en lotes para no saturar memoria
        total_sent = 0
        total_failed = 0
        
        # Pre-cargar el texto del template para usarlo en el historial de cada contacto
        template_body_text = None
        try:
            templates_result = whatsapp_api.get_templates()
            for t in templates_result.get("templates", []):
                if t.get("name") == camp.template_name and t.get("language") == camp.template_language:
                    for comp in t.get("components", []):
                        if comp.get("type") == "BODY":
                            template_body_text = comp.get("text", "")
                            break
                    break
        except Exception as e:
            logger.warning(f"No se pudo obtener texto del template para campaña {cid}: {e}")

        while True:
            # Cargar solo un lote de logs pendientes a la vez
            logs = CampaignLog.query.filter_by(
                campaign_id=cid, 
                status='pending'
            ).limit(BATCH_SIZE).all()
            
            if not logs:
                break  # No hay más logs pendientes
            
            for log in logs:
                try:
                    # Construir componentes con variables dinámicas
                    components = None
                    if camp.variables:
                        # Usar contact_id para obtener el contacto (o la relación directa)
                        contact = log.contact or Contact.query.get(log.contact_id)

                        # Separar variables por componente (body vs header)
                        body_vars = {}
                        header_vars = {}
                        
                        for key, field in camp.variables.items():
                            if '-' in key:
                                # New format: "body-1", "header-1"
                                comp, idx = key.split('-', 1)
                                if comp == 'header':
                                    header_vars[int(idx)] = field
                                else:
                                    body_vars[int(idx)] = field
                            else:
                                # Old format: "1", "2" (backwards compat, treat as body)
                                body_vars[int(key)] = field
                        
                        components = []
                        
                        # Build header parameters
                        if header_vars:
                            header_params = []
                            for idx in sorted(header_vars.keys()):
                                field = header_vars[idx]
                                value = "-"
                                if field == 'phone_number':
                                    value = contact.phone_number if contact else log.contact_phone
                                elif contact:
                                    val = getattr(contact, field, None)
                                    if val:
                                        value = str(val)
                                header_params.append({"type": "text", "text": value})
                            components.append({"type": "header", "parameters": header_params})
                        
                        # Build body parameters
                        if body_vars:
                            body_params = []
                            for idx in sorted(body_vars.keys()):
                                field = body_vars[idx]
                                value = "-"
                                if field == 'phone_number':
                                    value = contact.phone_number if contact else log.contact_phone
                                elif contact:
                                    val = getattr(contact, field, None)
                                    if val:
                                        value = str(val)
                                body_params.append({"type": "text", "text": value})
                            components.append({"type": "body", "parameters": body_params})
                        
                        if not components:
                            components = None

                    result = whatsapp_api.send_template_message(
                        log.contact_phone,
                        camp.template_name,
                        camp.template_language,
                        components=components
                    )
                    
                    if result.get('success'):
                        log.status = 'sent'
                        log.message_id = result.get('message_id')
                        total_sent += 1
                        wa_id = result.get('message_id')
                        if wa_id:
                            # Resolver el contenido real del template con las variables del contacto
                            content_preview = template_body_text or f'[Template: {camp.template_name}]'
                            if template_body_text and camp.variables:
                                # Reemplazar {{1}}, {{2}}, etc. con los valores reales
                                resolved_text = template_body_text
                                # Recoger body_vars del contacto actual
                                body_vars_local = {}
                                for key, field in camp.variables.items():
                                    if '-' in key:
                                        comp_name, idx = key.split('-', 1)
                                        if comp_name == 'body':
                                            body_vars_local[int(idx)] = field
                                    else:
                                        body_vars_local[int(key)] = field
                                
                                contact_for_preview = log.contact or Contact.query.get(log.contact_id)
                                for idx, field in body_vars_local.items():
                                    value = "-"
                                    if field == 'phone_number':
                                        value = contact_for_preview.phone_number if contact_for_preview else log.contact_phone
                                    elif contact_for_preview:
                                        val = getattr(contact_for_preview, field, None)
                                        if val:
                                            value = str(val)
                                    resolved_text = resolved_text.replace(f'{{{{{idx}}}}}', value)
                                content_preview = resolved_text
                            
                            new_msg = Message(
                                wa_message_id=wa_id,
                                phone_number=log.contact_phone,
                                direction='outbound',
                                message_type='template',
                                content=content_preview,
                                timestamp=datetime.utcnow()
                            )
                            db.session.add(new_msg)
                    else:
                        log.status = 'failed'
                        log.error_detail = str(result.get('error') or result)
                        total_failed += 1
                except Exception as e:
                    log.status = 'failed'
                    log.error_detail = str(e)
                    total_failed += 1

                db.session.commit()
                time_module.sleep(1)  # Rate limiting
            
            # Log de progreso cada lote
            logger.info(f"📊 Campaña {cid}: Lote procesado. Enviados: {total_sent}, Fallidos: {total_failed}")

        camp.status = 'completed'
        camp.completed_at = datetime.utcnow()
        db.session.commit()
        logger.info(f"✅ Campaña {cid} completada. Total enviados: {total_sent}, fallidos: {total_failed}")

def run_scheduler():
    """Scheduler para verificar campañas programadas, categorizar conversaciones, auto-taggear y enviar follow-ups."""
    from conversation_categorizer import run_categorization
    from auto_tagger import run_auto_tagger
    from followup_sender import run_followup_sender
    categorize_counter = 0

    while True:
        try:
            with app.app_context():
                now = datetime.utcnow()
                # Buscar campañas programadas que ya deberían salir
                # skip_locked=True evita que el scheduler intente procesar algo que ya está bloqueado por el usuario
                pending = Campaign.query.filter(
                    Campaign.status == 'scheduled',
                    Campaign.scheduled_at <= now
                ).with_for_update(skip_locked=True).all()
                
                for camp in pending:
                    logger.info(f"🚀 Ejecutando campaña programada: {camp.name}")
                    
                    # Contar contactos con el tag
                    contact_count = Contact.query.filter(
                        Contact.tags.any(Tag.id == camp.tag_id)
                    ).count()
                    
                    if contact_count == 0:
                        camp.status = 'failed'
                        camp.completed_at = now
                        logger.warning(f"Campaña {camp.name} fallida: Sin contactos")
                        db.session.commit()
                        continue
                        
                    # Pasar a sending
                    camp.status = 'sending'
                    camp.started_at = now
                    db.session.commit()
                    
                    # Crear logs - SUPER OPTIMIZADO CON SQL DIRECTO + ON CONFLICT
                    try:
                        db.session.execute(text("""
                            INSERT INTO whatsapp_campaign_logs (campaign_id, contact_id, contact_phone, status, created_at)
                            SELECT :cid, c.id, c.phone_number, 'pending', :now
                            FROM whatsapp_contacts c
                            JOIN whatsapp_contact_tags ct ON c.id = ct.contact_id
                            WHERE ct.tag_id = :tid
                            ON CONFLICT (campaign_id, contact_id) DO NOTHING
                        """), {'cid': camp.id, 'tid': camp.tag_id, 'now': now})
                        db.session.commit()
                    except Exception as e:
                        db.session.rollback()
                        logger.error(f"Error con ON CONFLICT en scheduler: {e}")
                        # Fallback sin ON CONFLICT
                        db.session.execute(text("""
                            INSERT INTO whatsapp_campaign_logs (campaign_id, contact_id, contact_phone, status, created_at)
                            SELECT :cid, c.id, c.phone_number, 'pending', :now
                            FROM whatsapp_contacts c
                            JOIN whatsapp_contact_tags ct ON c.id = ct.contact_id
                            WHERE ct.tag_id = :tid
                            AND NOT EXISTS (
                                SELECT 1 FROM whatsapp_campaign_logs cl 
                                WHERE cl.campaign_id = :cid AND cl.contact_id = c.id
                            )
                        """), {'cid': camp.id, 'tid': camp.tag_id, 'now': now})
                        db.session.commit()
                    
                    # Lanzar thread de envío
                    t = threading.Thread(target=send_campaign_bg, args=(app.app_context(), camp.id))
                    t.daemon = True
                    t.start()
                    
        except Exception as e:
            logger.error(f"Error en scheduler: {e}")
        
        # Run conversation categorization every 5 minutes (every 5 loops)
        categorize_counter += 1
        if categorize_counter >= 5:
            categorize_counter = 0
            try:
                run_categorization(app.app_context())
            except Exception as e:
                logger.error(f"Error en categorización: {e}")
            try:
                run_auto_tagger(app.app_context())
            except Exception as e:
                logger.error(f"Error en auto tagger: {e}")

        # Follow-up sender: corre cada minuto
        try:
            run_followup_sender(app.app_context())
        except Exception as e:
            logger.error(f"Error en follow-up sender: {e}")

        time_module.sleep(60) # Revisar cada minuto

# Iniciar scheduler
scheduler_thread = threading.Thread(target=run_scheduler)
scheduler_thread.daemon = True
scheduler_thread.start()



@app.route("/api/campaigns/<int:campaign_id>/status", methods=["GET"])
def api_campaign_status(campaign_id):
    """Obtiene el estado en tiempo real de una campaña."""
    campaign = Campaign.query.get(campaign_id)
    if not campaign:
        return jsonify({'error': 'Campaña no encontrada'}), 404

    logs = CampaignLog.query.filter_by(campaign_id=campaign_id).all()
    total = len(logs)
    sent = sum(1 for l in logs if l.status in ('sent', 'delivered', 'read'))
    failed = sum(1 for l in logs if l.status == 'failed')
    pending = sum(1 for l in logs if l.status == 'pending')

    return jsonify({
        'id': campaign.id,
        'status': campaign.status,
        'total': total,
        'sent': sent,
        'failed': failed,
        'pending': pending,
        'started_at': format_utc_iso(campaign.started_at),
        'completed_at': format_utc_iso(campaign.completed_at)
    })

@app.route("/api/campaigns/<int:campaign_id>/stats_preview", methods=["GET"]) # Renamed to avoid conflict
def api_get_campaign_stats_preview(campaign_id):
    """Obtiene detalles completos de una campaña (Preview)."""
    campaign = Campaign.query.get(campaign_id)
    if not campaign:
        return jsonify({'error': 'Campaña no encontrada'}), 404

    logs = CampaignLog.query.filter_by(campaign_id=campaign_id).all()
    total = len(logs)
    sent = sum(1 for l in logs if l.status in ('sent', 'delivered', 'read'))
    read = sum(1 for l in logs if l.status == 'read')
    failed = sum(1 for l in logs if l.status == 'failed')
    
    # Preview de logs (últimos 50)
    preview_logs = []
    for l in logs[-50:]:
        contact = l.contact or Contact.query.get(l.contact_id)
        preview_logs.append({
            'phone': l.contact_phone,
            'name': contact.name if contact else '',
            'status': l.status,
            'error': l.error_detail
        })

    return jsonify({
        'id': campaign.id,
        'name': campaign.name,
        'status': campaign.status,
        'template_name': campaign.template_name,
        'tag_name': campaign.tag.name if campaign.tag else '??',
        'created_at': format_utc_iso(campaign.created_at),
        'scheduled_at': format_utc_iso(campaign.scheduled_at),
        'started_at': format_utc_iso(campaign.started_at),
        'completed_at': format_utc_iso(campaign.completed_at),
        'stats': {
            'total': total,
            'sent': sent,
            'read': read,
            'failed': failed
        },
        'logs_preview': preview_logs,
        'variables': campaign.variables
    })

@app.route("/api/campaigns/<int:campaign_id>")
def api_campaign_details(campaign_id):
    """API para obtener detalles y estadísticas de una campaña."""
    try:
        campaign = Campaign.query.get_or_404(campaign_id)

        # Estadísticas agregadas
        total_logs = CampaignLog.query.filter_by(campaign_id=campaign_id).count()

        # Contar por estados
        logs_stats = db.session.query(
            CampaignLog.status, func.count(CampaignLog.id)
        ).filter(
            CampaignLog.campaign_id == campaign_id
        ).group_by(CampaignLog.status).all()

        stats_map = {s: c for s, c in logs_stats}
        sent_count = stats_map.get('sent', 0)
        delivered_count = stats_map.get('delivered', 0)
        read_count = stats_map.get('read', 0)
        failed_count = stats_map.get('failed', 0)

        # 'Enviados' para la UI incluye todo lo que salió exitosamente (sent, delivered, read)
        total_successful = sent_count + delivered_count + read_count

        # ========== CALCULAR RESPUESTAS A LA CAMPAÑA ==========
        # Respuestas = mensajes inbound recibidos después del envío de la campaña
        # Ventana de tiempo: 48 horas después del envío del mensaje de campaña
        response_window_hours = 48

        # Obtener todos los logs exitosos con sus teléfonos y timestamps
        successful_logs = db.session.query(
            CampaignLog.contact_phone,
            CampaignLog.created_at
        ).filter(
            CampaignLog.campaign_id == campaign_id,
            CampaignLog.status.in_(['sent', 'delivered', 'read']),
            CampaignLog.created_at.isnot(None)
        ).all()

        # Para cada contacto, buscar si respondió dentro de la ventana de tiempo
        responded_phones = set()
        total_responses = 0

        for phone, sent_time in successful_logs:
            if not sent_time:
                continue

            # Buscar mensajes inbound de este teléfono después del envío
            response_cutoff = sent_time + timedelta(hours=response_window_hours)

            response_count = Message.query.filter(
                Message.phone_number == phone,
                Message.direction == 'inbound',
                Message.timestamp > sent_time,
                Message.timestamp <= response_cutoff
            ).count()

            if response_count > 0:
                responded_phones.add(phone)
                total_responses += response_count

        unique_responders = len(responded_phones)
        response_rate = round((unique_responders / total_successful * 100) if total_successful > 0 else 0, 1)

        # Logs preview (últimos 50)
        # Logs preview con PAGINACIÓN
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)

        pagination = db.session.query(
            CampaignLog, Contact.name
        ).outerjoin(
            Contact, CampaignLog.contact_id == Contact.id
        ).filter(
            CampaignLog.campaign_id == campaign_id
        ).order_by(CampaignLog.created_at.desc()).paginate(page=page, per_page=per_page, error_out=False)

        logs_preview = []
        for log, contact_name in pagination.items:
            logs_preview.append({
                'phone': log.contact_phone,
                'name': contact_name,
                'status': log.status,
                'error': log.error_detail,
                'created_at': format_utc_iso(log.created_at)
            })

        # Obtener detalles del template desde WhatsApp API
        template_content = None
        try:
            templates_data = whatsapp_api.get_templates()
            if templates_data and 'templates' in templates_data:
                # Buscar el template por nombre y lenguaje
                matching_template = next(
                    (t for t in templates_data['templates']
                     if t.get('name') == campaign.template_name and
                        t.get('language') == campaign.template_language),
                    None
                )
                if matching_template:
                    template_content = {
                        'name': matching_template.get('name'),
                        'language': matching_template.get('language'),
                        'status': matching_template.get('status'),
                        'components': matching_template.get('components', [])
                    }
        except Exception as e:
            logger.warning(f"No se pudo obtener contenido del template: {e}")

        # Calcular métricas adicionales
        # 1. Respuestas promedio por contacto (engagement depth)
        avg_responses_per_contact = round(total_responses / unique_responders, 2) if unique_responders > 0 else 0

        # 2. Score de efectividad (0-10)
        # Fórmula: (Tasa de entrega × 0.3) + (Tasa de lectura × 0.4) + (Tasa de respuesta × 0.3)
        delivery_rate = (total_successful / total_logs * 100) if total_logs > 0 else 0
        read_rate = (read_count / total_successful * 100) if total_successful > 0 else 0
        # response_rate ya está calculado arriba

        effectiveness_score = round(
            (delivery_rate / 100 * 3) +  # Max 3 puntos
            (read_rate / 100 * 4) +       # Max 4 puntos
            (response_rate / 100 * 3),    # Max 3 puntos
            1
        )

        return jsonify({
            'id': campaign.id,
            'name': campaign.name,
            'status': campaign.status,
            'template_name': campaign.template_name,
            'template_content': template_content,
            'tag_name': campaign.tag.name if campaign.tag else 'N/A',
            'created_at': format_utc_iso(campaign.created_at),
            'started_at': format_utc_iso(campaign.started_at),
            'completed_at': format_utc_iso(campaign.completed_at),
            'scheduled_at': format_utc_iso(campaign.scheduled_at),
            'stats': {
                'total': total_logs,
                'sent': total_successful,
                'read': read_count,
                'failed': failed_count,
                'unique_responders': unique_responders,
                'total_responses': total_responses,
                'response_rate': response_rate,
                'avg_responses_per_contact': avg_responses_per_contact,
                'effectiveness_score': effectiveness_score,
                'delivery_rate': round(delivery_rate, 1),
                'read_rate': round(read_rate, 1)
            },
            'logs_preview': logs_preview,
            'pagination': {
                'page': pagination.page,
                'per_page': pagination.per_page,
                'total_pages': pagination.pages,
                'total_items': pagination.total,
                'has_next': pagination.has_next,
                'has_prev': pagination.has_prev
            }
        })

    except Exception as e:
        logger.error(f"Error getting campaign details: {e}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/campaigns/<int:campaign_id>/export", methods=["GET"])
def api_export_campaign_stats(campaign_id):
    """Exporta reporte de campaña a Excel."""
    try:
        campaign = Campaign.query.get(campaign_id)
        if not campaign:
            return jsonify({'error': 'Campaña no encontrada'}), 404
            
        # Optimization: Fetch all data in a single query
        results = db.session.query(
            CampaignLog.contact_phone,
            CampaignLog.status,
            CampaignLog.error_detail,
            CampaignLog.message_id,
            Contact.contact_id,
            Contact.name,
            Contact.first_name,
            Contact.last_name
        ).outerjoin(Contact, CampaignLog.contact_id == Contact.id)\
         .filter(CampaignLog.campaign_id == campaign_id).all()
        
        data = []
        for row in results:
            data.append({
                'Telefono': row.contact_phone,
                'ID Cliente': row.contact_id or '',
                'Nombre Completo': row.name or '',
                'Nombre': row.first_name or '',
                'Apellido': row.last_name or '',
                'Estado Mensaje': row.status,
                'Error': row.error_detail or '',
                'Mensaje ID': row.message_id or ''
            })
            
        df = pd.DataFrame(data)
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Reporte')
            
        output.seek(0)
        
        filename = f"reporte_{campaign.name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        # Limpiar nombre de archivo
        filename = "".join([c for c in filename if c.isalnum() or c in (' ', '.', '_')]).strip().replace(' ', '_')
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Error exportando campaña: {e}")
        return jsonify({'error': str(e)}), 500

# ==================== CONVERSATION TOPICS ====================

@app.route("/topics")
def topics_page():
    """Página para gestionar temas de conversación."""
    return render_template('topics.html')


@app.route("/api/conversation-topics", methods=["GET"])
def api_list_conversation_topics():
    """Lista todos los temas de conversación."""
    topics = ConversationTopic.query.order_by(ConversationTopic.name).all()
    return jsonify([t.to_dict() for t in topics])


@app.route("/api/conversation-topics", methods=["POST"])
def api_create_conversation_topic():
    """Crea un nuevo tema de conversación."""
    data = request.get_json()
    
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'El nombre es requerido'}), 400
    
    # Verificar si ya existe
    existing = ConversationTopic.query.filter_by(name=name).first()
    if existing:
        return jsonify({'success': False, 'error': 'Ya existe un tema con ese nombre'}), 400
    
    topic = ConversationTopic(
        name=name,
        description=data.get('description', '').strip() or None,
        keywords=data.get('keywords', []),
        color=data.get('color', 'blue')
    )
    
    db.session.add(topic)
    db.session.commit()
    
    return jsonify({'success': True, 'topic': topic.to_dict()})


@app.route("/api/conversation-topics/<int:topic_id>", methods=["PUT"])
def api_update_conversation_topic(topic_id):
    """Actualiza un tema de conversación."""
    topic = ConversationTopic.query.get_or_404(topic_id)
    data = request.get_json()
    
    if 'name' in data:
        new_name = data['name'].strip()
        if new_name and new_name != topic.name:
            # Verificar duplicados
            existing = ConversationTopic.query.filter_by(name=new_name).first()
            if existing:
                return jsonify({'success': False, 'error': 'Ya existe un tema con ese nombre'}), 400
            topic.name = new_name
    
    if 'description' in data:
        topic.description = data['description'].strip() or None
    
    if 'keywords' in data:
        topic.keywords = data['keywords']
    
    if 'color' in data:
        topic.color = data['color']
    
    db.session.commit()
    return jsonify({'success': True, 'topic': topic.to_dict()})


@app.route("/api/conversation-topics/<int:topic_id>", methods=["DELETE"])
def api_delete_conversation_topic(topic_id):
    """Elimina un tema de conversación."""
    topic = ConversationTopic.query.get_or_404(topic_id)
    
    # Desasociar sesiones pero no eliminarlas
    ConversationSession.query.filter_by(topic_id=topic_id).update({'topic_id': None})
    
    db.session.delete(topic)
    db.session.commit()
    
    return jsonify({'success': True})


# ==================== CONVERSATION NOTES ====================

@app.route("/api/conversations/<phone>/notes", methods=["GET"])
def api_get_notes(phone):
    """Lista notas internas de una conversación."""
    notes = ConversationNote.query.filter_by(
        phone_number=phone
    ).order_by(ConversationNote.created_at.desc()).all()
    return jsonify({'notes': [n.to_dict() for n in notes]})


@app.route("/api/conversations/<phone>/notes", methods=["POST"])
def api_create_note(phone):
    """Crea una nota interna para una conversación."""
    data = request.get_json()
    content = (data.get('content') or '').strip()
    if not content:
        return jsonify({'error': 'El contenido es requerido'}), 400

    note = ConversationNote(
        phone_number=phone,
        content=content
    )
    db.session.add(note)
    db.session.commit()
    return jsonify({'success': True, 'note': note.to_dict()})


@app.route("/api/conversations/notes/<int:note_id>", methods=["DELETE"])
def api_delete_note(note_id):
    """Elimina una nota interna."""
    note = ConversationNote.query.get_or_404(note_id)
    db.session.delete(note)
    db.session.commit()
    return jsonify({'success': True})


# ==================== EXPORT CONVERSATION PDF ====================

@app.route("/api/conversations/<phone>/export-pdf")
def api_export_conversation_pdf(phone):
    """Exporta una conversación completa a PDF."""
    from fpdf import FPDF

    # Obtener contacto
    contact = Contact.query.filter_by(phone_number=phone).first()
    contact_name = contact.name if contact else phone

    # Obtener mensajes
    messages = Message.query.filter_by(
        phone_number=phone
    ).order_by(Message.timestamp.asc()).all()

    if not messages:
        return jsonify({'error': 'No hay mensajes para exportar'}), 404

    # Obtener notas internas
    notes = ConversationNote.query.filter_by(
        phone_number=phone
    ).order_by(ConversationNote.created_at.desc()).all()

    # Zona horaria
    tz = pytz.timezone('America/Argentina/Buenos_Aires')

    # Crear PDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # --- Encabezado ---
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 10, 'Conversacion de WhatsApp', ln=True, align='C')
    pdf.set_font('Helvetica', '', 11)
    pdf.cell(0, 7, f'Contacto: {_safe_text(contact_name)}', ln=True)
    pdf.cell(0, 7, f'Telefono: {phone}', ln=True)

    if contact and contact.contact_id:
        pdf.cell(0, 7, f'Client ID: {contact.contact_id}', ln=True)

    now_ar = datetime.utcnow().replace(tzinfo=pytz.utc).astimezone(tz)
    pdf.cell(0, 7, f'Exportado: {now_ar.strftime("%d/%m/%Y %H:%M")}', ln=True)
    pdf.cell(0, 7, f'Total de mensajes: {len(messages)}', ln=True)
    pdf.ln(5)

    # Línea separadora
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    # --- Mensajes ---
    pdf.set_font('Helvetica', 'B', 12)
    pdf.cell(0, 8, 'Mensajes', ln=True)
    pdf.ln(3)

    current_date = None

    for msg in messages:
        # Convertir timestamp a Argentina
        ts = msg.timestamp
        if ts.tzinfo is None:
            ts = ts.replace(tzinfo=pytz.utc)
        ts_ar = ts.astimezone(tz)

        # Separador por día
        msg_date = ts_ar.strftime('%d/%m/%Y')
        if msg_date != current_date:
            current_date = msg_date
            pdf.ln(3)
            pdf.set_font('Helvetica', 'B', 9)
            pdf.set_text_color(120, 120, 120)
            pdf.cell(0, 6, f'--- {msg_date} ---', ln=True, align='C')
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)

        # Dirección
        if msg.direction == 'inbound':
            sender = _safe_text(contact_name)
            pdf.set_text_color(0, 100, 0)
        else:
            sender = 'Bot / Operador'
            pdf.set_text_color(0, 50, 150)

        time_str = ts_ar.strftime('%H:%M')

        # Sender + hora
        pdf.set_font('Helvetica', 'B', 9)
        pdf.cell(0, 5, f'{sender}  [{time_str}]', ln=True)

        # Contenido
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Helvetica', '', 9)
        content = msg.content or f'[{msg.message_type}]'
        content = _safe_text(content)

        # Multi-line cell para textos largos
        pdf.multi_cell(0, 5, content)
        pdf.ln(2)

    # --- Notas internas ---
    if notes:
        pdf.ln(5)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Notas internas del equipo', ln=True)
        pdf.ln(3)

        for note in notes:
            ts = note.created_at
            if ts and ts.tzinfo is None:
                ts = ts.replace(tzinfo=pytz.utc)
            ts_ar = ts.astimezone(tz) if ts else None

            pdf.set_font('Helvetica', 'I', 9)
            pdf.set_text_color(100, 100, 100)
            date_str = ts_ar.strftime('%d/%m/%Y %H:%M') if ts_ar else ''
            pdf.cell(0, 5, date_str, ln=True)

            pdf.set_text_color(0, 0, 0)
            pdf.set_font('Helvetica', '', 9)
            pdf.multi_cell(0, 5, _safe_text(note.content))
            pdf.ln(3)

    # Generar PDF
    pdf_bytes = pdf.output()
    safe_name = contact_name.replace(' ', '_')[:30] if contact_name else phone
    filename = f'conversacion_{safe_name}_{now_ar.strftime("%Y%m%d")}.pdf'

    return send_file(
        io.BytesIO(pdf_bytes),
        mimetype='application/pdf',
        download_name=filename,
        as_attachment=True
    )


def _safe_text(text):
    """Limpia texto para PDF (remueve caracteres no soportados por latin-1)."""
    if not text:
        return ''
    return text.encode('latin-1', errors='replace').decode('latin-1')


# ==================== CONVERSATION SESSIONS ====================

@app.route("/sessions")
def sessions_page():
    """Página para ver sesiones de conversación categorizadas."""
    return render_template('sessions.html')


@app.route("/api/conversation-sessions", methods=["GET"])
def api_list_conversation_sessions():
    """Lista sesiones de conversación con filtros."""
    # Filters
    topic_id = request.args.get('topic_id', type=int)
    rating = request.args.get('rating')
    search = request.args.get('search', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)

    query = ConversationSession.query

    if topic_id:
        query = query.filter(ConversationSession.topic_id == topic_id)
    if rating:
        # Mapeo para manejar variantes con/sin acentos
        rating_variants = {
            'problematica': ['problematica', 'problemática'],
            'excelente': ['excelente'],
            'buena': ['buena'],
            'neutral': ['neutral'],
            'mala': ['mala']
        }
        variants = rating_variants.get(rating, [rating])
        query = query.filter(ConversationSession.rating.in_(variants))

    # Filtro por estado de asistencia humana
    status_filter = request.args.get('status')
    if status_filter == 'unanswered':
        query = query.filter(ConversationSession.has_unanswered_questions == True)
    elif status_filter == 'escalated':
        query = query.filter(ConversationSession.escalated_to_human == True)

    # Búsqueda por texto
    if search:
        # Buscar por número de teléfono, resumen o nombre de contacto
        query = query.outerjoin(Contact, ConversationSession.phone_number == Contact.phone_number)
        search_filter = or_(
            ConversationSession.phone_number.ilike(f'%{search}%'),
            ConversationSession.summary.ilike(f'%{search}%'),
            Contact.name.ilike(f'%{search}%')
        )
        query = query.filter(search_filter)

    # Order by most recent first
    query = query.order_by(ConversationSession.ended_at.desc())

    # Paginate
    paginated = query.paginate(page=page, per_page=per_page, error_out=False)

    # Obtener nombres de contactos en batch
    phones = [s.phone_number for s in paginated.items]
    contacts_map = {}
    if phones:
        contacts = Contact.query.filter(Contact.phone_number.in_(phones)).all()
        contacts_map = {c.phone_number: c.name for c in contacts}

    # Agregar nombre de contacto a cada sesión
    sessions_data = []
    for s in paginated.items:
        session_dict = s.to_dict()
        session_dict['contact_name'] = contacts_map.get(s.phone_number)
        sessions_data.append(session_dict)

    return jsonify({
        'sessions': sessions_data,
        'total': paginated.total,
        'pages': paginated.pages,
        'current_page': page
    })


@app.route("/api/conversation-sessions/<int:session_id>", methods=["PUT"])
def api_update_conversation_session(session_id):
    """Recategoriza una sesión manualmente."""
    session = ConversationSession.query.get_or_404(session_id)
    data = request.get_json()
    
    if 'topic_id' in data:
        session.topic_id = data['topic_id'] if data['topic_id'] else None
    
    if 'rating' in data:
        session.rating = data['rating']
    
    session.auto_categorized = False  # Mark as manually edited
    db.session.commit()
    
    return jsonify({'success': True, 'session': session.to_dict()})


@app.route("/api/categorize/force", methods=["POST"])
def api_force_categorize():
    """Fuerza la categorización inmediata de conversaciones."""
    import conversation_categorizer

    data = request.get_json(silent=True) or {}
    phone = data.get('phone')

    def run_force():
        # Guardar valor original y restaurar después
        original_inactivity = conversation_categorizer.INACTIVITY_MINUTES
        conversation_categorizer.INACTIVITY_MINUTES = 0
        try:
            conversation_categorizer.run_categorization(
                app.app_context(),
                force_phone=phone
            )
        finally:
            conversation_categorizer.INACTIVITY_MINUTES = original_inactivity

    t = threading.Thread(target=run_force)
    t.daemon = True
    t.start()

    msg = f'Clasificación forzada iniciada para {phone}' if phone else 'Clasificación forzada iniciada para todas las conversaciones'
    logger.info(f"⚡ [CATEGORIZER] {msg}")

    return jsonify({
        'success': True,
        'message': msg
    })


# ==========================================
# RAG DOCUMENTS API
# ==========================================

@app.route("/chatbot")
def chatbot_page():
    """Página de gestión del chatbot."""
    documents = RagDocument.query.order_by(RagDocument.created_at.desc()).all()
    chatbot_enabled = ChatbotConfig.get('enabled', 'true') == 'true'
    system_prompt = ChatbotConfig.get('system_prompt', '')
    return render_template('chatbot.html', documents=documents, chatbot_enabled=chatbot_enabled, system_prompt=system_prompt)


@app.route("/reengagement")
def reengagement_page():
    """Página de re-engagement automático."""
    vis_tag_ids = get_visible_tag_ids(g.current_user)
    tags_q = Tag.query.filter_by(is_active=True)
    if vis_tag_ids is not None:
        tags_q = tags_q.filter(Tag.id.in_(vis_tag_ids)) if vis_tag_ids else tags_q.filter(False)
    tags = tags_q.order_by(Tag.name).all()
    auto_tagger_enabled = ChatbotConfig.get('auto_tagger_enabled', 'true') == 'true'
    return render_template('reengagement.html', tags=tags, auto_tagger_enabled=auto_tagger_enabled)

@app.route("/api/auto-tagger/stats", methods=["GET"])
def api_auto_tagger_stats():
    """Stats del sistema de re-engagement para el dashboard."""
    from sqlalchemy import func
    now = datetime.utcnow()
    last_7d = now - timedelta(days=7)

    # Logs últimos 7 días
    logs_7d = AutoTagLog.query.filter(AutoTagLog.created_at >= last_7d).all()
    tagged_7d  = sum(1 for l in logs_7d if l.result == 'tagged')
    skipped_7d = sum(1 for l in logs_7d if l.result == 'skipped')
    error_7d   = sum(1 for l in logs_7d if l.result == 'error')

    # Enrollments
    total_enrollments  = FollowUpEnrollment.query.count()
    active_enrollments = FollowUpEnrollment.query.filter_by(status='pending').count()
    finished           = FollowUpEnrollment.query.filter_by(status='finished').count()
    cancelled          = FollowUpEnrollment.query.filter_by(status='cancelled').count()

    # Tasa de respuesta: cancelados (respondieron) / (cancelados + finalizados)
    responded = cancelled
    total_closed = cancelled + finished
    response_rate = round((responded / total_closed * 100), 1) if total_closed > 0 else 0

    return jsonify({
        'tagged_7d': tagged_7d,
        'skipped_7d': skipped_7d,
        'error_7d': error_7d,
        'total_enrollments': total_enrollments,
        'active_enrollments': active_enrollments,
        'finished_enrollments': finished,
        'cancelled_enrollments': cancelled,
        'response_rate': response_rate
    })

@app.route("/api/auto-tagger/logs", methods=["GET"])
def api_auto_tagger_logs():
    """Log de actividad del auto-tagger."""
    limit = min(int(request.args.get('limit', 50)), 200)
    result_filter = request.args.get('result')
    q = AutoTagLog.query
    if result_filter:
        q = q.filter_by(result=result_filter)
    logs = q.order_by(AutoTagLog.created_at.desc()).limit(limit).all()
    return jsonify([l.to_dict() for l in logs])

@app.route("/api/auto-tagger/run", methods=["POST"])
def api_run_auto_tagger():
    """Corre el auto-tagger manualmente de forma asíncrona."""
    import threading
    from auto_tagger import run_auto_tagger
    ctx = app.app_context()
    t = threading.Thread(target=run_auto_tagger, args=(ctx,))
    t.daemon = True
    t.start()
    return jsonify({'success': True, 'message': 'Auto-tagger iniciado'})


@app.route("/api/auto-tagger/toggle", methods=["POST"])
def api_toggle_auto_tagger():
    """Enciende/apaga el auto tagger."""
    try:
        db.session.rollback()
        current = ChatbotConfig.get('auto_tagger_enabled', 'true')
        new_value = 'false' if current == 'true' else 'true'
        ChatbotConfig.set('auto_tagger_enabled', new_value)
        return jsonify({'success': True, 'enabled': new_value == 'true'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route("/api/rag/documents", methods=["GET"])
def api_list_rag_documents():
    """Lista todos los documentos RAG."""
    documents = RagDocument.query.order_by(RagDocument.created_at.desc()).all()
    return jsonify({'documents': [d.to_dict() for d in documents]})


@app.route("/api/rag/documents", methods=["POST"])
def api_upload_rag_document():
    """Sube un documento para RAG."""
    import hashlib
    from whatsapp_service import get_s3_client, ensure_rag_bucket_exists

    if 'file' not in request.files:
        return jsonify({'error': 'No se envió ningún archivo'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío'}), 400

    # Validar tipo de archivo
    allowed_extensions = {'pdf', 'xlsx', 'xls', 'csv', 'docx', 'doc', 'txt'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed_extensions:
        return jsonify({'error': f'Tipo de archivo no permitido. Permitidos: {", ".join(allowed_extensions)}'}), 400

    try:
        # Leer contenido del archivo
        file_content = file.read()
        file_size = len(file_content)

        # Calcular hash SHA256
        file_hash = hashlib.sha256(file_content).hexdigest()

        # Verificar si ya existe un documento con el mismo nombre
        existing = RagDocument.query.filter_by(original_filename=file.filename).first()

        if existing:
            # Verificar si el contenido cambió
            if existing.file_hash == file_hash:
                return jsonify({
                    'error': 'El archivo ya existe y no ha sido modificado',
                    'document': existing.to_dict()
                }), 409

            # Archivo modificado - actualizar
            action = 'updated'
            doc = existing
            doc.file_hash = file_hash
            doc.file_size = file_size
            doc.status = 'pending'
            doc.error_message = None
        else:
            # Nuevo archivo
            action = 'created'
            filename = f"{file_hash[:8]}_{file.filename}"
            doc = RagDocument(
                filename=filename,
                original_filename=file.filename,
                file_type=ext,
                file_size=file_size,
                file_hash=file_hash,
                minio_path=f"rag-documents/{filename}",
                status='pending'
            )
            db.session.add(doc)

        # Subir a MinIO
        ensure_rag_bucket_exists()
        s3 = get_s3_client()
        bucket = Config.MINIO_BUCKET_RAG

        # Determinar content type
        content_types = {
            'pdf': 'application/pdf',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'xls': 'application/vnd.ms-excel',
            'csv': 'text/csv',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'txt': 'text/plain'
        }

        s3.put_object(
            Bucket=bucket,
            Key=doc.filename,
            Body=file_content,
            ContentType=content_types.get(ext, 'application/octet-stream')
        )

        db.session.commit()

        # Llamar webhook de n8n para vectorizar
        webhook_url = Config.N8N_WEBHOOK_VECTORIZE
        if webhook_url:
            try:
                # Construir URL pública de MinIO
                protocol = "https" if Config.MINIO_USE_SSL else "http"
                minio_url = f"{protocol}://{Config.MINIO_ENDPOINT}/{bucket}/{doc.filename}"

                # URL de callback para que n8n actualice el estado
                callback_url = f"{Config.FLASK_BASE_URL}/api/rag/documents/{doc.id}/status"

                requests.post(webhook_url, json={
                    'action': action,
                    'document_id': doc.id,
                    'filename': doc.filename,
                    'original_filename': doc.original_filename,
                    'file_type': doc.file_type,
                    'minio_bucket': bucket,
                    'minio_path': doc.minio_path,
                    'minio_url': minio_url,
                    'callback_url': callback_url
                }, timeout=5)
                doc.status = 'processing'
                db.session.commit()
            except Exception as e:
                logger.warning(f"No se pudo notificar a n8n: {e}")

        return jsonify({
            'success': True,
            'action': action,
            'document': doc.to_dict()
        })

    except Exception as e:
        logger.error(f"Error subiendo documento RAG: {e}")
        return jsonify({'error': str(e)}), 500



@app.route("/api/rag/documents/<int:doc_id>", methods=["DELETE"])
def api_delete_rag_document(doc_id):
    """Elimina un documento RAG."""
    from whatsapp_service import get_s3_client

    doc = RagDocument.query.get_or_404(doc_id)

    try:
        # Eliminar de MinIO
        s3 = get_s3_client()
        bucket = Config.MINIO_BUCKET_RAG

        try:
            s3.delete_object(Bucket=bucket, Key=doc.filename)
        except Exception as e:
            logger.warning(f"No se pudo eliminar de MinIO: {e}")

        # Llamar webhook de n8n para limpiar vectores
        webhook_url = Config.N8N_WEBHOOK_DELETE
        if webhook_url:
            try:
                requests.post(webhook_url, json={
                    'action': 'deleted',
                    'document_id': doc.id,
                    'filename': doc.filename,
                    'original_filename': doc.original_filename
                }, timeout=5)
            except Exception as e:
                logger.warning(f"No se pudo notificar a n8n: {e}")

        # Eliminar de la BD
        db.session.delete(doc)
        db.session.commit()

        return jsonify({'success': True, 'message': 'Documento eliminado'})

    except Exception as e:
        logger.error(f"Error eliminando documento RAG: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/rag/documents/<int:doc_id>/download")
def api_download_rag_document(doc_id):
    """Descarga un documento RAG."""
    from whatsapp_service import get_s3_client

    doc = RagDocument.query.get_or_404(doc_id)

    try:
        s3 = get_s3_client()
        bucket = Config.MINIO_BUCKET_RAG

        response = s3.get_object(Bucket=bucket, Key=doc.filename)
        content_type = response.get('ContentType', 'application/octet-stream')

        return send_file(
            io.BytesIO(response['Body'].read()),
            mimetype=content_type,
            download_name=doc.original_filename,
            as_attachment=True
        )

    except Exception as e:
        logger.error(f"Error descargando documento RAG: {e}")
        return jsonify({'error': str(e)}), 500


@app.route("/api/rag/documents/<int:doc_id>/status", methods=["PUT"])
def api_update_rag_document_status(doc_id):
    """Actualiza el estado de un documento (llamado por n8n)."""
    doc = RagDocument.query.get_or_404(doc_id)
    data = request.get_json()

    if 'status' in data:
        doc.status = data['status']
    if 'error_message' in data:
        doc.error_message = data['error_message']

    db.session.commit()
    return jsonify({'success': True, 'document': doc.to_dict()})


# ==========================================
# CHATBOT CONFIG API
# ==========================================

@app.route("/api/chatbot/config", methods=["GET"])
def api_get_chatbot_config():
    """Obtiene la configuración del chatbot."""
    enabled = ChatbotConfig.get('enabled', 'true')
    return jsonify({
        'enabled': enabled == 'true',
        'webhook_vectorize': Config.N8N_WEBHOOK_VECTORIZE or '',
        'webhook_delete': Config.N8N_WEBHOOK_DELETE or ''
    })


@app.route("/api/chatbot/config", methods=["PUT"])
def api_update_chatbot_config():
    """Actualiza la configuración del chatbot."""
    data = request.get_json()

    if 'enabled' in data:
        ChatbotConfig.set('enabled', 'true' if data['enabled'] else 'false')

    return jsonify({'success': True})


@app.route("/api/chatbot/system-prompt", methods=["GET"])
def api_get_system_prompt():
    """Obtiene el system prompt del chatbot."""
    return jsonify({'system_prompt': ChatbotConfig.get('system_prompt', '')})


@app.route("/api/chatbot/system-prompt", methods=["POST"])
def api_save_system_prompt():
    """Guarda el system prompt del chatbot."""
    data = request.get_json()
    prompt = data.get('system_prompt', '').strip()
    ChatbotConfig.set('system_prompt', prompt)
    return jsonify({'success': True})


@app.route("/api/chatbot/toggle", methods=["POST"])
def api_toggle_chatbot():
    """Enciende/apaga el chatbot y el workflow de n8n."""
    try:
        db.session.rollback()  # Limpiar cualquier transaccion colgada
        current = ChatbotConfig.get('enabled', 'true')
        new_value = 'false' if current == 'true' else 'true'
        ChatbotConfig.set('enabled', new_value)
    except Exception as e:
        db.session.rollback()
        logger.error(f"❌ Error al guardar estado del chatbot en DB: {e}")
        return jsonify({'success': False, 'error': f'Error de base de datos: {str(e)}'}), 500

    n8n_result = None

    # También activar/desactivar el workflow de n8n
    if Config.N8N_CHATBOT_WORKFLOW_ID and Config.N8N_API_URL and Config.N8N_API_KEY:
        try:
            action = 'activate' if new_value == 'true' else 'deactivate'
            url = f"{Config.N8N_API_URL}/workflows/{Config.N8N_CHATBOT_WORKFLOW_ID}/{action}"
            headers = {
                "X-N8N-API-KEY": Config.N8N_API_KEY,
                "Content-Type": "application/json"
            }
            response = requests.post(url, headers=headers, timeout=10)
            if response.status_code == 200:
                n8n_result = 'ok'
                logger.info(f"✅ Workflow n8n {action}d: {Config.N8N_CHATBOT_WORKFLOW_ID}")
            else:
                n8n_result = f'error: {response.status_code}'
                logger.warning(f"⚠️ Error al {action} workflow n8n: {response.text}")
        except Exception as e:
            n8n_result = f'error: {str(e)}'
            logger.error(f"❌ Error conectando a n8n: {e}")

    return jsonify({
        'success': True,
        'enabled': new_value == 'true',
        'n8n_workflow': n8n_result
    })


# ==========================================
# N8N WORKFLOW CONTROL API
# ==========================================

def n8n_api_request(method, endpoint, data=None):
    """Helper para hacer requests a la API de n8n."""
    if not Config.N8N_API_URL or not Config.N8N_API_KEY:
        return {'error': 'n8n API no configurada'}, 500
    
    url = f"{Config.N8N_API_URL}{endpoint}"
    headers = {
        "X-N8N-API-KEY": Config.N8N_API_KEY,
        "Content-Type": "application/json"
    }
    
    try:
        if method == 'GET':
            response = requests.get(url, headers=headers, timeout=10)
        elif method == 'POST':
            response = requests.post(url, headers=headers, json=data or {}, timeout=10)
        elif method == 'PATCH':
            response = requests.patch(url, headers=headers, json=data or {}, timeout=10)
        else:
            return {'error': f'Método {method} no soportado'}, 400
        
        if response.status_code >= 400:
            return {'error': response.text}, response.status_code
        
        return response.json(), 200
    except requests.exceptions.Timeout:
        return {'error': 'Timeout conectando a n8n'}, 504
    except Exception as e:
        return {'error': str(e)}, 500


@app.route("/api/n8n/workflows", methods=["GET"])
def api_list_n8n_workflows():
    """Lista todos los workflows de n8n."""
    result, status = n8n_api_request('GET', '/workflows')
    return jsonify(result), status


@app.route("/api/n8n/workflows/<workflow_id>", methods=["GET"])
def api_get_n8n_workflow(workflow_id):
    """Obtiene detalles de un workflow específico."""
    result, status = n8n_api_request('GET', f'/workflows/{workflow_id}')
    return jsonify(result), status


@app.route("/api/n8n/workflows/<workflow_id>/activate", methods=["POST"])
def api_activate_n8n_workflow(workflow_id):
    """Activa un workflow de n8n."""
    result, status = n8n_api_request('POST', f'/workflows/{workflow_id}/activate')
    if status == 200:
        return jsonify({'success': True, 'message': 'Workflow activado', 'workflow': result})
    return jsonify(result), status


@app.route("/api/n8n/workflows/<workflow_id>/deactivate", methods=["POST"])
def api_deactivate_n8n_workflow(workflow_id):
    """Desactiva un workflow de n8n."""
    result, status = n8n_api_request('POST', f'/workflows/{workflow_id}/deactivate')
    if status == 200:
        return jsonify({'success': True, 'message': 'Workflow desactivado', 'workflow': result})
    return jsonify(result), status


@app.route("/api/n8n/workflows/<workflow_id>/toggle", methods=["POST"])
def api_toggle_n8n_workflow(workflow_id):
    """Alterna el estado de un workflow (activo/inactivo)."""
    # Primero obtener estado actual
    workflow, status = n8n_api_request('GET', f'/workflows/{workflow_id}')
    if status != 200:
        return jsonify(workflow), status
    
    is_active = workflow.get('active', False)
    
    # Alternar
    if is_active:
        result, status = n8n_api_request('POST', f'/workflows/{workflow_id}/deactivate')
        action = 'desactivado'
    else:
        result, status = n8n_api_request('POST', f'/workflows/{workflow_id}/activate')
        action = 'activado'
    
    if status == 200:
        return jsonify({
            'success': True, 
            'message': f'Workflow {action}',
            'active': not is_active,
            'workflow': result
        })
    return jsonify(result), status


# ==========================================
# DOCX TEXT EXTRACTION API (para n8n)
# ==========================================

@app.route("/api/extract-docx", methods=["POST"])
def api_extract_docx():
    """Extrae texto de un archivo DOCX. Usado por n8n para RAG."""
    from docx import Document
    
    if 'file' not in request.files:
        return jsonify({'error': 'No se envió ningún archivo'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío'}), 400
    
    try:
        doc = Document(io.BytesIO(file.read()))
        
        # Extraer todos los párrafos
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # También extraer texto de tablas
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
        
        # Combinar todo el texto
        full_text = '\n'.join(paragraphs)
        if table_text:
            full_text += '\n\n--- Tablas ---\n' + '\n'.join(table_text)
        
        return jsonify({
            'success': True,
            'data': full_text,
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables)
        })
        
    except Exception as e:
        logger.error(f"Error extrayendo texto de DOCX: {e}")
        return jsonify({'error': str(e)}), 500


# =====================================================
# CATALOG
# =====================================================

def _sync_catalog_to_db(catalog_id):
    """Sincroniza productos de Meta a la tabla local. Retorna (count, error)."""
    from models import CatalogProduct, ChatbotConfig
    from whatsapp_service import whatsapp_api
    result = whatsapp_api.sync_catalog_products(catalog_id)
    if "error" in result:
        return 0, result["error"]
    products = result.get("products", [])
    now = datetime.utcnow()
    for p in products:
        rid = p.get("retailer_id") or p.get("id")
        if not rid:
            continue
        raw_price = p.get("price")
        try:
            if isinstance(raw_price, (int, float)):
                price = float(raw_price)
            elif raw_price:
                # Meta puede mandar "$15.000,00" (formato AR) o "15000.00" (formato US)
                cleaned = str(raw_price).strip().replace('$', '').replace(' ', '')
                # Detectar formato: si tiene coma después del último punto → AR
                if ',' in cleaned and cleaned.rfind(',') > cleaned.rfind('.'):
                    cleaned = cleaned.replace('.', '').replace(',', '.')
                else:
                    cleaned = cleaned.replace(',', '')
                price = float(cleaned)
            else:
                price = None
        except (ValueError, TypeError):
            price = None
        existing = CatalogProduct.query.get(rid)
        if existing:
            existing.wa_product_id = p.get("id", existing.wa_product_id)
            existing.name = p.get("name", existing.name)
            existing.description = p.get("description", existing.description)
            existing.price = price
            existing.currency = p.get("currency", existing.currency)
            existing.availability = p.get("availability", existing.availability)
            existing.image_url = p.get("image_url", existing.image_url)
            existing.synced_at = now
        else:
            db.session.add(CatalogProduct(
                retailer_id=rid,
                wa_product_id=p.get("id"),
                name=p.get("name"),
                description=p.get("description"),
                price=price,
                currency=p.get("currency"),
                availability=p.get("availability", "in stock"),
                image_url=p.get("image_url"),
                synced_at=now,
            ))
    # Eliminar productos locales que ya no existen en Meta
    synced_ids = {p.get("retailer_id") or p.get("id") for p in products if p.get("retailer_id") or p.get("id")}
    local_products = CatalogProduct.query.all()
    for lp in local_products:
        if lp.retailer_id not in synced_ids:
            db.session.delete(lp)

    db.session.commit()
    return len(products), None


@app.route("/catalog")
def catalog_page():
    if not g.current_user or not g.current_user.has_permission('catalog'):
        return redirect(url_for('login'))
    return render_template('catalog.html')


@app.route("/bot-audios")
def bot_audios_page():
    if not g.current_user or not g.current_user.is_admin:
        return redirect(url_for('login'))
    return render_template('bot_audios.html')


@app.route("/api/catalog/upload-image", methods=["POST"])
def api_catalog_upload_image():
    """Sube una imagen de producto a MinIO y devuelve la URL pública absoluta."""
    from whatsapp_service import get_s3_client, ensure_bucket_exists
    if 'file' not in request.files:
        return jsonify({"error": "No se recibió archivo"}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({"error": "Archivo sin nombre"}), 400
    allowed = {'image/jpeg', 'image/png', 'image/webp', 'image/gif'}
    mime = f.mimetype or ''
    if mime not in allowed:
        return jsonify({"error": "Solo se permiten imágenes (jpg, png, webp)"}), 400
    ext_map = {'image/jpeg': '.jpg', 'image/png': '.png', 'image/webp': '.webp', 'image/gif': '.gif'}
    ext = ext_map.get(mime, '.jpg')
    filename = f"catalog-images/{uuid.uuid4().hex}{ext}"
    try:
        s3 = get_s3_client()
        ensure_bucket_exists()
        file_bytes = f.read()
        s3.put_object(
            Bucket=Config.MINIO_BUCKET,
            Key=filename,
            Body=file_bytes,
            ContentType=mime
        )
        base = Config.FLASK_BASE_URL.rstrip('/')
        public_url = f"{base}/media/{filename}"
        return jsonify({"url": public_url})
    except Exception as e:
        logger.error(f"Error subiendo imagen de catálogo a MinIO: {e}")
        return jsonify({"error": "Error al subir la imagen"}), 500


@app.route("/api/catalog/detect", methods=["GET"])
def api_catalog_detect():
    """Auto-detecta el catalog_id del WABA."""
    from whatsapp_service import whatsapp_api
    result = whatsapp_api.get_catalogs()
    if "error" in result:
        return jsonify(result), 500
    catalogs = result.get("catalogs", [])
    saved = ChatbotConfig.get("catalog_id")
    return jsonify({"catalogs": catalogs, "saved_catalog_id": saved})


@app.route("/api/catalog/set", methods=["POST"])
def api_catalog_set():
    """Guarda el catalog_id elegido y dispara la primera sync."""
    data = request.json or {}
    catalog_id = data.get("catalog_id", "").strip()
    if not catalog_id:
        return jsonify({"error": "catalog_id requerido"}), 400
    ChatbotConfig.set("catalog_id", catalog_id)
    count, err = _sync_catalog_to_db(catalog_id)
    if err:
        return jsonify({"error": err}), 500
    return jsonify({"success": True, "synced": count})


@app.route("/api/catalog/sync", methods=["POST"])
def api_catalog_sync():
    """Sincroniza el catálogo ahora."""
    catalog_id = ChatbotConfig.get("catalog_id")
    if not catalog_id:
        return jsonify({"error": "Catálogo no configurado"}), 400
    count, err = _sync_catalog_to_db(catalog_id)
    if err:
        return jsonify({"error": err}), 500
    return jsonify({"success": True, "synced": count})


@app.route("/api/bot/catalog", methods=["GET"])
def api_bot_catalog():
    """Endpoint público para el bot de n8n — devuelve nombre, precio y stock."""
    from models import CatalogProduct
    products = CatalogProduct.query.filter_by(availability='in stock').order_by(CatalogProduct.name).all()
    return jsonify([
        {
            "nombre": p.name,
            "precio": float(p.price) if p.price is not None else None,
            "moneda": p.currency or "ARS",
            "stock": "disponible" if p.availability == "in stock" else "sin stock"
        }
        for p in products
    ])


@app.route("/api/bot/audios", methods=["GET"])
def api_bot_audios_list():
    """Endpoint para el bot de n8n — devuelve lista de audios disponibles."""
    from models import BotAudio
    audios = BotAudio.query.order_by(BotAudio.created_at).all()
    return jsonify([a.to_dict() for a in audios])


@app.route("/api/bot/send-audio", methods=["POST"])
def api_bot_send_audio():
    """Endpoint para el bot de n8n — envía un audio pregrabado a un usuario."""
    data = request.json
    if not data:
        return jsonify({"error": "No data"}), 400
    audio_id = data.get("audio_id")
    phone_number = str(data.get("phone_number", "")).replace("+", "").strip()
    if not audio_id or not phone_number:
        return jsonify({"error": "audio_id y phone_number son requeridos"}), 400

    from models import BotAudio
    audio = BotAudio.query.get(audio_id)
    if not audio:
        return jsonify({"error": "Audio no encontrado"}), 404

    # Leer el archivo directamente desde MinIO usando el cliente S3
    try:
        from whatsapp_service import get_s3_client
        from config import Config
        # file_url es "/media/bot_audios/..." — la key en MinIO es sin el prefijo /media/
        minio_key = audio.file_url.lstrip('/media/').lstrip('static/media/')
        s3 = get_s3_client()
        obj = s3.get_object(Bucket=Config.MINIO_BUCKET, Key=minio_key)
        file_bytes = obj['Body'].read()
    except Exception as e:
        logger.error(f"Error descargando audio {audio_id} desde MinIO: {e}")
        return jsonify({"error": "No se pudo descargar el audio"}), 500

    # Subir a WhatsApp y enviar
    filename = f"{audio.nombre.replace(' ', '_')}.mp3"
    upload_result = whatsapp_api.upload_media(file_bytes, audio.mime_type, filename)
    if not upload_result.get("success"):
        return jsonify({"error": "Error subiendo audio a WhatsApp: " + upload_result.get("error", "")}), 500

    send_result = whatsapp_api.send_media_message(phone_number, "audio", upload_result["media_id"])
    if not send_result.get("success"):
        return jsonify({"error": "Error enviando audio: " + send_result.get("error", "")}), 500

    # Guardar el mensaje en la BD para que aparezca en el chat del dashboard
    try:
        from models import Message
        msg = Message(
            wa_message_id=send_result.get("wa_message_id"),
            phone_number=phone_number,
            direction="outbound",
            message_type="audio",
            media_url=audio.file_url,
            content=None,
            sent_by="bot",
        )
        db.session.add(msg)
        db.session.commit()
    except Exception as e:
        logger.warning(f"No se pudo guardar el mensaje de audio en BD: {e}")

    logger.info(f"🎵 Audio '{audio.nombre}' enviado a {phone_number} por el bot")
    return jsonify({"success": True, "audio": audio.nombre})


@app.route("/api/bot/audios", methods=["POST"])
def api_bot_audios_upload():
    """Sube un nuevo audio a la biblioteca del bot."""
    nombre = request.form.get("nombre", "").strip()
    descripcion = request.form.get("descripcion", "").strip()
    file = request.files.get("file")

    if not nombre or not descripcion or not file:
        return jsonify({"error": "nombre, descripcion y file son requeridos"}), 400

    file_bytes = file.read()
    if len(file_bytes) == 0:
        return jsonify({"error": "El archivo está vacío"}), 400

    mime_type = file.content_type or "audio/mpeg"
    original_filename = file.filename or f"{nombre}.mp3"

    # Convertir a MP3 si es necesario
    WHATSAPP_AUDIO_OK = {'audio/aac', 'audio/mp4', 'audio/mpeg', 'audio/amr', 'audio/ogg', 'audio/opus'}
    base_mime = mime_type.split(';')[0].strip().lower()
    if base_mime.startswith("audio/") and base_mime not in WHATSAPP_AUDIO_OK:
        try:
            import av, io
            input_buf = io.BytesIO(file_bytes)
            output_buf = io.BytesIO()
            with av.open(input_buf) as in_container:
                in_stream = in_container.streams.audio[0]
                with av.open(output_buf, 'w', format='mp3') as out_container:
                    out_stream = out_container.add_stream('libmp3lame', rate=44100)
                    out_stream.layout = 'mono'
                    for frame in in_container.decode(in_stream):
                        frame.pts = None
                        for packet in out_stream.encode(frame):
                            out_container.mux(packet)
                    for packet in out_stream.encode(None):
                        out_container.mux(packet)
            file_bytes = output_buf.getvalue()
            mime_type = 'audio/mpeg'
            original_filename = original_filename.rsplit('.', 1)[0] + '.mp3'
        except Exception as e:
            logger.warning(f"Error convirtiendo audio: {e}")

    # Subir a MinIO
    file_url = whatsapp_api.upload_to_minio(file_bytes, mime_type, f"bot_audios/{original_filename}")
    if not file_url:
        return jsonify({"error": "Error subiendo archivo a MinIO"}), 500

    from models import BotAudio
    audio = BotAudio(nombre=nombre, descripcion=descripcion, file_url=file_url, mime_type=mime_type)
    db.session.add(audio)
    db.session.commit()
    return jsonify({"success": True, "audio": audio.to_dict()})


@app.route("/api/bot/audios/<int:audio_id>", methods=["DELETE"])
def api_bot_audios_delete(audio_id):
    """Elimina un audio de la biblioteca."""
    from models import BotAudio
    audio = BotAudio.query.get_or_404(audio_id)
    db.session.delete(audio)
    db.session.commit()
    return jsonify({"success": True})


@app.route("/api/catalog/products", methods=["GET"])
def api_catalog_products_list():
    from models import CatalogProduct
    products = CatalogProduct.query.order_by(CatalogProduct.name).all()
    last_sync = None
    if products:
        synced_times = [p.synced_at for p in products if p.synced_at]
        if synced_times:
            last_sync = max(synced_times).isoformat()
    catalog_id = ChatbotConfig.get("catalog_id")
    return jsonify({
        "products": [p.to_dict() for p in products],
        "last_sync": last_sync,
        "catalog_id": catalog_id,
    })


@app.route("/api/catalog/products", methods=["POST"])
def api_catalog_products_create():
    from models import CatalogProduct
    from whatsapp_service import whatsapp_api
    catalog_id = ChatbotConfig.get("catalog_id")
    if not catalog_id:
        return jsonify({"error": "Catálogo no configurado"}), 400
    data = request.json or {}
    retailer_id = data.get("retailer_id", "").strip()
    name = data.get("name", "").strip()
    if not retailer_id or not name:
        return jsonify({"error": "retailer_id y name son requeridos"}), 400
    if CatalogProduct.query.get(retailer_id):
        return jsonify({"error": f"Ya existe un producto con Retailer ID '{retailer_id}'"}), 409
    price = data.get("price")
    currency = data.get("currency", "ARS")
    description = data.get("description", "")
    availability = data.get("availability", "in stock")
    image_url = data.get("image_url") or None

    # Crear en Meta
    meta_result = whatsapp_api.create_catalog_product(
        catalog_id, retailer_id, name, price or 0, currency, description, availability, image_url=image_url
    )
    if "error" in meta_result:
        return jsonify(meta_result), 500

    # Guardar en local
    product = CatalogProduct(
        retailer_id=retailer_id,
        wa_product_id=meta_result.get("data", {}).get("id"),
        name=name,
        description=description,
        price=price,
        currency=currency,
        availability=availability,
        image_url=image_url,
        synced_at=datetime.utcnow(),
    )
    db.session.add(product)
    db.session.commit()
    return jsonify({"success": True, "product": product.to_dict()}), 201


@app.route("/api/catalog/products/<retailer_id>", methods=["PUT"])
def api_catalog_products_update(retailer_id):
    from models import CatalogProduct
    from whatsapp_service import whatsapp_api
    product = CatalogProduct.query.get_or_404(retailer_id)
    data = request.json or {}

    name = data.get("name", product.name)
    price = data.get("price", product.price)
    currency = data.get("currency", product.currency)
    description = data.get("description", product.description)
    availability = data.get("availability", product.availability)
    image_url = data.get("image_url", product.image_url) or None

    # Actualizar en Meta si tenemos wa_product_id
    if product.wa_product_id:
        meta_result = whatsapp_api.update_catalog_product(
            product.wa_product_id, name=name, price=price,
            currency=currency, description=description, availability=availability,
            image_url=image_url
        )
        if "error" in meta_result:
            return jsonify(meta_result), 500

    product.name = name
    product.price = price
    product.currency = currency
    product.description = description
    product.availability = availability
    product.image_url = image_url
    product.synced_at = datetime.utcnow()
    db.session.commit()
    return jsonify({"success": True, "product": product.to_dict()})


@app.route("/api/catalog/products/<retailer_id>", methods=["DELETE"])
def api_catalog_products_delete(retailer_id):
    from models import CatalogProduct
    from whatsapp_service import whatsapp_api
    product = CatalogProduct.query.get_or_404(retailer_id)

    if product.wa_product_id:
        whatsapp_api.delete_catalog_product(product.wa_product_id)
        # Ignoramos error de Meta: si ya fue eliminado allá, igual lo borramos localmente

    db.session.delete(product)
    db.session.commit()
    return jsonify({"success": True})


# =====================================================
# ORDERS — lógica de etiquetas automáticas
# =====================================================

def _ensure_tag(name, color='green'):
    """Obtiene o crea una tag del sistema."""
    tag = Tag.query.filter_by(name=name).first()
    if not tag:
        tag = Tag(name=name, color=color, is_system=True, is_active=True)
        db.session.add(tag)
        db.session.flush()
    return tag


def _record_tag_history(contact_id, tag, action, source, created_by=None):
    """Registra un evento de etiqueta en el historial."""
    try:
        from models import ContactTagHistory
        entry = ContactTagHistory(
            contact_id=contact_id,
            tag_id=tag.id,
            tag_name_snapshot=tag.name,
            action=action,
            source=source,
            created_by=created_by or source
        )
        db.session.add(entry)
    except Exception as e:
        logger.warning(f"No se pudo registrar historial de etiqueta: {e}")


def _apply_order_tags(contact_id):
    """Agrega 'Con pedido' y 'Comprador' al contacto."""
    contact = Contact.query.get(contact_id)
    if not contact:
        return
    tag_con_pedido = _ensure_tag('Con pedido', 'yellow')
    tag_comprador = _ensure_tag('Comprador', 'blue')
    existing = {t.id for t in contact.tags}
    if tag_con_pedido.id not in existing:
        contact.tags.append(tag_con_pedido)
        _record_tag_history(contact.id, tag_con_pedido, 'added', 'system')
    if tag_comprador.id not in existing:
        contact.tags.append(tag_comprador)
        _record_tag_history(contact.id, tag_comprador, 'added', 'system')
    db.session.commit()


def _maybe_remove_con_pedido(contact_id):
    """
    Quita 'Con pedido' si el contacto no tiene órdenes activas restantes.
    """
    from models import Order, ACTIVE_ORDER_STATUSES
    contact = Contact.query.get(contact_id)
    if not contact:
        return
    active_count = Order.query.filter(
        Order.contact_id == contact_id,
        Order.status.in_(ACTIVE_ORDER_STATUSES)
    ).count()
    if active_count == 0:
        tag = Tag.query.filter_by(name='Con pedido').first()
        if tag and tag in contact.tags:
            contact.tags.remove(tag)
            _record_tag_history(contact.id, tag, 'removed', 'system')
            db.session.commit()


# =====================================================
# ORDERS — endpoints
# =====================================================

@app.route("/orders")
def orders_page():
    if not g.current_user or not g.current_user.has_permission('orders'):
        return redirect(url_for('login'))
    return render_template('orders.html')


@app.route("/api/orders", methods=["GET"])
def api_orders_list():
    from models import Order
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403

    status = request.args.get("status")
    payment_status = request.args.get("payment_status")
    payment_method = request.args.get("payment_method")
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")
    date_type = request.args.get("date_type", "created")  # "created" | "delivery"
    search = request.args.get("search", "").strip()
    vista = request.args.get("vista")  # "unseen" | "seen" | "" = todas
    phone = request.args.get("phone", "").strip()
    page = int(request.args.get("page", 1))
    per_page = int(request.args.get("per_page", 30))

    q = Order.query.order_by(Order.created_at.desc())

    if phone:
        q = q.filter(Order.phone_number == phone)

    if status:
        q = q.filter(Order.status == status)
    if payment_status:
        q = q.filter(Order.payment_status == payment_status)
    if payment_method:
        q = q.filter(Order.payment_method == payment_method)
    if date_from:
        try:
            from datetime import date as _date
            if date_type == "delivery":
                q = q.filter(Order.delivery_date >= _date.fromisoformat(date_from))
            else:
                q = q.filter(Order.created_at >= datetime.fromisoformat(date_from))
        except ValueError:
            pass
    if date_to:
        try:
            from datetime import date as _date
            if date_type == "delivery":
                q = q.filter(Order.delivery_date <= _date.fromisoformat(date_to))
            else:
                q = q.filter(Order.created_at <= datetime.fromisoformat(date_to))
        except ValueError:
            pass
    if search:
        like = f"%{search}%"
        q = q.join(Contact, Order.contact_id == Contact.id, isouter=True).filter(
            or_(Contact.name.ilike(like), Order.phone_number.ilike(like))
        )
    if vista == "unseen":
        q = q.filter(Order.seen_at.is_(None))
    elif vista == "seen":
        q = q.filter(Order.seen_at.isnot(None))

    total = q.count()
    orders = q.offset((page - 1) * per_page).limit(per_page).all()
    return jsonify({
        "orders": [o.to_dict() for o in orders],
        "total": total,
        "page": page,
        "per_page": per_page,
    })


@app.route("/api/orders/export", methods=["GET"])
def api_orders_export():
    """Exporta órdenes filtradas al formato Excel de rutas."""
    import io, openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from models import Order, Contact

    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403

    status = request.args.get("status")
    payment_status = request.args.get("payment_status")
    payment_method = request.args.get("payment_method")
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")
    date_type = request.args.get("date_type", "created")
    search = request.args.get("search", "").strip()
    vista = request.args.get("vista")

    q = Order.query.order_by(Order.created_at.desc())

    if status:
        q = q.filter(Order.status == status)
    if payment_status:
        q = q.filter(Order.payment_status == payment_status)
    if payment_method:
        q = q.filter(Order.payment_method == payment_method)
    if date_from:
        try:
            from datetime import date as _date
            if date_type == "delivery":
                q = q.filter(Order.delivery_date >= _date.fromisoformat(date_from))
            else:
                q = q.filter(Order.created_at >= datetime.fromisoformat(date_from))
        except ValueError:
            pass
    if date_to:
        try:
            from datetime import date as _date
            if date_type == "delivery":
                q = q.filter(Order.delivery_date <= _date.fromisoformat(date_to))
            else:
                q = q.filter(Order.created_at <= datetime.fromisoformat(date_to))
        except ValueError:
            pass
    if search:
        like = f"%{search}%"
        q = q.join(Contact, Order.contact_id == Contact.id, isouter=True).filter(
            or_(Contact.name.ilike(like), Order.phone_number.ilike(like))
        )
    if vista == "unseen":
        q = q.filter(Order.seen_at.is_(None))
    elif vista == "seen":
        q = q.filter(Order.seen_at.isnot(None))

    orders = q.all()

    # Crear workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Órdenes"

    columns = [
        "Address Line 1",
        "Address Line 2",
        "City",
        "State",
        "Postal Code",
        "Extra info (Optional)",
    ]

    # Header con estilo
    header_fill = PatternFill(start_color="13EC25", end_color="13EC25", fill_type="solid")
    header_font = Font(bold=True, color="000000")
    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # Datos
    for row_idx, o in enumerate(orders, 2):
        extra_parts = []
        if o.recipient_name:
            extra_parts.append(f"Persona que recibe: {o.recipient_name}")
        if o.recipient_phone:
            extra_parts.append(f"Celular: {o.recipient_phone}")
        if o.notes:
            extra_parts.append(o.notes)
        extra_info = " | ".join(extra_parts)

        row_data = [
            o.plus_code or "",      # Address Line 1
            o.address or "",        # Address Line 2
            o.city or "",           # City
            o.province or "",       # State
            o.postal_code or "",    # Postal Code
            extra_info,             # Extra info
        ]
        for col_idx, value in enumerate(row_data, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)

    # Ajustar anchos
    col_widths = [20, 35, 20, 20, 15, 50]
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    from flask import send_file
    filename = f"ordenes_{date_from or 'all'}_{date_to or 'all'}.xlsx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/api/orders", methods=["POST"])
def api_orders_create():
    from models import Order, OrderItem, CatalogProduct
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403

    data = request.json or {}
    phone_number = data.get("phone_number", "").strip()
    if not phone_number:
        return jsonify({"error": "phone_number requerido"}), 400

    contact = Contact.query.filter_by(phone_number=phone_number).first()

    from datetime import date as _date
    def _parse_date(v):
        try: return _date.fromisoformat(v) if v else None
        except: return None

    order = Order(
        contact_id=contact.id if contact else None,
        phone_number=phone_number,
        source="manual",
        status=data.get("status", "pendiente"),
        payment_status=data.get("payment_status", "sin_pagar"),
        payment_method=data.get("payment_method"),
        currency=data.get("currency", "ARS"),
        shipping_address=data.get("shipping_address"),
        notes=data.get("notes"),
        delivery_date=_parse_date(data.get("delivery_date")),
        delivery_time=data.get("delivery_time") or None,
        earliest_arrival_time=data.get("earliest_arrival_time") or None,
        latest_arrival_time=data.get("latest_arrival_time") or None,
        recipient_name=data.get("recipient_name") or None,
        recipient_phone=data.get("recipient_phone") or None,
        latitude=float(data["latitude"]) if data.get("latitude") not in (None, "") else None,
        longitude=float(data["longitude"]) if data.get("longitude") not in (None, "") else None,
        plus_code=data.get("plus_code") or None,
        address=data.get("address") or None,
        city=data.get("city") or None,
        province=data.get("province") or None,
        postal_code=data.get("postal_code") or None,
        created_by_id=g.current_user.id,
        last_edited_by_id=g.current_user.id,
    )
    db.session.add(order)
    db.session.flush()

    # Items
    items_data = data.get("items", [])
    total = 0
    for item_data in items_data:
        rid = item_data.get("retailer_id", "")
        product = CatalogProduct.query.get(rid)
        unit_price = float(item_data.get("unit_price", product.price if product else 0) or 0)
        qty = int(item_data.get("quantity", 1))
        item = OrderItem(
            order_id=order.id,
            retailer_id=rid,
            product_name=item_data.get("product_name") or (product.name if product else rid),
            quantity=qty,
            unit_price=unit_price,
            currency=item_data.get("currency", order.currency),
        )
        db.session.add(item)
        total += unit_price * qty

    # Total: usar el provisto o calcular de items
    if data.get("total") is not None:
        order.total = float(data["total"])
    elif items_data:
        order.total = total

    db.session.commit()

    if contact:
        _apply_order_tags(contact.id)

    logger.info(f"🛍️ Orden manual {order.order_number} creada por {g.current_user.username}")
    return jsonify({"success": True, "order": order.to_dict()}), 201


@app.route("/api/orders/<int:order_id>", methods=["GET"])
def api_orders_get(order_id):
    from models import Order
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403
    order = Order.query.get_or_404(order_id)
    return jsonify(order.to_dict())


@app.route("/api/orders/<int:order_id>", methods=["PUT"])
def api_orders_update(order_id):
    from models import Order, OrderItem, CatalogProduct, ACTIVE_ORDER_STATUSES
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403
    order = Order.query.get_or_404(order_id)
    data = request.json or {}

    old_status = order.status

    for field in ("status", "payment_status", "payment_method", "currency",
                  "shipping_address", "notes", "delivery_time",
                  "earliest_arrival_time", "latest_arrival_time",
                  "recipient_name", "recipient_phone",
                  "plus_code", "address", "city", "province", "postal_code"):
        if field in data:
            setattr(order, field, data[field] or None)

    for field in ("latitude", "longitude"):
        if field in data:
            try:
                setattr(order, field, float(data[field]) if data[field] not in (None, "") else None)
            except (ValueError, TypeError):
                pass

    if "delivery_date" in data:
        try:
            from datetime import date as _date
            order.delivery_date = _date.fromisoformat(data["delivery_date"]) if data["delivery_date"] else None
        except Exception:
            pass

    if "total" in data:
        order.total = float(data["total"]) if data["total"] is not None else None

    # Recalcular items si vienen
    if "items" in data:
        OrderItem.query.filter_by(order_id=order.id).delete()
        total = 0
        for item_data in data["items"]:
            rid = item_data.get("retailer_id", "")
            product = CatalogProduct.query.get(rid)
            unit_price = float(item_data.get("unit_price", product.price if product else 0) or 0)
            qty = int(item_data.get("quantity", 1))
            item = OrderItem(
                order_id=order.id,
                retailer_id=rid,
                product_name=item_data.get("product_name") or (product.name if product else rid),
                quantity=qty,
                unit_price=unit_price,
                currency=item_data.get("currency", order.currency),
            )
            db.session.add(item)
            total += unit_price * qty
        # Solo auto-calcular si el usuario no mandó total explícito
        if "total" not in data:
            order.total = total

    order.last_edited_by_id = g.current_user.id
    order.updated_at = datetime.utcnow()
    db.session.commit()

    # Evaluar etiquetas si cambió el estado
    if order.contact_id and old_status != order.status:
        if order.status in ACTIVE_ORDER_STATUSES:
            _apply_order_tags(order.contact_id)
        else:
            _maybe_remove_con_pedido(order.contact_id)

    return jsonify({"success": True, "order": order.to_dict()})


@app.route("/api/orders/<int:order_id>/terminate", methods=["POST"])
def api_orders_terminate(order_id):
    from models import Order
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403
    order = Order.query.get_or_404(order_id)
    order.status = "terminado"
    order.terminated_at = datetime.utcnow()
    order.terminated_by_id = g.current_user.id
    order.last_edited_by_id = g.current_user.id
    order.updated_at = datetime.utcnow()
    db.session.commit()

    if order.contact_id:
        _maybe_remove_con_pedido(order.contact_id)

    return jsonify({"success": True, "order": order.to_dict()})


@app.route("/api/orders/<int:order_id>/seen", methods=["POST"])
def api_orders_seen(order_id):
    from models import Order
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"error": "Sin permiso"}), 403
    order = Order.query.get_or_404(order_id)
    if not order.seen_at:
        order.seen_at = datetime.utcnow()
        order.seen_by_id = g.current_user.id
        db.session.commit()
    return jsonify({"success": True})


@app.route("/api/orders/unseen-count", methods=["GET"])
def api_orders_unseen_count():
    from models import Order
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"count": 0})
    count = Order.query.filter(Order.seen_at.is_(None)).count()
    return jsonify({"count": count})


@app.route("/api/orders/latest-unseen-id", methods=["GET"])
def api_orders_latest_unseen_id():
    """Para polling de toast: devuelve el id de la última orden de WA no vista."""
    from models import Order
    if not g.current_user or not g.current_user.has_permission('orders'):
        return jsonify({"id": None})
    order = Order.query.filter(
        Order.seen_at.is_(None),
        Order.source == 'whatsapp'
    ).order_by(Order.created_at.desc()).first()
    return jsonify({"id": order.id if order else None,
                    "order_number": order.order_number if order else None,
                    "contact_name": (order.contact.name if order and order.contact else order.phone_number) if order else None,
                    "items_summary": ", ".join(f"{i.product_name} x{i.quantity}" for i in (order.items[:2] if order else [])),
                    "total": float(order.total) if order and order.total else None})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=Config.PORT, debug=False)
